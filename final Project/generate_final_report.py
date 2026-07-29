"""Generate the NTB Event Management System Final Defense Report (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRATCH = "/tmp/claude-1000/-home-notcool-Desktop-8th-sem/b91906ac-a674-4be3-a3d1-f14fd37b493d/scratchpad"
DIAGRAMS = f"{SCRATCH}/diagrams"
SCREENSHOTS = f"{SCRATCH}/screenshots"

SUPERVISOR_NAME = "[Supervisor's Name]"
COLLEGE_NAME = "[College Name]"
REG_NO = "[T.U. Registration Number]"
SUBMISSION_DATE = "August 2026"

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.right_margin = Inches(1)
section.left_margin = Inches(1.25)

# ── Base style setup (Times New Roman everywhere, incl. built-in Heading styles) ──
base = doc.styles["Normal"]
base.font.name = "Times New Roman"
base.font.size = Pt(12)
base.paragraph_format.line_spacing = 1.5
rpr = base.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Times New Roman")

heading_sizes = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 13}
for name, size in heading_sizes.items():
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.space_before = Pt(18 if name == "Heading 1" else 12)
    st.paragraph_format.space_after = Pt(8 if name == "Heading 1" else 6)

cap_style = doc.styles["Caption"]
cap_style.font.name = "Times New Roman"
cap_style.font.size = Pt(11)
cap_style.font.italic = True
cap_style.font.color.rgb = RGBColor(0, 0, 0)
cap_style.paragraph_format.space_before = Pt(4)
cap_style.paragraph_format.space_after = Pt(14)

# ── Field / caption helpers ──────────────────────────────────────────────────
def _fldchar(run, kind):
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    run._r.append(fld)

def _instr(run, instr_text):
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

def add_field(paragraph, instr_text, placeholder="Right-click and select 'Update Field' to generate.",
              size=12, bold=False, italic=False):
    _fldchar(paragraph.add_run(), "begin")
    _instr(paragraph.add_run(), instr_text)
    _fldchar(paragraph.add_run(), "separate")
    result_run = paragraph.add_run(placeholder)
    set_run_font(result_run, size=size, bold=bold, italic=italic)
    _fldchar(paragraph.add_run(), "end")

_fig_n = [0]
_tbl_n = [0]

def add_seq_caption(kind, text):
    """kind: 'Figure' or 'Table'. Inserts an auto-numbering SEQ field + label."""
    counter = _fig_n if kind == "Figure" else _tbl_n
    counter[0] += 1
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead_run = p.add_run(f"{kind} ")
    set_run_font(lead_run, size=11, italic=True)
    _fldchar(p.add_run(), "begin")
    _instr(p.add_run(), f"SEQ {kind} \\* ARABIC")
    _fldchar(p.add_run(), "separate")
    num_run = p.add_run(str(counter[0]))
    set_run_font(num_run, size=11, italic=True)
    _fldchar(p.add_run(), "end")
    tail_run = p.add_run(f": {text}")
    set_run_font(tail_run, size=11, italic=True)
    return p

# ── Style helpers ────────────────────────────────────────────────────────────
def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False,
         space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=heading_sizes.get(f"Heading {level}", 12), bold=True)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p

def page_break():
    doc.add_page_break()

def add_figure(image_path, caption_text, width_inches=5.9, placeholder_text=None):
    if image_path and os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    else:
        box = doc.add_table(rows=1, cols=1)
        box.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = box.rows[0].cells[0]
        cell.width = Inches(width_inches)
        para_in_cell = cell.paragraphs[0]
        para_in_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para_in_cell.add_run(placeholder_text or "[Screenshot placeholder]")
        set_run_font(r, size=11, italic=True, color=(120, 120, 120))
        for row in box.rows:
            row.height = Inches(2.2)
    add_seq_caption("Figure", caption_text)

def find_screenshot(*keywords):
    if not os.path.isdir(SCREENSHOTS):
        return None
    files = sorted(os.listdir(SCREENSHOTS))
    for f in files:
        low = f.lower()
        if low.startswith("tmp"):
            continue
        if all(k.lower() in low for k in keywords):
            return os.path.join(SCREENSHOTS, f)
    return None

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, headers):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cell.paragraphs[0].runs[0], size=11, bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for cell, txt in zip(row, row_data):
            cell.text = str(txt)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if cell.paragraphs[0].runs:
                set_run_font(cell.paragraphs[0].runs[0], size=11)
    return table

def set_page_number_format(sect, fmt, start=None):
    sectPr = sect._sectPr
    existing = sectPr.find(qn("w:pgNumType"))
    if existing is not None:
        sectPr.remove(existing)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)

def set_footer_page_number(footer, align=WD_ALIGN_PARAGRAPH.CENTER):
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = align
    add_field(p, "PAGE")

# Front matter (cover through Abbreviations): lower-roman page numbers, cover page unnumbered.
section.different_first_page_header_footer = True
set_footer_page_number(section.footer)
set_page_number_format(section, "lowerRoman", start=1)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
def center_run(text, size=12, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p

center_run("TRIBHUVAN UNIVERSITY", size=16, bold=True, space_after=6)
center_run("Faculty of Humanities and Social Sciences", size=13)
center_run(COLLEGE_NAME, size=13)
doc.add_paragraph()
center_run("A PROJECT REPORT", size=14, bold=True)
center_run("Final Defense Report", size=13, italic=True)
center_run("Department of Computer Application", size=13)
center_run(COLLEGE_NAME, size=13)
doc.add_paragraph()
para(
    "Submitted in partial fulfillment of the requirements for the Bachelor in Computer Application",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
doc.add_paragraph()
center_run("NTB EVENT MANAGEMENT SYSTEM", size=18, bold=True)
center_run("A Smart Event Management Platform for Nepal Tourism Board", size=13, italic=True)
doc.add_paragraph()
doc.add_paragraph()
center_run("Submitted By:", size=12, bold=True)
center_run("Anjal Joshi", size=12)
center_run(f"T.U. Registration Number: {REG_NO}", size=12)
center_run("anjaljoshi6@gmail.com", size=12)
doc.add_paragraph()
center_run("Under the Supervision of", size=12)
center_run(SUPERVISOR_NAME, size=12, bold=True)
doc.add_paragraph()
center_run(SUBMISSION_DATE, size=12)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SUPERVISOR'S RECOMMENDATION
# ════════════════════════════════════════════════════════════════════════════
center_run("TRIBHUVAN UNIVERSITY", size=13)
center_run("Faculty of Humanities and Social Sciences", size=12)
center_run(COLLEGE_NAME, size=12)
doc.add_paragraph()
heading("SUPERVISOR'S RECOMMENDATION", 1)
para(
    f"I hereby recommend that this project report prepared under my supervision by Anjal Joshi "
    f"(T.U. Registration Number: {REG_NO}) entitled “NTB Event Management System” in partial "
    f"fulfillment of the requirements for the degree of Bachelor in Computer Application "
    f"is recommended for the final evaluation."
)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("……………………………………………")
set_run_font(run, size=12)
center_run("", size=4, space_after=0)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
run2 = p2.add_run("SIGNATURE")
set_run_font(run2, size=13, bold=True)
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
run3 = p3.add_run(SUPERVISOR_NAME)
set_run_font(run3, size=12)
p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
run4 = p4.add_run("SUPERVISOR")
set_run_font(run4, size=13, bold=True)
p5 = doc.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
run5 = p5.add_run(COLLEGE_NAME)
set_run_font(run5, size=12)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# LETTER OF APPROVAL
# ════════════════════════════════════════════════════════════════════════════
center_run("TRIBHUVAN UNIVERSITY", size=13)
center_run("Faculty of Humanities and Social Sciences", size=12)
center_run(COLLEGE_NAME, size=12)
doc.add_paragraph()
heading("LETTER OF APPROVAL", 1)
para(
    f"This is to certify that this project report prepared by Anjal Joshi (T.U. Registration "
    f"Number: {REG_NO}) entitled “NTB Event Management System” in partial fulfillment of the "
    f"requirements for the degree of Bachelor in Computer Application has been evaluated. "
    f"In our opinion, it is satisfactory in the scope and quality of a project for the required degree."
)
doc.add_paragraph()
doc.add_paragraph()

for role in ["Project Supervisor", "Internal Examiner", "External Examiner", "Program Coordinator"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("……………………………………………")
    set_run_font(run, size=12)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(role)
    set_run_font(run2, size=12, bold=True)
    doc.add_paragraph()

page_break()

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1)
para(
    "The NTB Event Management System is a complete, production-grade web platform built for the "
    "Nepal Tourism Board (NTB) to manage, publish, and promote events across Nepal. The system "
    "solves the lack of a centralized, intelligent platform where tourism-related festivals, "
    "meetings, holidays, and cultural events can be efficiently managed by administrators and "
    "discovered by the public. Beyond standard CRUD operations, the system incorporates four "
    "custom-built AI modules implemented natively in C#, without relying on any third-party AI "
    "service: a content-based recommendation engine (TF-IDF weighted vectors and Cosine "
    "Similarity), a smart search module (BM25 probabilistic ranking), an automated tag suggestion "
    "service (TF-IDF keyword extraction with Levenshtein fuzzy matching against existing tags), "
    "and a weighted popularity scoring model for event ranking. The backend is built using "
    "ASP.NET Core 8 following Clean Architecture principles with PostgreSQL as the database and a "
    "hybrid Entity Framework Core / Dapper data-access layer, while the frontend is built using "
    "SvelteKit and TypeScript acting as a backend-for-frontend. The platform also supports "
    "role-based access control (SuperAdmin, Admin, Client), JWT authentication with refresh-token "
    "rotation, a dual AD/BS calendar, and a full guest-invitation and QR-code check-in workflow. "
    "By the final defense, all four AI modules, the invitation/check-in feature, and a suite of "
    "24 automated unit and system tests (xUnit + Moq) have been implemented, integrated end-to-end, "
    "and verified. All 24 tests pass, and the AI-powered search, recommendation, and tag-suggestion "
    "features have been confirmed against real seeded data."
)
para(
    "Keywords: Event Management System, TF-IDF, Cosine Similarity, BM25, Content-Based "
    "Recommendation, Clean Architecture, ASP.NET Core, SvelteKit, PostgreSQL, JWT Authentication, "
    "QR Check-in."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════════
heading("Acknowledgement", 1)
para(
    f"I would like to express my sincere gratitude to my supervisor, {SUPERVISOR_NAME}, for the "
    f"continuous guidance, feedback, and support provided throughout the development of this "
    f"project. Their direction was instrumental in shaping both the technical depth and the "
    f"presentation of this work."
)
para(
    f"I would also like to thank the Department of Computer Application, {COLLEGE_NAME}, and all "
    "the faculty members who provided valuable feedback during the mid-term defense, which "
    "directly shaped the scope and quality of the AI modules and testing work completed for this "
    "final submission. Finally, I extend my thanks to my friends and family for their continuous "
    "encouragement throughout this project."
)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Anjal Joshi")
set_run_font(run, size=12, bold=True)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (real Word field; press Ctrl+A then F9 in Word/LibreOffice
# to populate/refresh once the document is finalized)
# ════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════════
heading("List of Figures", 1)
lof_p = doc.add_paragraph()
add_field(lof_p, 'TOC \\h \\z \\c "Figure"')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ════════════════════════════════════════════════════════════════════════════
heading("List of Tables", 1)
lot_p = doc.add_paragraph()
add_field(lot_p, 'TOC \\h \\z \\c "Table"')

page_break()

# ════════════════════════════════════════════════════════════════════════════
# ABBREVIATIONS
# ════════════════════════════════════════════════════════════════════════════
heading("Abbreviations", 1)
abbr_rows = [
    ("AD", "Anno Domini (English Calendar)"),
    ("API", "Application Programming Interface"),
    ("BFF", "Backend-for-Frontend"),
    ("BM25", "Best Match 25 (probabilistic ranking function)"),
    ("BS", "Bikram Sambat (Nepali Calendar)"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DTO", "Data Transfer Object"),
    ("EF Core", "Entity Framework Core"),
    ("EMS", "Event Management System"),
    ("ERD", "Entity Relationship Diagram"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("JWT", "JSON Web Token"),
    ("NTB", "Nepal Tourism Board"),
    ("ORM", "Object-Relational Mapper"),
    ("QR", "Quick Response (code)"),
    ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"),
    ("SPA", "Single Page Application"),
    ("SQL", "Structured Query Language"),
    ("SSR", "Server-Side Rendering"),
    ("TF-IDF", "Term Frequency-Inverse Document Frequency"),
    ("UI/UX", "User Interface / User Experience"),
]
styled_table(["Abbreviation", "Full Form"], abbr_rows)

# Body (Chapter 1 onward): new section, decimal page numbers restarting at 1.
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
body_section.different_first_page_header_footer = False
body_section.page_height = Cm(29.7)
body_section.page_width = Cm(21.0)
body_section.top_margin = Inches(1)
body_section.bottom_margin = Inches(1)
body_section.right_margin = Inches(1)
body_section.left_margin = Inches(1.25)
set_footer_page_number(body_section.footer)
set_page_number_format(body_section, "decimal", start=1)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 1: Introduction", 1)

heading("1.1 Introduction", 2)
para(
    "Nepal is a country of extraordinary cultural richness and natural diversity. Every year, "
    "hundreds of festivals, tourism events, national holidays, and community gatherings are "
    "organized across the country. The Nepal Tourism Board (NTB), as the official government "
    "body responsible for promoting Nepal as a tourism destination, plays a central role in "
    "organizing and publicizing these events to both domestic and international audiences."
)
para(
    "However, managing and publishing this large volume of events through disconnected or manual "
    "processes creates significant operational inefficiencies. There is no unified, modern platform "
    "that brings event creation, publication, calendar management, and intelligent event discovery "
    "under one roof. The NTB Event Management System is built to solve exactly this problem."
)
para(
    "This system is a full-stack web application where authorized administrators can create, "
    "update, and manage events of various types (festivals, meetings, holidays, and general "
    "events) and invite guests to them with QR-code based check-in. The public can browse, "
    "filter, and explore these events through an interactive calendar and listing interface. "
    "What sets this platform apart is the integration of self-built AI modules that enhance "
    "search quality, surface relevant events to users, and assist administrators with intelligent "
    "tag suggestions during event creation. By the final defense, every one of these modules has "
    "been implemented, integrated end-to-end, and covered by an automated test suite."
)

heading("1.2 Problem Statement", 2)
para(
    "Nepal's tourism event landscape suffers from fragmentation and lack of digital structure. "
    "The core problems that motivated this project are:"
)
bullet("No centralized digital platform exists for NTB to manage all event types from a single admin interface.")
bullet("The public has no reliable, filterable, and searchable directory of NTB-published events.")
bullet("Event discovery is passive, since users must scroll through lists with no intelligence behind what is surfaced.")
bullet("Manual content tagging is inconsistent, making category-based filtering less reliable.")
bullet("There is no way to track event popularity or surface trending or high-impact events automatically.")
bullet("Inviting and verifying guests at event venues is typically a manual, paper-based process prone to errors.")
para(
    "These gaps result in missed opportunities for tourism promotion, lower public awareness, and "
    "high administrative overhead. The NTB Event Management System directly addresses each of these "
    "problems through structured event management workflows, embedded AI-driven discovery features, "
    "and a digital guest invitation/check-in flow."
)

heading("1.3 Objectives", 2)
para("The primary objectives of this project are as follows:")
numbered("To design and develop a complete, production-grade event management platform for the Nepal Tourism Board.")
numbered("To implement a role-based access control system that supports Super Admin, Admin, and Client roles with appropriate permissions.")
numbered("To build an interactive calendar interface that displays events in both month-view and list-view with dual-calendar support (AD and BS).")
numbered("To implement a custom-built Content-Based Recommendation Engine using TF-IDF and Cosine Similarity that suggests relevant events without relying on third-party AI APIs.")
numbered("To develop a smart search module powered by the BM25 ranking algorithm, replacing naive keyword matching with a probabilistic relevance model.")
numbered("To create an automated tag suggestion feature using keyword extraction so administrators receive intelligent recommendations while creating events.")
numbered("To implement a Weighted Popularity Scoring model that ranks events based on rating, attendance, recency, and featured status.")
numbered("To implement a guest invitation and QR-code based check-in system for event attendance verification.")
numbered("To ensure the platform follows Clean Architecture principles on the backend, ensuring maintainability and testability, and to validate correctness through automated unit and system testing.")

heading("1.4 Scope and Limitation", 2)
para("The scope of the NTB Event Management System covers the following areas:")
bullet("Full event lifecycle management: draft, publish, and archive.")
bullet("Multiple event types: Festival, Meeting, Holiday, and General Event.")
bullet("Image upload and media management for events.")
bullet("Category and tag management with support for many-to-many relationships.")
bullet("User management with role-based access: Super Admin, Admin, and Client.")
bullet("JWT-based authentication with refresh token rotation.")
bullet("Public-facing event discovery with calendar view, filtering, and smart search.")
bullet("AI-assisted content recommendation, search ranking, and tag suggestion.")
bullet("Support for both AD (English calendar) and BS (Bikram Sambat) date formats.")
bullet("Guest invitation via email with QR code generation, and door-side scan/verify check-in.")

para("The current limitations of the system include:")
bullet("The recommendation engine is content-based; collaborative filtering (based on user behaviour history) is planned for future phases.")
bullet("The system does not currently support event ticketing or payment processing.")
bullet("Mobile native applications (iOS/Android) are outside the current scope; the platform is optimized as a responsive web application.")
bullet("The AI modules currently run at request time; background pre-computation and caching are planned as a future enhancement.")

heading("1.5 Development Methodology", 2)
para(
    "This project follows an Agile-Incremental development methodology. Rather than planning and "
    "building everything at once, the system has been developed in iterative sprints, with each "
    "sprint delivering a working vertical slice of functionality that was reviewed and refined "
    "based on supervisor feedback before the next sprint began."
)
add_figure(f"{DIAGRAMS}/fig_agile.png", "Agile (Scrum-based) Development Methodology", width_inches=5.2)
para("The development was divided into the following major increments:")
bullet("Increment 1 (Core Infrastructure): Database schema design, Clean Architecture project setup, JWT authentication, user roles, and basic event CRUD.")
bullet("Increment 2 (Frontend Foundation): SvelteKit project setup, admin panel layout, event creation/editing forms, and image upload.")
bullet("Increment 3 (Calendar and Public View): Interactive calendar with month/list toggle, public event discovery page, filtering, and category/region-based navigation.")
bullet("Increment 4 (AI Modules): TF-IDF and Cosine Similarity based recommendation engine, BM25 search ranking, keyword extraction with Levenshtein-based tag suggestion, and weighted popularity scoring, all implemented and integrated end-to-end.")
bullet("Increment 5 (Guest Invitations and QR Check-in): Invitation/guest data model, QR code generation, email delivery, and the admin door-side scan-and-verify flow.")
bullet("Increment 6 (Testing and Documentation): A 24-case automated test suite (xUnit + Moq) covering authentication, event CRUD, and all four AI modules, followed by result analysis and final report writing.")
para("The project's schedule, spanning both the mid-term and final defense milestones, is summarized in the Gantt chart below.")
add_figure(f"{DIAGRAMS}/fig_gantt.png", "Project Gantt Chart", width_inches=5.9)
para(
    "This incremental approach allowed continuous delivery of working functionality while enabling "
    "course corrections based on testing and supervisor feedback throughout the project lifecycle, "
    "most notably, the AI modules and the automated test suite, which were the primary gap closed "
    "between the mid-term and final defense."
)

heading("1.6 Report Organization", 2)
para("This report is organized into five chapters:")
bullet("Chapter 1 (Introduction): Provides the background, problem statement, objectives, scope, and development methodology.")
bullet("Chapter 2 (Background Study and Literature Review): Covers the fundamental concepts and reviews similar systems and research work.")
bullet("Chapter 3 (System Analysis and Design): Presents the requirement analysis, feasibility study, system diagrams, interface design, and algorithm details.")
bullet("Chapter 4 (Implementation and Testing): Details the tools used, module implementations, test cases, and result analysis.")
bullet("Chapter 5 (Conclusion and Future Recommendations): Summarizes the completed work and outlines the future work planned.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – BACKGROUND STUDY AND LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 2: Background Study and Literature Review", 1)

heading("2.1 Background Study", 2)
para(
    "This chapter provides the theoretical foundation behind the key concepts used in the "
    "NTB Event Management System, with a particular focus on the AI and algorithmic components."
)

heading("2.1.1 Event Management Systems", 3)
para(
    "An Event Management System (EMS) is a software platform that facilitates the planning, "
    "organization, publication, and administration of events. Modern EMS platforms go beyond "
    "simple CRUD operations by integrating features such as calendar visualization, multi-role "
    "user access control, media management, and intelligent content discovery. In the context "
    "of a national tourism body like NTB, an EMS must be robust, scalable, and capable of "
    "serving both internal administrators and the general public simultaneously."
)

heading("2.1.2 Content-Based Filtering and TF-IDF", 3)
para(
    "Content-based filtering is a recommendation technique that suggests items similar to those "
    "a user has shown interest in, based on the properties of the items themselves rather than "
    "the behaviour of other users. In the context of event recommendation, if a user views a "
    "Dashain Festival event, the system should surface other festival events with similar topics, "
    "locations, or keywords."
)
para(
    "Term Frequency-Inverse Document Frequency (TF-IDF) is a numerical statistic that reflects "
    "how important a word is to a document within a collection. The TF (Term Frequency) component "
    "measures how often a word appears in a document, while the IDF (Inverse Document Frequency) "
    "component reduces the weight of words that appear in many documents (common words) and "
    "increases the weight of rare, domain-specific words. Together, they produce a vector "
    "representation of a document that captures its meaningful content."
)

heading("2.1.3 Cosine Similarity", 3)
para(
    "Cosine Similarity is a metric used to measure the similarity between two vectors by "
    "computing the cosine of the angle between them in a multi-dimensional space. When two "
    "TF-IDF vectors are compared using cosine similarity, the result is a score between 0 and 1, "
    "where 1 means identical content and 0 means completely unrelated content. This is the "
    "mathematical core of the recommendation engine in this system."
)

heading("2.1.4 BM25 – Probabilistic Information Retrieval", 3)
para(
    "BM25 (Best Match 25), also known as Okapi BM25, is a ranking function used in information "
    "retrieval to rank documents based on their relevance to a given search query. Unlike a simple "
    "keyword match (ILIKE in SQL), BM25 accounts for the frequency of query terms in each document, "
    "the inverse document frequency of each term, and the document length, meaning longer documents "
    "are not unfairly rewarded for containing query terms more times. BM25 is the foundation of "
    "modern search engines like Elasticsearch and Lucene, and its implementation in this project "
    "provides significantly more relevant search results than plain pattern matching."
)

heading("2.1.5 Weighted Scoring Models", 3)
para(
    "A weighted scoring model is a simple but effective technique for ranking items based on "
    "multiple criteria simultaneously. Each criterion (e.g., rating, attendance, recency, featured "
    "status) is assigned a weight reflecting its relative importance. The final score is the "
    "weighted sum of all criteria. This technique is widely used in recommendation systems to "
    "surface popular or high-quality items without requiring complex machine learning infrastructure."
)

heading("2.1.6 Keyword Extraction", 3)
para(
    "Keyword extraction is the process of automatically identifying the most relevant and "
    "descriptive words or phrases from a piece of text. In this system, when an administrator "
    "is creating an event, the title and description are analysed to extract the top keywords "
    "using a TF-IDF-based approach against the existing event corpus. These keywords are then "
    "matched against the available tags in the system, using Levenshtein edit-distance for "
    "fuzzy matching, to produce intelligent tag suggestions."
)

heading("2.2 Literature Review", 2)
para(
    "Several existing event management and tourism platforms have been studied to understand the "
    "current state of the art and identify areas where the NTB system can provide better value."
)

heading("2.2.1 Eventbrite", 3)
para(
    "Eventbrite is one of the world's largest event management platforms. It supports event "
    "discovery, ticketing, and registration. Its recommendation system is primarily collaborative "
    "filtering-based, leveraging purchase history and user behaviour. While powerful, Eventbrite's "
    "approach requires large volumes of user behavioural data to function well, something that "
    "a new platform serving a national tourism body would not have at launch. The NTB system's "
    "content-based approach using TF-IDF and Cosine Similarity is more suitable for a cold-start "
    "scenario where user data is limited."
)

heading("2.2.2 Nepal's Existing Tourism Information Portals", 3)
para(
    "Nepal's existing government tourism portals (such as tourismdepartment.gov.np) primarily "
    "serve static informational content. They lack dynamic event management features, interactive "
    "calendars, admin workflows, or any form of intelligent event discovery. This confirms the "
    "gap that the NTB Event Management System is designed to fill."
)

heading("2.2.3 Research on BM25 in Web Applications", 3)
para(
    "Research by Robertson and Zaragoza (2009) established BM25 as a robust, parameter-tunable "
    "ranking algorithm for text retrieval. Subsequent work has shown that BM25 consistently "
    "outperforms simple TF-based models in precision and recall for domain-specific search tasks. "
    "Implementing BM25 natively within the application layer (rather than delegating to a "
    "full-text search engine) is a deliberate design choice in this project, demonstrating "
    "algorithm-level understanding rather than reliance on infrastructure abstractions."
)

heading("2.2.4 TF-IDF in Recommendation Systems", 3)
para(
    "Lops et al. (2011) provide a comprehensive survey of content-based filtering techniques "
    "and identify TF-IDF as a foundational and highly interpretable method for building item "
    "profiles in recommendation systems. The survey notes that TF-IDF based content-based "
    "filtering performs particularly well for text-rich domains, which aligns directly with "
    "the event management domain where events have rich titles, descriptions, categories, and tags."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – SYSTEM ANALYSIS AND DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 3: System Analysis and Design", 1)

heading("3.1 System Analysis", 2)

heading("3.1.1 Functional Requirements", 3)
para("The functional requirements of the NTB Event Management System are organized by user role:")

para("Administrator Requirements:", bold=True)
bullet("Admin can log in securely using email and password with JWT-based session management.")
bullet("Admin can create, edit, and delete events with full details including title, description, category, type, date (AD and BS), location, region, price, images, highlights, tags, and organizer information.")
bullet("Admin can change the lifecycle status of events: Draft, Published, or Archived.")
bullet("Admin can manage categories and tags, including creating and organizing tag-category associations, and can accept AI-suggested tags directly in the event form.")
bullet("Admin can manage system users and assign roles.")
bullet("Admin can view a calendar showing all events with color-coded type indicators.")
bullet("Admin can invite guests to an event by email, which generates and sends a QR code, and can scan/verify guest QR codes at check-in.")

para("Public User Requirements:", bold=True)
bullet("Public users can browse events on the homepage with featured event highlights, ranked by popularity score.")
bullet("Public users can filter events by type, category, region, and date range.")
bullet("Public users can use the smart search (BM25-ranked) to find events by title, location, or description.")
bullet("Public users can view events in both calendar and list view.")
bullet("Public users can view detailed event information including location, organizer, highlights, and images.")
bullet("Public users receive smart event recommendations (TF-IDF + Cosine Similarity) based on the event they are currently viewing.")
bullet("Invited guests can view their invitation landing page and present their QR code for check-in.")

heading("3.1.2 Non-Functional Requirements", 3)
bullet("Performance: API response time should be under 200ms for standard queries; search and recommendation results should return within 500ms.")
bullet("Security: All admin endpoints are protected by JWT authentication. Passwords are hashed using BCrypt. Refresh token rotation prevents token replay attacks.")
bullet("Scalability: The Clean Architecture design ensures business logic is decoupled from infrastructure, enabling horizontal scaling when needed.")
bullet("Usability: The frontend must be responsive and work seamlessly on both desktop and mobile devices.")
bullet("Maintainability: Clean Architecture layers (Domain, Application, Infrastructure, API) ensure clear separation of concerns.")
bullet("Data Integrity: Dual-calendar date support (AD/BS) must be consistently maintained across all event records.")
bullet("Testability: All core services (auth, event CRUD, and the four AI modules) are covered by an automated unit/system test suite.")

heading("3.1.3 Use Case Diagram", 3)
para(
    "The use case diagram below summarizes how the two main actor groups, public users/clients "
    "and administrators, interact with the system."
)
add_figure(f"{DIAGRAMS}/fig_usecase.png", "Use Case Diagram", width_inches=5.9)

heading("3.2 Feasibility Analysis", 2)

heading("3.2.1 Technical Feasibility", 3)
para(
    "The system is built on mature, well-supported technologies. ASP.NET Core 8 is a "
    "production-grade web framework maintained by Microsoft. PostgreSQL is a battle-tested, "
    "open-source relational database. SvelteKit is a modern, performant web framework. "
    "The AI algorithms (TF-IDF, Cosine Similarity, BM25) are mathematically well-defined "
    "and have been implemented in pure C# without any external AI library dependencies, "
    "confirming technical feasibility in practice, not just in theory."
)

heading("3.2.2 Operational Feasibility", 3)
para(
    "NTB administrators are familiar with web-based administration tools. The system's admin "
    "interface is designed to be intuitive and requires minimal training. The public-facing "
    "interface is simple and discoverable. From an operational standpoint, the system reduces "
    "manual event publication effort and eliminates the need for external tools or third-party "
    "content management platforms."
)

heading("3.2.3 Economic Feasibility", 3)
para(
    "The entire technology stack is open-source with no licensing fees. The system can be "
    "hosted on cloud providers (such as AWS, Azure, or DigitalOcean) at a low monthly cost. "
    "Development effort is entirely student-driven. The long-term cost savings from eliminating "
    "manual event management workflows and reducing dependence on external platforms make the "
    "system economically viable for NTB."
)

heading("3.3 System Design", 2)

heading("3.3.1 System Architecture Overview", 3)
para(
    "The NTB Event Management System follows Clean Architecture on the backend, organized into "
    "four distinct layers, with the SvelteKit frontend acting as a backend-for-frontend (BFF) "
    "that proxies requests to the API rather than exposing the access token to the browser directly."
)
add_figure(f"{DIAGRAMS}/fig_architecture.png", "Layered System Architecture", width_inches=5.9)
bullet("Domain Layer: Contains the core business entities (Event, User, Category, Tag, RefreshToken, EventHighlight, Guest, Invitation) and enumerations (EventType, EventLifecycleStatus, UserRole). This layer has no dependencies on any external framework.")
bullet("Application Layer: Contains service interfaces, repository contracts, DTOs, and business logic, including the TF-IDF vectorizer, BM25 ranker, popularity scorer, and tag suggestion service. It depends only on the Domain layer.")
bullet("Infrastructure Layer: Implements the repository contracts and services defined in the Application layer. It uses Entity Framework Core for CRUD operations, Dapper for performance-critical list queries, BCrypt for password hashing, QRCoder for QR generation, and SMTP for email delivery.")
bullet("API Layer: Hosts the ASP.NET Core controllers, handles HTTP request/response, JWT configuration, CORS, Swagger documentation, and dependency injection wiring.")

heading("3.3.2 Database Design", 3)
para("The primary entities and their relationships in the PostgreSQL database are shown below.")
add_figure(f"{DIAGRAMS}/fig_erd.png", "Database Schema Design (Entity Relationship Diagram)", width_inches=5.9)
bullet("events: Stores all event data including title, description, dates (AD and BS), location, region, type, status, images, price, rating, organizer information, and highlighted features as a JSON column.")
bullet("users: Stores registered user accounts with hashed passwords and role assignments.")
bullet("refresh_tokens: Stores refresh tokens per user session for secure token rotation.")
bullet("categories: Stores event categories managed by administrators.")
bullet("tags: Stores descriptive tags that can be associated with events and categories.")
bullet("event_tags: Junction table linking events to tags (many-to-many).")
bullet("categories_tags: Junction table linking categories to their default tags.")
bullet("tags_jn (Audit): Tracks tag association changes for audit purposes.")
bullet("invitation_guests: One row per unique invited person (normalized by email), reused across events.")
bullet("event_invitations: Event–guest association holding the invite token, status (pending/sent/verified/expired/cancelled), expiry, and check-in audit fields.")
bullet("invitation_scans: One audit row per QR scan attempt at check-in.")

heading("3.3.3 API Design", 3)
para("The backend exposes the following RESTful API endpoints:")
bullet("POST /api/auth/login: Authenticate user and receive JWT + refresh token.")
bullet("POST /api/auth/refresh: Rotate refresh token and issue a new access token.")
bullet("POST /api/auth/logout: Invalidate the current refresh token.")
bullet("GET /api/events: Get paginated, filtered list of events (public).")
bullet("GET /api/events/{id}: Get a single event by ID.")
bullet("GET /api/events/{slug}: Get a single event by slug.")
bullet("POST /api/events: Create a new event (Admin).")
bullet("PUT /api/events/{id}: Update an event (Admin).")
bullet("DELETE /api/events/{id}: Delete an event (Admin).")
bullet("GET /api/events/{id}/recommendations: Get AI-powered recommendations for an event.")
bullet("GET /api/events/search?q={query}: Smart BM25-ranked search.")
bullet("POST /api/events/suggest-tags: Get AI-suggested tags for a title and description.")
bullet("GET /api/categories: List all categories.")
bullet("GET /api/tags: List all tags, optionally filtered by category.")
bullet("POST /api/events/{eventId}/invitations: Invite a guest (generates QR + emails them).")
bullet("GET /api/events/{eventId}/invitations: List an event's invitations.")
bullet("POST /api/invitations/scan: Look up an invitation by scanned token (no consume).")
bullet("POST /api/invitations/{id}/verify: Confirm check-in and expire the QR.")
bullet("GET /api/invitations/by-token/{token}: Public guest landing data.")

heading("3.3.4 Interface Design (UI/UX)", 3)
para(
    "The following screenshots, captured from the running application, illustrate the primary "
    "public and administrative interfaces."
)

add_figure(find_screenshot("public_landing"), "Public Homepage: Featured Events", width_inches=5.9,
           placeholder_text="Public homepage screenshot")
add_figure(find_screenshot("search_festival"), "Smart Search Results (BM25-ranked)", width_inches=5.9,
           placeholder_text="Search results screenshot")
add_figure(find_screenshot("event_details_modal"), "Event Details Modal", width_inches=5.9,
           placeholder_text="Event details modal screenshot")
add_figure(find_screenshot("recommendations"), "“You Might Also Like”: AI-Powered Recommendations", width_inches=5.9,
           placeholder_text="Recommendations panel screenshot")
add_figure(find_screenshot("public_calendar"), "Public Interactive Calendar (AD/BS Toggle)", width_inches=5.9,
           placeholder_text="Public calendar screenshot")
add_figure(find_screenshot("admin_login"), "Admin Login Page", width_inches=5.5,
           placeholder_text="Admin login screenshot")
add_figure(find_screenshot("admin_dashboard"), "Admin Dashboard", width_inches=5.9,
           placeholder_text="Admin dashboard screenshot")
add_figure(find_screenshot("admin_events_list"), "Admin Events Management List", width_inches=5.9,
           placeholder_text="Admin events list screenshot")
add_figure(find_screenshot("admin_create_event"), "Admin: Create Event Form (Category, Status, Tags)", width_inches=5.9,
           placeholder_text="Admin create-event screenshot")
add_figure(find_screenshot("tag_suggestions"), "Admin Event Create: AI Tag Suggestion in Action", width_inches=5.9,
           placeholder_text="Admin create-event screenshot with tag suggestion")
add_figure(find_screenshot("admin_calendar"), "Admin Calendar Setup (AD/BS, category filters)", width_inches=5.9,
           placeholder_text="Admin calendar screenshot")

heading("3.4 Algorithm Details", 2)
para(
    "This section describes the four AI and algorithmic modules implemented natively within "
    "the NTB Event Management System. All algorithms are implemented in C# within the "
    "Application/Infrastructure layers and do not depend on any external AI API or service."
)

heading("3.4.1 Content-Based Event Recommendation Engine (TF-IDF + Cosine Similarity)", 3)
para(
    "The recommendation engine answers the question: given the event a user is currently viewing, "
    "which other events are most similar in content? This is a classic content-based filtering "
    "problem solved using TF-IDF vectorization and Cosine Similarity, implemented in "
    "RecommendationService using the shared TfIdfVectorizer utility."
)

para("Step 1: Build the Event Corpus", bold=True)
para(
    "Each event in the database is represented as a text document by concatenating its most "
    "meaningful text fields: title (weighted 3×), summary (weighted 2×), long description, "
    "category name, region, and associated tags. Title repetition amplifies its influence "
    "on the final vector; this is a standard technique in content-based systems."
)

para("Step 2: Compute TF-IDF Vectors", bold=True)
para(
    "For each event document, the system tokenizes the text (via TextTokenizer), removes stop "
    "words (common words such as 'the', 'is', 'in' which carry no meaning), and computes TF-IDF "
    "scores for each remaining token. The vocabulary is built from all events in the corpus."
)
para("The TF-IDF formula used is:")
para("    TF(t, d)  = (count of term t in document d) / (total terms in document d)")
para("    IDF(t)    = log( (1 + N) / (1 + df(t)) ) + 1")
para("    TF-IDF(t, d) = TF(t, d) × IDF(t)")
para(
    "Where N is the total number of event documents and df(t) is the number of documents "
    "that contain term t. The +1 smoothing in the IDF formula prevents zero values for "
    "terms that appear in every document."
)

para("Step 3: Compute Cosine Similarity", bold=True)
para(
    "Given the TF-IDF vector of the target event (the one being viewed) and the TF-IDF vectors "
    "of all other events, cosine similarity is computed as:"
)
para("    similarity(A, B) = (A · B) / (||A|| × ||B||)")
para(
    "Where A · B is the dot product of the two vectors and ||A|| and ||B|| are their "
    "magnitudes. This gives a score between 0 and 1 for each pair. Events are then sorted by "
    "descending similarity score, and the top-N results (default: 5) are returned as "
    "recommendations, restricted to published events only."
)

para("Why This Is Not a SQL Query:", bold=True)
para(
    "This algorithm operates entirely in the C# application layer. The database is used only to "
    "retrieve event data. The TF-IDF computation, vector construction, and similarity scoring are "
    "all done in memory using mathematical operations, not SQL filters or joins. This is verified "
    "end-to-end by the system test Recommendations_ForFestivalEvent_ReturnsSimilarCategoryEventsOnly."
)

heading("3.4.2 Smart Search with BM25 Ranking", 3)
para(
    "The smart search module (SearchRankingService) replaces the naive ILIKE pattern matching "
    "with a proper probabilistic ranking model called BM25 (Best Match 25), which is the same "
    "ranking function used by Elasticsearch, Solr, and many major search engines internally. "
    "The admin panel's existing ILIKE-based filter (BuildWhereClause) was intentionally left "
    "untouched; the BM25 search is a separate, additive public search path exposed at "
    "GET /api/events/search."
)

para("How BM25 Works:", bold=True)
para("Given a user's search query Q containing terms q1, q2, ..., qn, BM25 scores each event document D as follows:")
para("    Score(D, Q) = Σ IDF(qi) × [ f(qi, D) × (k1 + 1) ] / [ f(qi, D) + k1 × (1 - b + b × |D| / avgdl) ]")
para("Where:")
bullet("f(qi, D) is the frequency of query term qi in document D.")
bullet("|D| is the length of document D in words.")
bullet("avgdl is the average document length across all events.")
bullet("k1 is a term saturation parameter (implemented value: 1.5) that controls how much repeated terms increase the score.")
bullet("b is a length normalization parameter (implemented value: 0.75) that controls how much document length affects scoring.")

para(
    "Each event's searchable text (title + summary + location + region) is tokenized and indexed "
    "into an in-memory inverted index built by Bm25Ranker. When a search query arrives, BM25 "
    "scores are computed in O(k) time where k is the number of matching terms."
)

para("Why BM25 Over ILIKE:", bold=True)
para(
    "A plain ILIKE query in SQL returns all events that simply contain the search term, with no "
    "relevance ranking. BM25 assigns higher scores to events where the query terms appear frequently "
    "but the document is not artificially long, and where the query terms are rare across the corpus "
    "(more discriminative). This is verified by the system test "
    "SmartSearch_DashainFestivalQuery_RanksFestivalEventsHigher, confirming that festival events rank "
    "above unrelated events for a festival-specific query."
)

heading("3.4.3 Automated Tag Suggestion via Keyword Extraction", 3)
para(
    "When an administrator is in the process of creating or editing an event, TagSuggestionService "
    "automatically analyses the event's title and description to suggest relevant tags. This "
    "reduces manual tagging effort and improves consistency across the event catalog."
)

para("How It Works:", bold=True)
para("The keyword extraction module uses a corpus-aware TF-IDF scoring approach:")
numbered("The input text (title + description) is tokenized and stop words are removed.")
numbered("TF-IDF scores are computed for each token in the input text against the published-event corpus.")
numbered("The top-K tokens by TF-IDF score are selected as candidate keywords.")
numbered("Each candidate keyword is matched against the system's existing tag list using normalized string comparison and Levenshtein edit distance for fuzzy matching.")
numbered("Matched tags are returned in ranked order as suggestions to the administrator via the 'Suggest tags' button on the event form.")

para(
    "This process is entirely self-contained. It uses the event corpus already stored in the "
    "database and the tag dictionary maintained by administrators. No external NLP service or "
    "API is called. It has been verified end-to-end, including a fuzzy-match case where the typo "
    "“festivle” correctly resolved to the existing tag “Festival”."
)

heading("3.4.4 Weighted Popularity Scoring for Event Ranking", 3)
para(
    "To surface the most valuable events to users on the homepage and in category listings, "
    "PopularityScoreService computes a Popularity Score for each event using a multi-factor "
    "weighted formula. This is not a machine learning model; it is a transparent, tunable "
    "scoring function that administrators can understand and adjust."
)

para("Scoring Formula:", bold=True)
para(
    "    PopularityScore(e) = (W_r × NormalizedRating) + (W_a × NormalizedAttendance) "
    "+ (W_f × FeaturedBoost) + (W_d × RecencyScore)"
)
para("Where the weights are:")
bullet("W_r = 0.35: Rating weight (user ratings are the strongest signal of quality).")
bullet("W_a = 0.25: Attendance/engagement weight.")
bullet("W_f = 0.20: Featured boost (admin-marked events get a boost).")
bullet("W_d = 0.20: Recency score (more recent events rank higher to keep listings fresh).")

para(
    "NormalizedRating is computed as (event_rating / max_rating). Since the domain model does not "
    "carry a numeric attendance field (AttendanceLabel is free text), NormalizedAttendance uses "
    "ReviewsLabel as a practical proxy. RecencyScore uses an exponential decay function: "
    "RecencyScore = e^(−λ × days_since_event_start), where λ = 0.01 ensures that events "
    "decay gradually over ~100 days. This formula is computed entirely in C# without any database "
    "aggregation query. EventDto.PopularityScore is populated on every event read, and the public "
    "homepage's featured events section (EventsLandingSection.svelte) sorts by this score descending."
)

para("Why This Matters:", bold=True)
para(
    "Without this scoring model, events would simply be sorted by creation date or start date. "
    "The weighted model ensures that high-rated, well-attended, recently-starting festivals "
    "appear prominently, which is exactly what NTB wants: showcasing Nepal's most impactful "
    "tourism events to visitors."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – IMPLEMENTATION AND TESTING
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 4: Implementation and Testing", 1)

heading("4.1 Implementation", 2)

heading("4.1.1 Tools and Technologies Used", 3)

add_seq_caption("Table", "Tools and Technologies Used")
styled_table(
    ["Component", "Technology", "Purpose"],
    [
        ("Backend Framework", "ASP.NET Core 8", "REST API, routing, middleware"),
        ("Architecture Pattern", "Clean Architecture", "Separation of concerns, testability"),
        ("Primary ORM", "Entity Framework Core 8", "CRUD operations, migrations"),
        ("Query Micro-ORM", "Dapper", "Performance-optimised list queries"),
        ("Database", "PostgreSQL 16", "Relational data storage"),
        ("Authentication", "JWT + Refresh Tokens", "Stateless, secure session management"),
        ("Password Hashing", "BCrypt.Net", "Secure credential storage"),
        ("QR Code Generation", "QRCoder", "Guest invitation QR codes"),
        ("Email Delivery", "SMTP (Mailcow-compatible)", "Invitation emails, dev-mode logging fallback"),
        ("Frontend Framework", "SvelteKit + TypeScript", "Reactive, performant UI"),
        ("Styling", "CSS Variables + Tailwind", "Responsive design system"),
        ("API Documentation", "Swagger / OpenAPI", "Interactive API explorer"),
        ("AI Algorithms", "Custom C# Modules", "TF-IDF, BM25, Cosine Similarity, Popularity Scoring, Levenshtein"),
        ("Testing", "xUnit + Moq", "Unit and system test automation"),
        ("Version Control", "Git / GitHub", "Source code management"),
    ],
)
doc.add_paragraph()

heading("4.1.2 Module Implementation Details", 3)

para("Authentication Module:", bold=True)
para(
    "The authentication module implements JWT access tokens with a short expiry (15 minutes) "
    "and long-lived refresh tokens stored in the database. On each successful refresh, the old "
    "token is invalidated and a new one is issued (token rotation). This prevents refresh token "
    "reuse attacks. The User entity carries a UserRole enum (SuperAdmin, Admin, Client) that is "
    "embedded in the JWT claims, enabling role-based route protection across all admin endpoints."
)

para("Event Management Module:", bold=True)
para(
    "Events are created and managed through the EventsController, which delegates to the "
    "IEventService in the Application layer. The EventRepository uses Entity Framework Core for "
    "CRUD operations and Dapper for the paginated list query. The list query uses a dynamic WHERE "
    "clause builder (BuildWhereClause) that constructs parameterized SQL based on the active "
    "filters, preventing SQL injection while maintaining query flexibility. Events support dual "
    "dates (DateAd/DateBs), multiple images stored as a JSON array, and structured highlights "
    "stored as a JSON column."
)

para("Calendar Module:", bold=True)
para(
    "The frontend calendar is built as a custom SvelteKit component with no external calendar "
    "library dependency. It supports month-view and list-view toggles. Events are color-coded "
    "by type (Festival, Meeting, Holiday, Event). The calendar supports both AD and BS date "
    "display through a toggle mechanism, with the real AD↔BS conversion logic implemented in "
    "the frontend's dateUtils.ts."
)

para("Tag and Category Module:", bold=True)
para(
    "Categories and tags have a many-to-many relationship through the CategoriesTags junction. "
    "Events also relate to tags through the EventTags junction. An audit table (TagsJn) tracks "
    "changes to tag associations over time. The tag suggestion AI module (described in Chapter 3) "
    "operates on top of this data structure, and its 'Suggest tags' button was added to the admin "
    "event create/edit form together with a visible Tags field."
)

para("Guest Invitation & QR Check-in Module:", bold=True)
para(
    "This module spans all layers: the Guest, Invitation, and InvitationScan entities (Domain), "
    "InvitationService and its contracts (Application), GuestRepository, InvitationRepository, "
    "QrCodeService, and SmtpEmailService (Infrastructure), and InvitationsController (API). The "
    "flow is: an admin invites a guest → the guest receives an email with a link and QR code "
    "→ the guest presents the QR at the venue → an admin scans it for a preview → the "
    "admin confirms verification, which consumes/expires the QR so it cannot be reused. Invitation "
    "status follows the lifecycle pending → sent → verified, with expired and cancelled as "
    "terminal states."
)

para("AI Modules: Implemented and Integrated", bold=True)
para(
    "The four AI modules described in Section 3.4 have been fully implemented as services within "
    "the Application and Infrastructure layers and are wired into the API and the frontend:"
)
bullet("RecommendationService (IRecommendationService): TF-IDF vectorization and Cosine Similarity, exposed via GET /api/events/{id}/recommendations and surfaced as the “You might also like” section in EventDetailsModal.svelte.")
bullet("SearchRankingService (ISearchRankingService): BM25 inverted index and scoring, exposed via GET /api/events/search, driven by the search box in PublicNav.svelte.")
bullet("TagSuggestionService (ITagSuggestionService): TF-IDF keyword extraction with Levenshtein-based fuzzy tag matching, exposed via POST /api/events/suggest-tags and the admin event form's “Suggest tags” button.")
bullet("PopularityScoreService (IPopularityScoreService): Weighted scoring formula with exponential recency decay; EventDto.PopularityScore powers the homepage's featured-events ordering.")
para(
    "The shared numerical building blocks (TfIdfVectorizer, Bm25Ranker, TextTokenizer, and "
    "LevenshteinDistance) live in NtbEvent.Application/Common and are reused across all four "
    "services, avoiding duplicated tokenization or scoring logic."
)

heading("4.2 Testing", 2)
para(
    "An automated test project (NtbEvent.Tests, xUnit + Moq) was added covering both unit-level "
    "and system-level (cross-service) behaviour. System tests are wired against in-memory fakes "
    "of the repository interfaces rather than a live PostgreSQL host, which keeps the suite fast "
    "and deterministic while still exercising real service composition (EventService + "
    "RecommendationService + SearchRankingService + TagSuggestionService working together)."
)

heading("4.2.1 Unit Test Cases", 3)
add_seq_caption("Table", "Unit Test Cases")
styled_table(
    ["Test Case ID", "Module", "Test", "Result"],
    [
        ("UT-01", "Authentication", "LoginAsync_ValidCredentials_ReturnsAccessAndRefreshTokens", "Pass"),
        ("UT-02", "Authentication", "LoginAsync_InvalidPassword_ThrowsUnauthorized", "Pass"),
        ("UT-03", "Authentication", "RefreshAsync_ExpiredToken_ThrowsUnauthorizedAndIssuesNoNewToken", "Pass"),
        ("UT-04", "Event CRUD", "CreateEventAsync_AllFields_PersistsWithAutoGeneratedSlug", "Pass"),
        ("UT-05", "Event CRUD", "UpdateEventAsync_StatusChangedToPublished_PersistsUpdatedStatus", "Pass"),
        ("UT-06", "Event CRUD", "DeleteEventAsync_ExistingId_RemovesEvent", "Pass"),
        ("UT-07", "AI – TF-IDF", "ComputeTfIdf_SingleDocument_TermsWeightedByFrequency", "Pass"),
        ("UT-08", "AI – Cosine Similarity", "CosineSimilarity_IdenticalDocuments_ReturnsOne", "Pass"),
        ("UT-09", "AI – Cosine Similarity", "CosineSimilarity_CompletelyDifferentDocuments_ReturnsNearZero", "Pass"),
        ("UT-10", "AI – BM25", "Rank_QueryTermInOneDocumentOnly_RanksThatDocumentFirst", "Pass"),
        ("UT-11", "AI – BM25", "Rank_NoMatchingTerms_ReturnsEmpty", "Pass"),
        ("UT-12", "AI – Tag Suggestion", "SuggestTagsAsync_TitleWithFestivalKeyword_ReturnsRelevantExistingTag", "Pass"),
        ("UT-13", "AI – Tag Suggestion", "SuggestTagsAsync_BlankTitleAndDescription_ReturnsEmpty", "Pass"),
        ("UT-14", "AI – Popularity Score", "ScoreAll_HighRatingAndFeatured_ScoresHigherThanNonFeatured", "Pass"),
        ("UT-15", "AI – Popularity Score", "RankByPopularity_ReturnsDescendingOrder", "Pass"),
        ("UT-16", "AI – Popularity Score", "ScoreAll_EmptyList_DoesNotThrow", "Pass"),
    ],
)
doc.add_paragraph()

heading("4.2.2 System Test Cases", 3)
add_seq_caption("Table", "System Test Cases")
styled_table(
    ["Test Case ID", "Scenario", "Result"],
    [
        ("ST-01", "Admin creates event, publishes it, and public user views it on homepage", "Pass"),
        ("ST-02", "Public user searches “Dashain festival” using smart search", "Pass"),
        ("ST-03", "Public user views a festival event and requests recommendations", "Pass"),
        ("ST-04", "Admin creates a new event; system suggests tags from title/description", "Pass"),
        ("ST-05", "Public user filters events by region", "Pass"),
        ("ST-06", "Unauthenticated request to an admin-only endpoint is rejected (not anonymous)", "Pass"),
        ("ST-07", "Event created without an explicit BS date still stores a distinct BS date", "Pass"),
        ("ST-08", "Admin archives an event; it no longer appears in the public listing", "Pass"),
    ],
)
doc.add_paragraph()
para(
    "Running dotnet test NtbEvent.Tests/NtbEvent.Tests.csproj executes all 16 unit tests and 8 "
    "system tests: Passed! – Failed: 0, Passed: 24, Skipped: 0, Total: 24 (completed in under "
    "one second)."
)

heading("4.3 Result Analysis", 2)
para(
    "By the final defense, the following components have been implemented, integrated, and "
    "verified:"
)
bullet("Full backend infrastructure: Clean Architecture project structure, PostgreSQL schema with all EF Core migrations, EF Core + Dapper hybrid data access.")
bullet("User authentication: JWT login, refresh token rotation, role-based access control (SuperAdmin, Admin, Client).")
bullet("Complete Event CRUD: Create, read, update, delete with dual-date support, image upload, tags, highlights, and lifecycle status management.")
bullet("Category and tag management with many-to-many associations and audit tracking.")
bullet("Admin dashboard with event listing, pagination, status management, and category setup.")
bullet("Public homepage with featured events (sorted by popularity score), stats display, hero section, and navigation.")
bullet("Interactive calendar with month/list toggle, event color coding, day event panel, and BS/AD date toggle.")
bullet("Event filtering by type, category, region, and date range.")
bullet("Guest invitation and QR-code check-in: invite → email + QR → guest presents QR → admin scans → verifies → QR is consumed.")
bullet("All four AI modules (TF-IDF/Cosine Similarity recommendations, BM25 smart search, tag suggestion, and popularity scoring) are fully implemented, wired to the frontend, and verified against the live database (12 seeded events across 6 categories and 22 tags, spanning Kathmandu, Pokhara, and Solukhumbu).")
bullet("A 24-case automated test suite (xUnit + Moq) covering authentication, event CRUD, and all four AI modules; all 24 tests pass.")
para(
    "Manual verification against the seeded dataset additionally confirmed that the fuzzy tag "
    "matcher correctly resolves the typo “festivle” to the existing “Festival” tag, that "
    "BM25 search and TF-IDF recommendations return sensible, relevant results, and that the "
    "popularity-sorted featured events section renders correctly on the public homepage. The "
    "system is functionally complete relative to the objectives defined in Chapter 1."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – CONCLUSION AND FUTURE RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 5: Conclusion and Future Recommendations", 1)

heading("5.1 Conclusion", 2)
para(
    "The NTB Event Management System represents a significant step forward in how Nepal's "
    "tourism events are managed, published, and discovered. This final phase has delivered a "
    "fully functional event management platform, covering the complete lifecycle from event "
    "creation by administrators to public discovery by users, plus a digital guest invitation "
    "and QR check-in workflow."
)
para(
    "The project goes beyond a standard CRUD application by fully implementing and integrating "
    "four distinct AI modules: a content-based recommendation engine using TF-IDF and Cosine "
    "Similarity, a smart search module using the BM25 ranking algorithm, an automated tag "
    "suggestion system using keyword extraction with Levenshtein fuzzy matching, and a weighted "
    "popularity scoring model for intelligent event surfacing. All four are built natively in C# "
    "without delegating intelligence to third-party AI APIs. This demonstrates genuine algorithmic "
    "implementation rather than API integration, and the modules are backed by a 24-case automated test "
    "suite that passes in full."
)
para(
    "The Clean Architecture backend ensures the codebase is maintainable and testable, the "
    "dual-calendar support (AD and BS) reflects the real-world requirement of serving a Nepali "
    "government context, and the guest invitation/QR check-in feature extends the platform beyond "
    "the scope originally defined at the mid-term. The system is feature-complete and "
    "deployment-ready."
)

heading("5.2 Future Recommendations", 2)
para("The following enhancements are recommended for future iterations of the platform:")
numbered("Collaborative Filtering: Extend the recommendation engine with a collaborative filtering component that incorporates user behaviour data (view history, bookmarks) once the platform has accumulated sufficient user interaction data.")
numbered("Caching Layer: Implement an in-memory cache (using IMemoryCache in ASP.NET Core) for pre-computed TF-IDF vectors and BM25 indices to further improve response times for AI-powered endpoints as the event catalog grows.")
numbered("Event Ticketing: Add registration and ticketing capabilities, allowing users to register for paid events directly through the platform.")
numbered("Mobile-Responsive Native App: Build a companion mobile application using a cross-platform framework, reusing the existing API layer.")
numbered("Analytics Dashboard: Provide NTB administrators with insights into event views, search queries, popular categories, and regional engagement.")
numbered("Automated Notifications: Extend the existing invitation email system with subscription-based notifications when new events matching a user's interests are published.")
numbered("Event Review System: Allow authenticated users to leave ratings and reviews for events they have attended, enriching the data that powers the popularity scoring model beyond its current proxy signals.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
heading("References", 1)

refs = [
    "[1] S. E. Robertson and H. Zaragoza, \"The Probabilistic Relevance Framework: BM25 and Beyond,\" Foundations and Trends in Information Retrieval, vol. 3, no. 4, pp. 333–389, 2009.",
    "[2] P. Lops, M. de Gemmis, and G. Semeraro, \"Content-based Recommender Systems: State of the Art and Trends,\" in Recommender Systems Handbook, F. Ricci et al., Eds., Springer, 2011, pp. 73–105.",
    "[3] G. Salton and C. Buckley, \"Term-weighting approaches in automatic text retrieval,\" Information Processing and Management, vol. 24, no. 5, pp. 513–523, 1988.",
    "[4] Microsoft, \"ASP.NET Core documentation,\" Microsoft Docs, 2024. [Online]. Available: https://learn.microsoft.com/aspnet/core",
    "[5] PostgreSQL Global Development Group, \"PostgreSQL 16 Documentation,\" 2024. [Online]. Available: https://www.postgresql.org/docs/16/",
    "[6] SvelteKit Contributors, \"SvelteKit Documentation,\" 2024. [Online]. Available: https://kit.svelte.dev/docs",
    "[7] R. Baeza-Yates and B. Ribeiro-Neto, Modern Information Retrieval: The Concepts and Technology behind Search, 2nd ed. ACM Press, 2011.",
    "[8] C. D. Manning, P. Raghavan, and H. Schütze, Introduction to Information Retrieval. Cambridge University Press, 2008.",
    "[9] V. I. Levenshtein, \"Binary codes capable of correcting deletions, insertions, and reversals,\" Soviet Physics Doklady, vol. 10, no. 8, pp. 707–710, 1966.",
    "[10] xUnit.net Contributors, \"xUnit.net Documentation,\" 2024. [Online]. Available: https://xunit.net/",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(ref)
    set_run_font(run, size=12)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDICES - additional screenshots
# ════════════════════════════════════════════════════════════════════════════
heading("Appendices", 1)
para("Additional screenshots supporting the interface design discussed in Chapter 3 are provided below.")

appendix_shots = [
    (find_screenshot("user_management"), "Admin: Create User with Module-Level Permissions", "Admin user management"),
    (find_screenshot("invitation_form"), "Admin: Guest Invitation Form (Event Check-in Stats)", "Guest invitation form"),
    (find_screenshot("swagger"), "Swagger API Documentation", "Swagger API explorer"),
]
any_appendix = False
for path, caption, placeholder in appendix_shots:
    if path:
        any_appendix = True
        add_figure(path, caption, width_inches=5.9)
if not any_appendix:
    para(
        "No additional screenshots beyond those in Section 3.3.4 were captured for this build; "
        "re-run generate_final_report.py after adding more files to the screenshots folder to "
        "populate this section.",
        italic=True,
    )

# ── settings: auto-update fields on open ────────────────────────────────────
settings = doc.settings.element
uf = OxmlElement("w:updateFields")
uf.set(qn("w:val"), "true")
settings.append(uf)

# ── Save ─────────────────────────────────────────────────────────────────────
path = "/home/notcool/Desktop/8th-sem/final Project/NTB_Event_Final_Defense_Report.docx"
doc.save(path)
print("Saved:", path)
print("Figures inserted:", _fig_n[0], "| Tables captioned:", _tbl_n[0])
