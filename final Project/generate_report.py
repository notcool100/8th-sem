from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.right_margin  = Inches(1)
section.left_margin   = Inches(1.25)

# ── Style helpers ────────────────────────────────────────────────────────────
def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         size=12, bold=False, italic=False, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.style = style
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(18)   # 1.5 line spacing
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def heading(text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("POKHARA UNIVERSITY")
set_run_font(run, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Faculty of Engineering and Technology")
set_run_font(run, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Department of Computer Science")
set_run_font(run, size=13)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Mid-Term Defense Report")
set_run_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Project III (CACS452)")
set_run_font(run, size=13, italic=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("NTB EVENT MANAGEMENT SYSTEM")
set_run_font(run, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Smart Event Management Platform for Nepal Tourism Board")
set_run_font(run, size=13, italic=True)

doc.add_paragraph()
doc.add_paragraph()

# Submitted by
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted By:")
set_run_font(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anjal Joshi")
set_run_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("anjaljoshi6@gmail.com")
set_run_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Year/Semester: IV / VIII")
set_run_font(run, size=12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted To:")
set_run_font(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Computer Science")
set_run_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Pokhara University")
set_run_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("June 2026")
set_run_font(run, size=12)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1)
para(
    "The NTB Event Management System is a web-based platform designed specifically for the "
    "Nepal Tourism Board (NTB) to manage, publish, and promote events across Nepal. The system "
    "addresses the lack of a centralized, intelligent platform where tourism-related festivals, "
    "meetings, holidays, and cultural events can be efficiently managed by administrators and "
    "discovered by the public. Beyond standard CRUD operations, the system incorporates a "
    "custom-built smart recommendation engine and an intelligent search module that are implemented "
    "natively within the application — without relying entirely on third-party AI services. "
    "The backend is built using ASP.NET Core 8 following Clean Architecture principles with "
    "PostgreSQL as the database, while the frontend is built using SvelteKit. The AI components "
    "use algorithms such as TF-IDF weighted Cosine Similarity for content-based event "
    "recommendations, BM25 probabilistic ranking for smart full-text search, a Weighted "
    "Popularity Scoring model for event ranking, and keyword extraction for automated tag "
    "suggestions. By the midterm, the core event management infrastructure, user authentication, "
    "calendar view, admin dashboard, and foundational AI modules have been successfully implemented."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)

toc_items = [
    ("Abstract", "ii"),
    ("Chapter 1: Introduction", "1"),
    ("  1.1  Introduction", "1"),
    ("  1.2  Problem Statement", "2"),
    ("  1.3  Objectives", "2"),
    ("  1.4  Scope and Limitation", "3"),
    ("  1.5  Development Methodology", "3"),
    ("  1.6  Report Organization", "4"),
    ("Chapter 2: Background Study and Literature Review", "5"),
    ("  2.1  Background Study", "5"),
    ("  2.2  Literature Review", "6"),
    ("Chapter 3: System Analysis and Design", "8"),
    ("  3.1  System Analysis", "8"),
    ("  3.2  Feasibility Analysis", "10"),
    ("  3.3  System Design", "11"),
    ("  3.4  Algorithm Details", "13"),
    ("Chapter 4: Implementation and Testing", "18"),
    ("  4.1  Implementation", "18"),
    ("  4.2  Testing", "21"),
    ("  4.3  Result Analysis", "22"),
    ("Chapter 5: Conclusion and Future Recommendations", "23"),
    ("  5.1  Conclusion", "23"),
    ("  5.2  Future Recommendations", "24"),
    ("References", "25"),
]

for item, page_num in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(item)
    set_run_font(run, size=12, bold=("Chapter" in item))
    tab_run = p.add_run(f"\t{page_num}")
    set_run_font(tab_run, size=12)
    pass  # simple TOC without dot leaders

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 1: Introduction", 1)

heading("1.1 Introduction", 2)
para(
    "Nepal is a country of extraordinary cultural richness and natural diversity. Every year, "
    "hundreds of festivals, tourism events, national holidays, and community gatherings are "
    "organized across the country. The Nepal Tourism Board (NTB), as the official government "
    "body responsible for promoting Nepal as a tourism destination, plays a central role in "
    "organizing and publicizing these events to both domestic and international audiences."
)
para(
    "However, managing and publishing this large volume of events through disconnected or manual "
    "processes creates significant operational inefficiencies. There is no unified, modern platform "
    "that brings event creation, publication, calendar management, and intelligent event discovery "
    "under one roof. The NTB Event Management System is built to solve exactly this problem."
)
para(
    "This system is a full-stack web application where authorized administrators can create, "
    "update, and manage events of various types — festivals, meetings, holidays, and general "
    "events. The public can browse, filter, and explore these events through an interactive "
    "calendar and listing interface. What sets this platform apart is the integration of "
    "self-built AI modules that enhance search quality, surface relevant events to users, "
    "and assist administrators with intelligent tag suggestions during event creation."
)

heading("1.2 Problem Statement", 2)
para(
    "Nepal's tourism event landscape suffers from fragmentation and lack of digital structure. "
    "The core problems that motivated this project are:"
)
bullet("No centralized digital platform exists for NTB to manage all event types from a single admin interface.")
bullet("The public has no reliable, filterable, and searchable directory of NTB-published events.")
bullet("Event discovery is passive — users must scroll through lists with no intelligence behind what is surfaced.")
bullet("Manual content tagging is inconsistent, making category-based filtering less reliable.")
bullet("There is no way to track event popularity or surface trending or high-impact events automatically.")
para(
    "These gaps result in missed opportunities for tourism promotion, lower public awareness, and "
    "high administrative overhead. The NTB Event Management System directly addresses each of these "
    "problems through structured event management workflows and embedded AI-driven discovery features."
)

heading("1.3 Objectives", 2)
para("The primary objectives of this project are as follows:")
numbered("To design and develop a complete, production-grade event management platform for the Nepal Tourism Board.")
numbered("To implement a role-based access control system that supports Super Admin, Admin, and Client roles with appropriate permissions.")
numbered("To build an interactive calendar interface that displays events in both month-view and list-view with dual-calendar support (AD and BS).")
numbered("To implement a custom-built Content-Based Recommendation Engine using TF-IDF and Cosine Similarity that suggests relevant events without relying on third-party AI APIs.")
numbered("To develop a smart search module powered by the BM25 ranking algorithm, replacing naive keyword matching with a probabilistic relevance model.")
numbered("To create an automated tag suggestion feature using keyword extraction so administrators receive intelligent recommendations while creating events.")
numbered("To implement a Weighted Popularity Scoring model that ranks events based on rating, attendance, recency, and featured status.")
numbered("To ensure the platform follows Clean Architecture principles on the backend, ensuring maintainability and testability.")

heading("1.4 Scope and Limitation", 2)
para("The scope of the NTB Event Management System covers the following areas:")
bullet("Full event lifecycle management: draft, publish, and archive.")
bullet("Multiple event types: Festival, Meeting, Holiday, and General Event.")
bullet("Image upload and media management for events.")
bullet("Category and tag management with support for many-to-many relationships.")
bullet("User management with role-based access: Super Admin, Admin, and Client.")
bullet("JWT-based authentication with refresh token rotation.")
bullet("Public-facing event discovery with calendar view, filtering, and smart search.")
bullet("AI-assisted content recommendation, search ranking, and tag suggestion.")
bullet("Support for both AD (English calendar) and BS (Bikram Sambat) date formats.")

para("The current limitations of the system include:")
bullet("The recommendation engine is content-based; collaborative filtering (based on user behaviour history) is planned for future phases.")
bullet("The system does not currently support event ticketing or payment processing.")
bullet("Mobile native applications (iOS/Android) are outside the current scope; the platform is optimized as a responsive web application.")
bullet("The AI modules currently run at request time; background pre-computation and caching are planned as a future enhancement.")

heading("1.5 Development Methodology", 2)
para(
    "This project follows an Agile-Incremental development methodology. Rather than planning and "
    "building everything at once, the system has been developed in iterative cycles, with each "
    "iteration delivering a working vertical slice of functionality."
)
para(
    "The development was divided into the following major increments:"
)
bullet("Increment 1 — Core Infrastructure: Database schema design, Clean Architecture project setup, JWT authentication, user roles, and basic event CRUD.")
bullet("Increment 2 — Frontend Foundation: SvelteKit project setup, admin panel layout, event creation/editing forms, and image upload.")
bullet("Increment 3 — Calendar & Public View: Interactive calendar with month/list toggle, public event discovery page, filtering, and category/region-based navigation.")
bullet("Increment 4 — AI Modules (In Progress): TF-IDF and Cosine Similarity based recommendation engine, BM25 search ranking, keyword extraction for tag suggestion, and popularity scoring.")
bullet("Increment 5 (Planned) — Polish & Testing: Comprehensive unit and system testing, result analysis, and user documentation.")

para(
    "This approach allows for continuous delivery of working functionality while enabling course "
    "corrections based on testing and supervisor feedback throughout the project lifecycle."
)

heading("1.6 Report Organization", 2)
para("This report is organized into five chapters:")
bullet("Chapter 1 — Introduction: Provides the background, problem statement, objectives, scope, and development methodology.")
bullet("Chapter 2 — Background Study and Literature Review: Covers the fundamental concepts and reviews similar systems and research work.")
bullet("Chapter 3 — System Analysis and Design: Presents the requirement analysis, feasibility study, system diagrams, and algorithm details.")
bullet("Chapter 4 — Implementation and Testing: Details the tools used, module implementations, test cases, and result analysis.")
bullet("Chapter 5 — Conclusion and Future Recommendations: Summarizes the progress and outlines the future work planned.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – BACKGROUND STUDY AND LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 2: Background Study and Literature Review", 1)

heading("2.1 Background Study", 2)
para(
    "This chapter provides the theoretical foundation behind the key concepts used in the "
    "NTB Event Management System, with a particular focus on the AI and algorithmic components."
)

heading("2.1.1 Event Management Systems", 3)
para(
    "An Event Management System (EMS) is a software platform that facilitates the planning, "
    "organization, publication, and administration of events. Modern EMS platforms go beyond "
    "simple CRUD operations by integrating features such as calendar visualization, multi-role "
    "user access control, media management, and intelligent content discovery. In the context "
    "of a national tourism body like NTB, an EMS must be robust, scalable, and capable of "
    "serving both internal administrators and the general public simultaneously."
)

heading("2.1.2 Content-Based Filtering and TF-IDF", 3)
para(
    "Content-based filtering is a recommendation technique that suggests items similar to those "
    "a user has shown interest in, based on the properties of the items themselves rather than "
    "the behaviour of other users. In the context of event recommendation, if a user views a "
    "Dashain Festival event, the system should surface other festival events with similar topics, "
    "locations, or keywords."
)
para(
    "Term Frequency-Inverse Document Frequency (TF-IDF) is a numerical statistic that reflects "
    "how important a word is to a document within a collection. The TF (Term Frequency) component "
    "measures how often a word appears in a document, while the IDF (Inverse Document Frequency) "
    "component reduces the weight of words that appear in many documents (common words) and "
    "increases the weight of rare, domain-specific words. Together, they produce a vector "
    "representation of a document that captures its meaningful content."
)

heading("2.1.3 Cosine Similarity", 3)
para(
    "Cosine Similarity is a metric used to measure the similarity between two vectors by "
    "computing the cosine of the angle between them in a multi-dimensional space. When two "
    "TF-IDF vectors are compared using cosine similarity, the result is a score between 0 and 1, "
    "where 1 means identical content and 0 means completely unrelated content. This is the "
    "mathematical core of the recommendation engine in this system."
)

heading("2.1.4 BM25 – Probabilistic Information Retrieval", 3)
para(
    "BM25 (Best Match 25), also known as Okapi BM25, is a ranking function used in information "
    "retrieval to rank documents based on their relevance to a given search query. Unlike a simple "
    "keyword match (ILIKE in SQL), BM25 accounts for the frequency of query terms in each document, "
    "the inverse document frequency of each term, and the document length — meaning longer documents "
    "are not unfairly rewarded for containing query terms more times. BM25 is the foundation of "
    "modern search engines like Elasticsearch and Lucene, and its implementation in this project "
    "provides significantly more relevant search results than plain pattern matching."
)

heading("2.1.5 Weighted Scoring Models", 3)
para(
    "A weighted scoring model is a simple but effective technique for ranking items based on "
    "multiple criteria simultaneously. Each criterion (e.g., rating, attendance, recency, featured "
    "status) is assigned a weight reflecting its relative importance. The final score is the "
    "weighted sum of all criteria. This technique is widely used in recommendation systems to "
    "surface popular or high-quality items without requiring complex machine learning infrastructure."
)

heading("2.1.6 Keyword Extraction", 3)
para(
    "Keyword extraction is the process of automatically identifying the most relevant and "
    "descriptive words or phrases from a piece of text. In this system, when an administrator "
    "is creating an event, the title and description are analysed to extract the top keywords "
    "using a TF-IDF-based approach against the existing event corpus. These keywords are then "
    "matched against the available tags in the system to produce intelligent tag suggestions."
)

heading("2.2 Literature Review", 2)
para(
    "Several existing event management and tourism platforms have been studied to understand the "
    "current state of the art and identify areas where the NTB system can provide better value."
)

heading("2.2.1 Eventbrite", 3)
para(
    "Eventbrite is one of the world's largest event management platforms. It supports event "
    "discovery, ticketing, and registration. Its recommendation system is primarily collaborative "
    "filtering-based, leveraging purchase history and user behaviour. While powerful, Eventbrite's "
    "approach requires large volumes of user behavioural data to function well — something that "
    "a new platform serving a national tourism body would not have at launch. The NTB system's "
    "content-based approach using TF-IDF and Cosine Similarity is more suitable for a cold-start "
    "scenario where user data is limited."
)

heading("2.2.2 Nepal's Existing Tourism Information Portals", 3)
para(
    "Nepal's existing government tourism portals (such as tourismdepartment.gov.np) primarily "
    "serve static informational content. They lack dynamic event management features, interactive "
    "calendars, admin workflows, or any form of intelligent event discovery. This confirms the "
    "gap that the NTB Event Management System is designed to fill."
)

heading("2.2.3 Research on BM25 in Web Applications", 3)
para(
    "Research by Robertson and Zaragoza (2009) established BM25 as a robust, parameter-tunable "
    "ranking algorithm for text retrieval. Subsequent work has shown that BM25 consistently "
    "outperforms simple TF-based models in precision and recall for domain-specific search tasks. "
    "Implementing BM25 natively within the application layer (rather than delegating to a "
    "full-text search engine) is a deliberate design choice in this project, demonstrating "
    "algorithm-level understanding rather than reliance on infrastructure abstractions."
)

heading("2.2.4 TF-IDF in Recommendation Systems", 3)
para(
    "Lops et al. (2011) provide a comprehensive survey of content-based filtering techniques "
    "and identify TF-IDF as a foundational and highly interpretable method for building item "
    "profiles in recommendation systems. The survey notes that TF-IDF based content-based "
    "filtering performs particularly well for text-rich domains — which aligns directly with "
    "the event management domain where events have rich titles, descriptions, categories, and tags."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – SYSTEM ANALYSIS AND DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 3: System Analysis and Design", 1)

heading("3.1 System Analysis", 2)

heading("3.1.1 Functional Requirements", 3)
para("The functional requirements of the NTB Event Management System are organized by user role:")

para("Administrator Requirements:", bold=True)
bullet("Admin can log in securely using email and password with JWT-based session management.")
bullet("Admin can create, edit, and delete events with full details including title, description, category, type, date (AD and BS), location, region, price, images, highlights, tags, and organizer information.")
bullet("Admin can change the lifecycle status of events: Draft, Published, or Archived.")
bullet("Admin can manage categories and tags, including creating and organizing tag-category associations.")
bullet("Admin can manage system users and assign roles.")
bullet("Admin can view a calendar showing all events with color-coded type indicators.")

para("Public User Requirements:", bold=True)
bullet("Public users can browse events on the homepage with featured event highlights.")
bullet("Public users can filter events by type, category, region, and date range.")
bullet("Public users can use the smart search to find events by title, location, or description.")
bullet("Public users can view events in both calendar and list view.")
bullet("Public users can view detailed event information including location, organizer, highlights, and images.")
bullet("Public users receive smart event recommendations based on the event they are currently viewing.")

heading("3.1.2 Non-Functional Requirements", 3)
bullet("Performance: API response time should be under 200ms for standard queries; search and recommendation results should return within 500ms.")
bullet("Security: All admin endpoints are protected by JWT authentication. Passwords are hashed using BCrypt. Refresh token rotation prevents token replay attacks.")
bullet("Scalability: The Clean Architecture design ensures business logic is decoupled from infrastructure, enabling horizontal scaling when needed.")
bullet("Usability: The frontend must be responsive and work seamlessly on both desktop and mobile devices.")
bullet("Maintainability: Clean Architecture layers (Domain, Application, Infrastructure, API) ensure clear separation of concerns.")
bullet("Data Integrity: Dual-calendar date support (AD/BS) must be consistently maintained across all event records.")

heading("3.2 Feasibility Analysis", 2)

heading("3.2.1 Technical Feasibility", 3)
para(
    "The system is built on mature, well-supported technologies. ASP.NET Core 8 is a "
    "production-grade web framework maintained by Microsoft. PostgreSQL is a battle-tested, "
    "open-source relational database. SvelteKit is a modern, performant web framework. "
    "The AI algorithms (TF-IDF, Cosine Similarity, BM25) are mathematically well-defined "
    "and can be implemented in pure C# without any external AI library dependencies. "
    "All members of the team have hands-on experience with these technologies, confirming "
    "technical feasibility."
)

heading("3.2.2 Operational Feasibility", 3)
para(
    "NTB administrators are familiar with web-based administration tools. The system's admin "
    "interface is designed to be intuitive and requires minimal training. The public-facing "
    "interface is simple and discoverable. From an operational standpoint, the system reduces "
    "manual event publication effort and eliminates the need for external tools or third-party "
    "content management platforms."
)

heading("3.2.3 Economic Feasibility", 3)
para(
    "The entire technology stack is open-source with no licensing fees. The system can be "
    "hosted on cloud providers (such as AWS, Azure, or DigitalOcean) at a low monthly cost. "
    "Development effort is entirely student-driven. The long-term cost savings from eliminating "
    "manual event management workflows and reducing dependence on external platforms make the "
    "system economically viable for NTB."
)

heading("3.3 System Design", 2)

heading("3.3.1 System Architecture Overview", 3)
para(
    "The NTB Event Management System follows Clean Architecture on the backend, organized into "
    "four distinct layers:"
)
bullet("Domain Layer: Contains the core business entities (Event, User, Category, Tag, RefreshToken, EventHighlight) and enumerations (EventType, EventLifecycleStatus, UserRole). This layer has no dependencies on any external framework.")
bullet("Application Layer: Contains service interfaces, repository contracts, DTOs, and business logic. It depends only on the Domain layer.")
bullet("Infrastructure Layer: Implements the repository contracts and services defined in the Application layer. It uses Entity Framework Core for CRUD operations and Dapper for performance-critical list queries.")
bullet("API Layer: Hosts the ASP.NET Core controllers, handles HTTP request/response, JWT configuration, CORS, Swagger documentation, and dependency injection wiring.")

heading("3.3.2 Database Design", 3)
para(
    "The primary entities and their relationships in the PostgreSQL database are:"
)
bullet("events: Stores all event data including title, description, dates (AD and BS), location, region, type, status, images, price, rating, organizer information, and highlighted features as a JSON column.")
bullet("users: Stores registered user accounts with hashed passwords and role assignments.")
bullet("refresh_tokens: Stores refresh tokens per user session for secure token rotation.")
bullet("categories: Stores event categories managed by administrators.")
bullet("tags: Stores descriptive tags that can be associated with events and categories.")
bullet("event_tags: Junction table linking events to tags (many-to-many).")
bullet("categories_tags: Junction table linking categories to their default tags.")
bullet("tags_jn (Audit): Tracks tag association changes for audit purposes.")

heading("3.3.3 API Design", 3)
para("The backend exposes the following RESTful API endpoints:")
bullet("POST /api/auth/login — Authenticate user and receive JWT + refresh token.")
bullet("POST /api/auth/refresh — Rotate refresh token and issue a new access token.")
bullet("POST /api/auth/logout — Invalidate the current refresh token.")
bullet("GET /api/events — Get paginated, filtered list of events (public).")
bullet("GET /api/events/{id} — Get a single event by ID.")
bullet("GET /api/events/{slug} — Get a single event by slug.")
bullet("POST /api/events — Create a new event (Admin).")
bullet("PUT /api/events/{id} — Update an event (Admin).")
bullet("DELETE /api/events/{id} — Delete an event (Admin).")
bullet("GET /api/events/{id}/recommendations — Get AI-powered recommendations for an event.")
bullet("GET /api/events/search?q={query} — Smart BM25-ranked search.")
bullet("GET /api/categories — List all categories.")
bullet("GET /api/tags — List all tags, optionally filtered by category.")
bullet("POST /api/events/suggest-tags — Get AI-suggested tags for a title and description.")

heading("3.4 Algorithm Details", 2)
para(
    "This section describes the four AI and algorithmic modules implemented natively within "
    "the NTB Event Management System. All algorithms are implemented in C# within the "
    "Application/Infrastructure layers and do not depend on any external AI API or service."
)

heading("3.4.1 Content-Based Event Recommendation Engine (TF-IDF + Cosine Similarity)", 3)
para(
    "The recommendation engine answers the question: given the event a user is currently viewing, "
    "which other events are most similar in content? This is a classic content-based filtering "
    "problem solved using TF-IDF vectorization and Cosine Similarity."
)

para("Step 1 — Build the Event Corpus:", bold=True)
para(
    "Each event in the database is represented as a text document by concatenating its most "
    "meaningful text fields: title (weighted 3×), summary (weighted 2×), long description, "
    "category name, region, and associated tags. Title repetition amplifies its influence "
    "on the final vector — this is a standard technique in content-based systems."
)

para("Step 2 — Compute TF-IDF Vectors:", bold=True)
para(
    "For each event document, the system tokenizes the text, removes stop words (common words "
    "such as 'the', 'is', 'in' which carry no meaning), and computes TF-IDF scores for each "
    "remaining token. The vocabulary is built from all events in the corpus."
)
para(
    "The TF-IDF formula used is:"
)
para(
    "    TF(t, d)  = (count of term t in document d) / (total terms in document d)",
    bold=False
)
para(
    "    IDF(t)    = log( (1 + N) / (1 + df(t)) ) + 1",
    bold=False
)
para(
    "    TF-IDF(t, d) = TF(t, d) × IDF(t)",
    bold=False
)
para(
    "Where N is the total number of event documents and df(t) is the number of documents "
    "that contain term t. The +1 smoothing in the IDF formula prevents zero values for "
    "terms that appear in every document."
)

para("Step 3 — Compute Cosine Similarity:", bold=True)
para(
    "Given the TF-IDF vector of the target event (the one being viewed) and the TF-IDF vectors "
    "of all other events, cosine similarity is computed as:"
)
para(
    "    similarity(A, B) = (A · B) / (||A|| × ||B||)",
    bold=False
)
para(
    "Where A · B is the dot product of the two vectors and ||A|| and ||B|| are their magnitudes. "
    "This gives a score between 0 and 1 for each pair. Events are then sorted by descending "
    "similarity score, and the top-N results (default: 5) are returned as recommendations."
)

para("Why This Is Not a SQL Query:", bold=True)
para(
    "This algorithm operates entirely in the C# application layer. The database is used only to "
    "retrieve event data. The TF-IDF computation, vector construction, and similarity scoring are "
    "all done in memory using mathematical operations — not SQL filters or joins."
)

heading("3.4.2 Smart Search with BM25 Ranking", 3)
para(
    "The smart search module replaces the naive ILIKE pattern matching with a proper probabilistic "
    "ranking model called BM25 (Best Match 25), which is the same ranking function used by "
    "Elasticsearch, Solr, and many major search engines internally."
)

para("How BM25 Works:", bold=True)
para(
    "Given a user's search query Q containing terms q1, q2, ..., qn, BM25 scores each event "
    "document D as follows:"
)
para(
    "    Score(D, Q) = Σ IDF(qi) × [ f(qi, D) × (k1 + 1) ] / [ f(qi, D) + k1 × (1 - b + b × |D| / avgdl) ]",
    bold=False
)
para(
    "Where:"
)
bullet("f(qi, D) is the frequency of query term qi in document D.")
bullet("|D| is the length of document D in words.")
bullet("avgdl is the average document length across all events.")
bullet("k1 is a term saturation parameter (default: 1.5) — controls how much repeated terms increase the score.")
bullet("b is a length normalization parameter (default: 0.75) — controls how much document length affects scoring.")

para(
    "Each event's searchable text (title + summary + location + region) is tokenized and indexed "
    "at application startup (or on demand) into an in-memory inverted index. When a search query "
    "arrives, BM25 scores are computed in O(k) time where k is the number of matching terms."
)

para("Why BM25 Over ILIKE:", bold=True)
para(
    "A plain ILIKE query in SQL returns all events that simply contain the search term, with no "
    "relevance ranking. BM25 assigns higher scores to events where the query terms appear frequently "
    "but the document is not artificially long, and where the query terms are rare across the corpus "
    "(more discriminative). This results in meaningfully better search relevance — the most "
    "important events float to the top of results rather than appearing in arbitrary insertion order."
)

heading("3.4.3 Automated Tag Suggestion via Keyword Extraction", 3)
para(
    "When an administrator is in the process of creating or editing an event, the system "
    "automatically analyses the event's title and description to suggest relevant tags. "
    "This reduces manual tagging effort and improves consistency across the event catalog."
)

para("How It Works:", bold=True)
para(
    "The keyword extraction module uses a corpus-aware TF-IDF scoring approach:"
)
numbered("The input text (title + description) is tokenized and stop words are removed.")
numbered("TF-IDF scores are computed for each token in the input text against the existing event corpus.")
numbered("The top-K tokens by TF-IDF score are selected as candidate keywords.")
numbered("Each candidate keyword is matched against the system's existing tag list using normalized string comparison and edit distance (Levenshtein distance for fuzzy matching).")
numbered("Matched tags are returned in ranked order as suggestions to the administrator.")

para(
    "This process is entirely self-contained — it uses the event corpus already stored in the "
    "database and the tag dictionary maintained by administrators. No external NLP service or "
    "API is called. The algorithm is implemented as a service within the Application layer."
)

heading("3.4.4 Weighted Popularity Scoring for Event Ranking", 3)
para(
    "To surface the most valuable events to users on the homepage and in category listings, "
    "the system computes a Popularity Score for each event using a multi-factor weighted formula. "
    "This is not a machine learning model — it is a transparent, tunable scoring function that "
    "administrators can understand and adjust."
)

para("Scoring Formula:", bold=True)
para(
    "    PopularityScore(e) = (W_r × NormalizedRating) + (W_a × NormalizedAttendance) "
    "+ (W_f × FeaturedBoost) + (W_d × RecencyScore)",
    bold=False
)
para("Where the weights are:")
bullet("W_r = 0.35  — Rating weight (user ratings are the strongest signal of quality).")
bullet("W_a = 0.25  — Attendance/engagement weight.")
bullet("W_f = 0.20  — Featured boost (admin-marked events get a boost).")
bullet("W_d = 0.20  — Recency score (more recent events rank higher to keep listings fresh).")

para(
    "NormalizedRating is computed as (event_rating / max_rating). RecencyScore uses an "
    "exponential decay function: RecencyScore = e^(−λ × days_since_event_start), where λ = 0.01 "
    "ensures that events decay gradually over ~100 days. This formula can be computed entirely "
    "in C# without any database aggregation query, and its output is used to sort featured "
    "events on the homepage."
)

para("Why This Matters:", bold=True)
para(
    "Without this scoring model, events would simply be sorted by creation date or start date. "
    "The weighted model ensures that high-rated, well-attended, recently-starting festivals "
    "appear prominently — which is exactly what NTB wants: showcasing Nepal's most impactful "
    "tourism events to visitors."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – IMPLEMENTATION AND TESTING
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 4: Implementation and Testing", 1)

heading("4.1 Implementation", 2)

heading("4.1.1 Tools and Technologies Used", 3)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["Component", "Technology", "Purpose"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Backend Framework", "ASP.NET Core 8", "REST API, routing, middleware"),
    ("Architecture Pattern", "Clean Architecture", "Separation of concerns, testability"),
    ("Primary ORM", "Entity Framework Core 8", "CRUD operations, migrations"),
    ("Query Micro-ORM", "Dapper", "Performance-optimised list queries"),
    ("Database", "PostgreSQL 16", "Relational data storage"),
    ("Authentication", "JWT + Refresh Tokens", "Stateless, secure session management"),
    ("Password Hashing", "BCrypt.Net", "Secure credential storage"),
    ("Frontend Framework", "SvelteKit + TypeScript", "Reactive, performant UI"),
    ("Styling", "CSS Variables + Tailwind", "Responsive design system"),
    ("API Documentation", "Swagger / OpenAPI", "Interactive API explorer"),
    ("AI Algorithms", "Custom C# Modules", "TF-IDF, BM25, Cosine Similarity, Scoring"),
    ("Version Control", "Git / GitHub", "Source code management"),
]
for row_data in rows_data:
    row = table.add_row().cells
    for cell, txt in zip(row, row_data):
        cell.text = txt
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

heading("4.1.2 Module Implementation Details", 3)

para("Authentication Module:", bold=True)
para(
    "The authentication module implements JWT access tokens with a short expiry (15 minutes) "
    "and long-lived refresh tokens stored in the database. On each successful refresh, the old "
    "token is invalidated and a new one is issued (token rotation). This prevents refresh token "
    "reuse attacks. The User entity carries a UserRole enum (SuperAdmin, Admin, Client) that is "
    "embedded in the JWT claims, enabling role-based route protection across all admin endpoints."
)

para("Event Management Module:", bold=True)
para(
    "Events are created and managed through the EventsController, which delegates to the "
    "IEventService in the Application layer. The EventRepository uses Entity Framework Core for "
    "CRUD operations and Dapper for the paginated list query. The list query uses a dynamic WHERE "
    "clause builder (BuildWhereClause) that constructs parameterized SQL based on the active "
    "filters — preventing SQL injection while maintaining query flexibility. Events support dual "
    "dates (DateAd/DateBs), multiple images stored as a JSON array, and structured highlights "
    "stored as a JSON column."
)

para("Calendar Module:", bold=True)
para(
    "The frontend calendar is built as a custom SvelteKit component with no external calendar "
    "library dependency. It supports month-view and list-view toggles. Events are color-coded "
    "by type (Festival, Meeting, Holiday, Event). The calendar supports both AD and BS date "
    "display through a toggle mechanism. Date utility functions handle month navigation, day "
    "cell generation, and event-to-date mapping."
)

para("Tag and Category Module:", bold=True)
para(
    "Categories and tags have a many-to-many relationship through the CategoriesTags junction. "
    "Events also relate to tags through the EventTags junction. An audit table (TagsJn) tracks "
    "changes to tag associations over time. The tag suggestion AI module (described in Chapter 3) "
    "operates on top of this data structure."
)

para("AI Modules (In Progress):", bold=True)
para(
    "The four AI modules described in Section 3.4 are currently being implemented as services "
    "within the Application and Infrastructure layers:"
)
bullet("RecommendationService: Implements TF-IDF vectorization and Cosine Similarity computation.")
bullet("SearchRankingService: Implements the BM25 inverted index and scoring function.")
bullet("TagSuggestionService: Implements keyword extraction with Levenshtein-based tag matching.")
bullet("PopularityScoreService: Implements the weighted scoring formula with exponential recency decay.")

heading("4.2 Testing", 2)

heading("4.2.1 Unit Test Cases", 3)

table2 = doc.add_table(rows=1, cols=4)
table2.style = "Table Grid"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for cell, txt in zip(hdr2, ["Test Case ID", "Module", "Description", "Expected Result"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

unit_tests = [
    ("UT-01", "Authentication", "Login with valid credentials", "JWT token and refresh token returned"),
    ("UT-02", "Authentication", "Login with invalid password", "401 Unauthorized response"),
    ("UT-03", "Authentication", "Refresh with expired token", "401 Unauthorized, new token not issued"),
    ("UT-04", "Event CRUD", "Create event with all fields", "Event persisted, slug auto-generated"),
    ("UT-05", "Event CRUD", "Update event status to Published", "Status updated in database"),
    ("UT-06", "Event CRUD", "Delete event by ID", "Event removed from database"),
    ("UT-07", "AI – TF-IDF", "Compute TF-IDF for single document", "Correct term weights returned"),
    ("UT-08", "AI – Cosine Similarity", "Identical documents", "Similarity score = 1.0"),
    ("UT-09", "AI – Cosine Similarity", "Completely different documents", "Similarity score near 0"),
    ("UT-10", "AI – BM25", "Query term found in one document only", "That document ranked first"),
    ("UT-11", "AI – Tag Suggestion", "Extract keywords from event title", "Relevant tags returned"),
    ("UT-12", "AI – Popularity Score", "Event with high rating + featured", "Higher score than non-featured"),
]
for row_data in unit_tests:
    row = table2.add_row().cells
    for cell, txt in zip(row, row_data):
        cell.text = txt
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

heading("4.2.2 System Test Cases", 3)

table3 = doc.add_table(rows=1, cols=3)
table3.style = "Table Grid"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for cell, txt in zip(hdr3, ["Test Case ID", "Scenario", "Expected Outcome"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True

sys_tests = [
    ("ST-01", "Admin creates event, publishes it, and public user views it on homepage", "Event appears in public listing with correct details"),
    ("ST-02", "Public user searches 'Dashain festival' using smart search", "Festival events ranked higher than unrelated events"),
    ("ST-03", "Public user views a festival event, scrolls to recommendations section", "Top 5 events from similar categories returned"),
    ("ST-04", "Admin creates new event, system suggests tags from title/description", "Relevant tags from existing database returned"),
    ("ST-05", "Public user filters events by region 'Kathmandu'", "Only Kathmandu events returned"),
    ("ST-06", "Admin attempts to access admin dashboard without login", "Redirected to login page"),
    ("ST-07", "Public user toggles calendar to BS date view", "Event dates displayed in Bikram Sambat"),
    ("ST-08", "Admin archives an event", "Event no longer appears in public listing"),
]
for row_data in sys_tests:
    row = table3.add_row().cells
    for cell, txt in zip(row, row_data):
        cell.text = txt
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

heading("4.3 Result Analysis", 2)
para(
    "By the midterm milestone, the following components have been successfully implemented "
    "and verified through manual testing and code review:"
)
bullet("Full backend infrastructure: Clean Architecture project structure, PostgreSQL schema with all migrations, EF Core + Dapper hybrid data access.")
bullet("User authentication: JWT login, refresh token rotation, role-based access control (SuperAdmin, Admin, Client).")
bullet("Complete Event CRUD: Create, read, update, delete with dual-date support, image upload, tags, highlights, and lifecycle status management.")
bullet("Category and tag management with many-to-many associations and audit tracking.")
bullet("Admin dashboard with event listing, pagination, status management, and category setup.")
bullet("Public homepage with featured events, stats display, hero section, and navigation.")
bullet("Interactive calendar with month/list toggle, event color coding, day event panel, and BS/AD date toggle.")
bullet("Event filtering by type, category, region, and date range.")
bullet("AI modules: TF-IDF vectorization and Cosine Similarity engine (implementation in progress), BM25 search module (in progress), tag suggestion service (in progress), popularity scoring model (design complete, implementation in progress).")
para(
    "The overall system is functioning as a complete event management platform. The AI modules "
    "are at the algorithmic design and early implementation stage, with integration into the "
    "frontend planned for the next sprint before the final defense."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – CONCLUSION AND FUTURE RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════════════════
heading("Chapter 5: Conclusion and Future Recommendations", 1)

heading("5.1 Conclusion", 2)
para(
    "The NTB Event Management System represents a significant step forward in how Nepal's "
    "tourism events are managed, published, and discovered. The midterm phase has successfully "
    "delivered a fully functional event management backend and a polished, responsive frontend — "
    "covering the complete lifecycle from event creation by administrators to public discovery "
    "by users."
)
para(
    "The project goes beyond a standard CRUD application by designing and beginning the "
    "implementation of four distinct AI modules: a content-based recommendation engine using "
    "TF-IDF and Cosine Similarity, a smart search module using the BM25 ranking algorithm, "
    "an automated tag suggestion system using keyword extraction, and a weighted popularity "
    "scoring model for intelligent event surfacing. All of these are built natively in C# "
    "without delegating intelligence to third-party AI APIs — demonstrating genuine "
    "algorithmic implementation rather than API integration."
)
para(
    "The Clean Architecture backend ensures the codebase is maintainable, testable, and ready "
    "for the additional feature work planned for the final submission. The dual-calendar support "
    "(AD and BS) reflects the real-world requirement of serving a Nepali government context. "
    "The system is on track to be fully feature-complete and deployment-ready by the final defense."
)

heading("5.2 Future Recommendations", 2)
para("The following enhancements are planned for the final submission phase and beyond:")
numbered("Complete AI Module Integration: Finalize the implementation and frontend integration of all four AI modules, with API endpoints exposed and tested.")
numbered("Collaborative Filtering: Extend the recommendation engine with a collaborative filtering component that incorporates user behaviour data (view history, bookmarks) once the platform has accumulated user interaction data.")
numbered("Caching Layer: Implement an in-memory cache (using IMemoryCache in ASP.NET Core) for pre-computed TF-IDF vectors and BM25 indices to improve response times for AI-powered endpoints.")
numbered("Event Ticketing: Add registration and ticketing capabilities, allowing users to register for events directly through the platform.")
numbered("Mobile-Responsive Native App: Build a companion mobile application using a cross-platform framework.")
numbered("Analytics Dashboard: Provide NTB administrators with insights into event views, search queries, popular categories, and regional engagement.")
numbered("Automated Notifications: Implement email and push notification support to alert subscribers when new events matching their interests are published.")
numbered("Event Review System: Allow authenticated users to leave ratings and reviews for events they have attended, enriching the data that powers the popularity scoring model.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
heading("References", 1)

refs = [
    "[1] S. E. Robertson and H. Zaragoza, \"The Probabilistic Relevance Framework: BM25 and Beyond,\" Foundations and Trends in Information Retrieval, vol. 3, no. 4, pp. 333–389, 2009.",
    "[2] P. Lops, M. de Gemmis, and G. Semeraro, \"Content-based Recommender Systems: State of the Art and Trends,\" in Recommender Systems Handbook, F. Ricci et al., Eds., Springer, 2011, pp. 73–105.",
    "[3] G. Salton and C. Buckley, \"Term-weighting approaches in automatic text retrieval,\" Information Processing and Management, vol. 24, no. 5, pp. 513–523, 1988.",
    "[4] Microsoft, \"ASP.NET Core documentation,\" Microsoft Docs, 2024. [Online]. Available: https://learn.microsoft.com/aspnet/core",
    "[5] PostgreSQL Global Development Group, \"PostgreSQL 16 Documentation,\" 2024. [Online]. Available: https://www.postgresql.org/docs/16/",
    "[6] SvelteKit Contributors, \"SvelteKit Documentation,\" 2024. [Online]. Available: https://kit.svelte.dev/docs",
    "[7] R. Baeza-Yates and B. Ribeiro-Neto, Modern Information Retrieval: The Concepts and Technology behind Search, 2nd ed. ACM Press, 2011.",
    "[8] C. D. Manning, P. Raghavan, and H. Schütze, Introduction to Information Retrieval. Cambridge University Press, 2008.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(ref)
    set_run_font(run, size=12)

# ── Save ─────────────────────────────────────────────────────────────────────
path = "/home/notcool/Desktop/ntb-event/NTB_Event_MidTerm_Defense_Report.docx"
doc.save(path)
print(f"Saved: {path}")
