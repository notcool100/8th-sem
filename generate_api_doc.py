"""Generate NTB Event API documentation as a DOCX file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── colour palette ──────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1E, 0x40, 0xAF)   # deep blue for headings
METHOD_GET   = RGBColor(0x16, 0xA3, 0x4A)   # green
METHOD_POST  = RGBColor(0x25, 0x63, 0xEB)   # blue
METHOD_PUT   = RGBColor(0xD9, 0x77, 0x06)   # amber
METHOD_DEL   = RGBColor(0xDC, 0x26, 0x26)   # red
BG_LIGHT     = "EFF6FF"                      # very light blue table header
BG_CODE      = "F3F4F6"                      # code block background
TEXT_DARK    = RGBColor(0x11, 0x18, 0x27)
TEXT_MUTED   = RGBColor(0x6B, 0x72, 0x80)

METHOD_COLOR = {
    "GET":    METHOD_GET,
    "POST":   METHOD_POST,
    "PUT":    METHOD_PUT,
    "DELETE": METHOD_DEL,
    "PATCH":  METHOD_PUT,
}

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **edges):
    """Add border to specific edges of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in edges:
            tag = OxmlElement(f"w:{edge}")
            for attr, val in edges[edge].items():
                tag.set(qn(attr), val)
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_paragraph_border_bottom(para):
    """Add a subtle bottom border to a paragraph (used for section dividers)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DBEAFE")
    pBdr.append(bottom)
    pPr.append(pBdr)


def font_of(run, size_pt, bold=False, italic=False, color: RGBColor = None, name="Calibri"):
    run.font.name   = name
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading(doc: Document, text: str, level: int):
    """Add a styled heading paragraph (not using built-in heading styles)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20 if level == 1 else 14 if level == 2 else 8)
    para.paragraph_format.space_after  = Pt(6  if level == 1 else 4)
    run = para.add_run(text)
    if level == 1:
        font_of(run, 22, bold=True, color=BRAND_BLUE)
        add_paragraph_border_bottom(para)
    elif level == 2:
        font_of(run, 16, bold=True, color=BRAND_BLUE)
    elif level == 3:
        font_of(run, 13, bold=True, color=TEXT_DARK)
    else:
        font_of(run, 11, bold=True, color=TEXT_DARK)
    return para


def body(doc: Document, text: str, color: RGBColor = None, italic=False, size=10.5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    font_of(run, size, italic=italic, color=color or TEXT_DARK)
    return para


def code_block(doc: Document, text: str):
    """Shaded monospaced paragraph for JSON / path examples."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent   = Cm(0.6)
    para.paragraph_format.space_before  = Pt(4)
    para.paragraph_format.space_after   = Pt(4)
    run = para.add_run(text)
    run.font.name   = "Courier New"
    run.font.size   = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # shade the paragraph background
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  BG_CODE)
    pPr.append(shd)
    return para


def method_badge(doc: Document, method: str, path: str, summary: str, auth: str = ""):
    """Render  [METHOD]  /path  summary  in one paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(2)

    # coloured method badge
    r_method = para.add_run(f" {method} ")
    r_method.font.name  = "Calibri"
    r_method.font.size  = Pt(10)
    r_method.font.bold  = True
    r_method.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # highlight colour (shading trick via XML)
    rPr  = r_method._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    # We'll use a coloured background via rPrChange trick — simpler: just color the text
    r_method.font.color.rgb = METHOD_COLOR.get(method, TEXT_DARK)
    r_method.font.bold = True

    # path
    r_path = para.add_run(f"  {path}")
    r_path.font.name  = "Courier New"
    r_path.font.size  = Pt(10.5)
    r_path.font.bold  = True
    r_path.font.color.rgb = TEXT_DARK

    # summary
    if summary:
        r_sum = para.add_run(f"   —   {summary}")
        font_of(r_sum, 10, italic=True, color=TEXT_MUTED)

    # auth note
    if auth:
        r_auth = para.add_run(f"   [{auth}]")
        font_of(r_auth, 9, color=RGBColor(0x92, 0x40, 0x0E))

    return para


def field_table(doc: Document, columns, rows):
    """
    columns: list of header strings
    rows: list of tuples matching columns
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = table.rows[0]
    for i, col in enumerate(columns):
        cell = hdr.cells[i]
        set_cell_bg(cell, BG_LIGHT)
        p   = cell.paragraphs[0]
        run = p.add_run(col)
        font_of(run, 9.5, bold=True, color=BRAND_BLUE)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            font_of(run, 9.5)
            if ci == 0:   # field name
                run.font.name = "Courier New"
                run.font.size = Pt(9)
    doc.add_paragraph()   # spacing after table


def status_table(doc: Document, statuses):
    """statuses: list of (code, description)."""
    table = doc.add_table(rows=1 + len(statuses), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, label in enumerate(["HTTP Status", "Meaning"]):
        cell = hdr.cells[i]
        set_cell_bg(cell, BG_LIGHT)
        run = cell.paragraphs[0].add_run(label)
        font_of(run, 9.5, bold=True, color=BRAND_BLUE)
    for ri, (code, desc) in enumerate(statuses):
        row   = table.rows[ri + 1]
        r_code = row.cells[0].paragraphs[0].add_run(str(code))
        r_code.font.name  = "Courier New"
        r_code.font.size  = Pt(9)
        r_desc = row.cells[1].paragraphs[0].add_run(desc)
        font_of(r_desc, 9.5)
    doc.add_paragraph()


# ─── document setup ──────────────────────────────────────────────────────────

def make_doc() -> Document:
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    return doc


# ─── COVER PAGE ──────────────────────────────────────────────────────────────

def cover(doc: Document):
    # spacer
    for _ in range(5):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("NTB EVENT")
    font_of(run, 36, bold=True, color=BRAND_BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("REST API Documentation")
    font_of(run, 22, color=TEXT_MUTED)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run("Version 1.0   ·   June 2026")
    font_of(run, 12, italic=True, color=TEXT_MUTED)

    doc.add_page_break()


# ─── TOC (manual) ────────────────────────────────────────────────────────────

def toc(doc: Document):
    heading(doc, "Table of Contents", 1)
    entries = [
        ("1.", "Overview"),
        ("2.", "Base URL & Authentication"),
        ("3.", "Common Response Envelope"),
        ("4.", "Auth — /api/auth"),
        ("5.", "Events — /api/events"),
        ("6.", "Categories — /api/categories"),
        ("7.", "Tags — /api/tags"),
        ("8.", "Invitations — /api/…/invitations"),
        ("9.", "Approvals — /api/approvals"),
        ("10.", "Notifications — /api/notifications"),
        ("11.", "Upload — /api/upload"),
        ("12.", "Calendar Filter Settings — /api/calendar-filter-settings"),
        ("13.", "Real-Time Hub (SignalR)"),
        ("14.", "Health Check"),
        ("15.", "Enumerations"),
        ("16.", "Error Handling"),
    ]
    for num, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{num}  {title}")
        font_of(r1, 11, color=BRAND_BLUE)
    doc.add_page_break()


# ─── SECTION 1: Overview ─────────────────────────────────────────────────────

def section_overview(doc: Document):
    heading(doc, "1. Overview", 1)
    body(doc,
         "The NTB Event API is a RESTful JSON API that powers the NTB (Nepal Tourism Board / National Tourism Board) "
         "event management platform. It provides endpoints for managing events, user accounts, invitations, "
         "approvals, notifications, file uploads and calendar-filter settings. A SignalR WebSocket hub delivers "
         "real-time notifications to connected admin clients.")
    body(doc,
         "All response bodies follow a uniform envelope structure (see Section 3). Dates are expressed in ISO 8601 "
         "UTC unless otherwise noted. The Bikram Sambat (BS) date fields contain human-readable Nepali-calendar "
         "strings (e.g. \"2082-02-28\").")


# ─── SECTION 2: Base URL & Auth ──────────────────────────────────────────────

def section_base_url(doc: Document):
    heading(doc, "2. Base URL & Authentication", 1)

    heading(doc, "2.1  Base URL", 2)
    body(doc, "Development default:")
    code_block(doc, "http://localhost:5000")
    body(doc, "Production base URL is configured via environment / deployment settings.")

    heading(doc, "2.2  Authentication", 2)
    body(doc,
         "The API uses JWT Bearer tokens. Most endpoints require the caller to supply a valid access token "
         "in the Authorization request header:")
    code_block(doc, "Authorization: Bearer <access_token>")
    body(doc,
         "Tokens are obtained via POST /api/auth/login and refreshed via POST /api/auth/refresh. "
         "Short-lived access tokens expire after a configurable period; long-lived refresh tokens are used to "
         "obtain new access tokens without re-authenticating.")

    heading(doc, "2.3  Role-Based Access Control", 2)
    body(doc, "Three built-in roles exist:")
    field_table(doc,
        ["Role", "Description"],
        [
            ("SuperAdmin", "Unrestricted access to all endpoints. Can approve/publish events directly without a review step."),
            ("Admin",      "Access to most management endpoints. Publishing events may require SuperAdmin approval depending on per-user permission settings."),
            ("(standard)", "No special role. Can access the authenticated /api/auth/me endpoint only."),
        ]
    )
    body(doc,
         "Fine-grained per-user permissions (CanView / CanCreate / CanUpdate / CanDelete / NeedsApproval) "
         "are set by a SuperAdmin via PUT /api/users/{id}/permissions.")

    heading(doc, "2.4  SignalR Token", 2)
    body(doc,
         "The SignalR hub at /hubs/notifications accepts the access token via the access_token query-string "
         "parameter when WebSocket transport is used (browser security restrictions prevent custom headers "
         "over native WebSocket):")
    code_block(doc, "ws://localhost:5000/hubs/notifications?access_token=<token>")


# ─── SECTION 3: Response Envelope ────────────────────────────────────────────

def section_envelope(doc: Document):
    heading(doc, "3. Common Response Envelope", 1)
    body(doc,
         "Every JSON response is wrapped in the following envelope. The generic type parameter T is "
         "replaced by the actual payload type documented per endpoint.")
    code_block(doc, """{
  "isSuccess":        true | false,
  "message":          "Human-readable message or null",
  "data":             <T> | null,
  "errors":           ["string", ...] | null,
  "requiresApproval": true | false | null
}""")
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("isSuccess",        "boolean",       "true when the operation succeeded."),
            ("message",          "string | null", "Optional plain-text message."),
            ("data",             "T | null",      "The response payload. null on failure."),
            ("errors",           "string[] | null","List of validation / business-rule errors."),
            ("requiresApproval", "boolean | null","true when an event was sent for approval instead of being published immediately."),
        ]
    )

    heading(doc, "3.1  Paginated Response", 2)
    body(doc, "Endpoints that return pages wrap their payload in PagedResult<T>:")
    code_block(doc, """{
  "isSuccess": true,
  "data": {
    "items":      [ <T>, ... ],
    "totalCount": 42,
    "page":       1,
    "pageSize":   20,
    "totalPages": 3,
    "hasPrevious": false,
    "hasNext":    true
  }
}""")


# ─── SECTION 4: Auth ─────────────────────────────────────────────────────────

def section_auth(doc: Document):
    heading(doc, "4. Auth  —  /api/auth", 1)
    body(doc, "Handles login, session refresh, logout and retrieving the current user's profile.")

    # ── 4.1 Login
    heading(doc, "4.1  POST /api/auth/login", 2)
    body(doc, "Authenticate a user and return an access + refresh token pair.", italic=True)
    method_badge(doc, "POST", "/api/auth/login", "Authenticate and issue tokens", "Public")

    heading(doc, "Request Body", 3)
    field_table(doc,
        ["Field", "Type", "Required", "Constraints", "Description"],
        [
            ("email",    "string", "Yes", "Valid email, ≤200 chars", "User's email address."),
            ("password", "string", "Yes", "—",                       "User's password (plaintext over HTTPS)."),
        ]
    )
    code_block(doc, '{\n  "email": "admin@ntb.gov.np",\n  "password": "S3cur3P@ss!"\n}')

    heading(doc, "Response  200 OK  —  ApiResponse<AuthSessionDto>", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("accessToken",              "string",   "Short-lived JWT access token."),
            ("accessTokenExpiresAtUtc",  "datetime", "UTC expiry of the access token."),
            ("refreshToken",             "string",   "Long-lived refresh token."),
            ("refreshTokenExpiresAtUtc", "datetime", "UTC expiry of the refresh token."),
            ("user",                     "AuthUserDto","The authenticated user's profile (see below)."),
        ]
    )

    heading(doc, "AuthUserDto", 4)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",          "long",              "Unique user ID."),
            ("fullName",    "string",            "User's display name."),
            ("email",       "string",            "User's email."),
            ("department",  "string",            "Department the user belongs to."),
            ("role",        "string",            "SuperAdmin | Admin | (empty for standard users)."),
            ("permissions", "UserPermissionDto[]","Fine-grained permission list."),
        ]
    )

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Login successful."),
        (401, "Invalid credentials."),
        (403, "Account is inactive / disabled."),
    ])

    # ── 4.2 Refresh
    heading(doc, "4.2  POST /api/auth/refresh", 2)
    body(doc, "Exchange a refresh token for a new access + refresh token pair.", italic=True)
    method_badge(doc, "POST", "/api/auth/refresh", "Refresh session", "Public")

    heading(doc, "Request Body", 3)
    field_table(doc,
        ["Field", "Type", "Required", "Description"],
        [("refreshToken", "string", "Yes", "The refresh token obtained during login or a previous refresh.")]
    )

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "New session issued."),
        (401, "Refresh token is invalid, expired or revoked."),
        (403, "Account is inactive."),
    ])

    # ── 4.3 Logout
    heading(doc, "4.3  POST /api/auth/logout", 2)
    body(doc, "Revoke the supplied refresh token (server-side token invalidation).", italic=True)
    method_badge(doc, "POST", "/api/auth/logout", "Revoke refresh token", "Public")

    heading(doc, "Request Body", 3)
    field_table(doc,
        ["Field", "Type", "Required", "Description"],
        [("refreshToken", "string | null", "No", "The refresh token to revoke. If omitted, only the current server session is closed.")]
    )

    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Logged out successfully.")])

    # ── 4.4 Me
    heading(doc, "4.4  GET /api/auth/me", 2)
    body(doc, "Return the profile of the currently authenticated user.", italic=True)
    method_badge(doc, "GET", "/api/auth/me", "Get current user", "Authenticated")

    heading(doc, "Response  200 OK  —  ApiResponse<AuthUserDto>", 3)
    body(doc, "Returns the AuthUserDto described in §4.1.")

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Profile returned."),
        (401, "Token missing, invalid or expired."),
    ])


# ─── SECTION 5: Events ───────────────────────────────────────────────────────

def section_events(doc: Document):
    heading(doc, "5. Events  —  /api/events", 1)
    body(doc,
         "Create, read, update and delete events. Public endpoints allow unauthenticated access; "
         "management endpoints are restricted to Admin / SuperAdmin roles.")

    # ── 5.1 GET /api/events
    heading(doc, "5.1  GET /api/events", 2)
    body(doc, "Retrieve events visible to the authenticated user (includes drafts by default).", italic=True)
    method_badge(doc, "GET", "/api/events", "List events", "Authenticated")

    heading(doc, "Query Parameters", 3)
    field_table(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ("category",      "string",  "—",    "Filter by category name."),
            ("fromDate",      "datetime","—",    "Return events on or after this UTC date."),
            ("toDate",        "datetime","—",    "Return events on or before this UTC date."),
            ("region",        "string",  "—",    "Filter by region."),
            ("search",        "string",  "—",    "Free-text search against title/summary."),
            ("status",        "string",  "—",    "Filter by lifecycle status (e.g. published, draft)."),
            ("includeDrafts", "boolean", "true", "Whether to include draft events."),
        ]
    )

    heading(doc, "Response  200 OK  —  ApiResponse<EventDto[]>", 3)
    body(doc, "Returns an array of EventDto (see §5.5 for the full schema).")

    # ── 5.2 GET /api/events/public
    heading(doc, "5.2  GET /api/events/public", 2)
    body(doc, "Same as §5.1 but publicly accessible; always returns only published events.", italic=True)
    method_badge(doc, "GET", "/api/events/public", "List published events", "Public")

    heading(doc, "Query Parameters", 3)
    field_table(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ("category", "string",  "—", "Filter by category name."),
            ("fromDate", "datetime","—", "Start of date range."),
            ("toDate",   "datetime","—", "End of date range."),
            ("region",   "string",  "—", "Filter by region."),
            ("search",   "string",  "—", "Free-text search."),
        ]
    )

    # ── 5.3 GET /api/events/public/list  (paged)
    heading(doc, "5.3  GET /api/events/public/list", 2)
    body(doc, "Paginated list of published events with sort support.", italic=True)
    method_badge(doc, "GET", "/api/events/public/list", "Paginated public events", "Public")

    heading(doc, "Query Parameters", 3)
    field_table(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ("category", "string",  "—",    "Filter by category."),
            ("fromDate", "datetime","—",    "Start of date range."),
            ("toDate",   "datetime","—",    "End of date range."),
            ("region",   "string",  "—",    "Filter by region."),
            ("search",   "string",  "—",    "Free-text search."),
            ("page",     "int",     "1",    "Page number (1-based)."),
            ("pageSize", "int",     "20",   "Items per page (1–100)."),
            ("sortBy",   "string",  "date", "Sort field: date | title."),
            ("sortDir",  "string",  "asc",  "Sort direction: asc | desc."),
        ]
    )

    heading(doc, "Response  200 OK  —  ApiResponse<PagedResult<EventDto>>", 3)
    body(doc, "See §3.1 for PagedResult wrapper fields.")

    # ── 5.4 GET /api/events/{id}
    heading(doc, "5.4  GET /api/events/{id}", 2)
    body(doc, "Retrieve a single event by its integer ID.", italic=True)
    method_badge(doc, "GET", "/api/events/{id}", "Get event by ID", "Public")

    heading(doc, "Path Parameters", 3)
    field_table(doc, ["Parameter", "Type", "Description"], [("id", "long", "The unique event identifier.")])

    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Event returned."), (404, "Event not found.")])

    # ── 5.5 EventDto schema
    heading(doc, "5.5  EventDto Schema", 2)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",               "long",              "Unique identifier."),
            ("slug",             "string",            "URL-friendly identifier."),
            ("title",            "string",            "Event title (≤160 chars)."),
            ("date_ad",          "datetime",          "Start date/time (AD, UTC)."),
            ("end_date_ad",      "datetime",          "End date/time (AD, UTC)."),
            ("date_bs",          "string",            "Start date in Bikram Sambat (e.g. \"2082-02-28\")."),
            ("end_date_bs",      "string",            "End date in Bikram Sambat."),
            ("color",            "string",            "Hex colour string for UI theming."),
            ("category",         "string",            "Category name."),
            ("type",             "string",            "event | news | etc."),
            ("status",           "string",            "Lifecycle status (see §15)."),
            ("summary",          "string",            "Short description (≤220 chars)."),
            ("longDescription",  "string",            "Full rich-text / HTML description."),
            ("location",         "string",            "Venue name (≤120 chars)."),
            ("region",           "string",            "Geographic region (≤120 chars)."),
            ("address",          "string",            "Full address (≤200 chars)."),
            ("dateRangeLabel",   "string",            "Human-friendly date range label."),
            ("durationLabel",    "string",            "Duration display string."),
            ("attendanceLabel",  "string",            "Expected attendance label."),
            ("attendanceNote",   "string",            "Additional attendance information."),
            ("entryType",        "string",            "Free / Paid / Registration."),
            ("price",            "decimal",           "Ticket price (0 for free events)."),
            ("rating",           "decimal",           "Average rating (0–5)."),
            ("reviewsLabel",     "string",            "Reviews count label."),
            ("tags",             "string[]",          "Associated tag names."),
            ("image",            "string[]",          "Image URL list."),
            ("mapImage",         "string",            "Map image URL."),
            ("organizer",        "string",            "Organizer name."),
            ("organizerSubtitle","string",            "Organizer secondary line."),
            ("organizerVerified","boolean",           "Whether the organizer is verified."),
            ("highlights",       "EventHighlightDto[]","Key highlights (see below)."),
            ("featured",         "boolean",           "Whether the event is featured/pinned."),
            ("readTime",         "string",            "Estimated reading time label."),
        ]
    )

    heading(doc, "EventHighlightDto", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("icon",        "string", "Icon identifier or emoji."),
            ("title",       "string", "Highlight heading (≤120 chars)."),
            ("description", "string", "Highlight body text (≤400 chars)."),
            ("tone",        "string", "Visual tone class (e.g. success, info)."),
        ]
    )

    # ── 5.6 POST  Create
    heading(doc, "5.6  POST /api/events", 2)
    body(doc, "Create a new event. Admins may be subject to an approval workflow.", italic=True)
    method_badge(doc, "POST", "/api/events", "Create event", "Admin | SuperAdmin")

    heading(doc, "Approval Workflow", 3)
    body(doc,
         "If an Admin (non-SuperAdmin) sets status to \"published\" and their account has the "
         "NeedsApproval flag set for the events module, the event is saved as \"pendingapproval\" "
         "and an EventApprovalRequest is created. requiresApproval: true is returned in the envelope.")

    heading(doc, "Request Body  —  SaveEventRequest", 3)
    field_table(doc,
        ["Field", "Type", "Required", "Constraints", "Description"],
        [
            ("title",            "string",  "Yes", "≤160",    "Event title."),
            ("summary",          "string",  "Yes", "≤220",    "Short description."),
            ("longDescription",  "string",  "Yes", "—",       "Full HTML/text description."),
            ("category",         "string",  "Yes", "≤80",     "Category name."),
            ("type",             "string",  "Yes", "≤30",     "\"event\" by default."),
            ("date_ad",          "datetime","Yes", "—",       "Start date (UTC)."),
            ("location",         "string",  "Yes", "≤120",    "Venue name."),
            ("region",           "string",  "Yes", "≤120",    "Geographic region."),
            ("organizer",        "string",  "Yes", "≤120",    "Organizer name."),
            ("slug",             "string",  "No",  "≤200",    "Auto-generated if omitted."),
            ("status",           "string",  "No",  "≤30",     "draft | published | pendingapproval."),
            ("end_date_ad",      "datetime","No",  "—",       "End date (UTC)."),
            ("date_bs",          "string",  "No",  "≤20",     "Start date in BS calendar."),
            ("end_date_bs",      "string",  "No",  "≤20",     "End date in BS calendar."),
            ("color",            "string",  "No",  "≤20",     "Hex colour."),
            ("address",          "string",  "No",  "≤200",    "Full address."),
            ("dateRangeLabel",   "string",  "No",  "≤140",    "Date range display."),
            ("durationLabel",    "string",  "No",  "≤120",    "Duration display."),
            ("attendanceLabel",  "string",  "No",  "≤120",    "Attendance display."),
            ("attendanceNote",   "string",  "No",  "—",       "Attendance note."),
            ("entryType",        "string",  "No",  "≤40",     "Free | Paid | Registration."),
            ("price",            "decimal", "No",  "0–1 000 000","Ticket price."),
            ("rating",           "decimal", "No",  "0–5",     "Initial rating."),
            ("reviewsLabel",     "string",  "No",  "≤30",     "Reviews label."),
            ("tags",             "string[]","No",  "—",       "Tag name list."),
            ("image",            "string[]","No",  "—",       "Image URL list."),
            ("mapImage",         "string",  "No",  "URL",     "Map image URL."),
            ("organizerSubtitle","string",  "No",  "≤160",    "Organizer secondary line."),
            ("organizerVerified","boolean", "No",  "—",       "Default: true."),
            ("highlights",       "SaveEventHighlightRequest[]","No","—","Highlight items."),
            ("featured",         "boolean", "No",  "—",       "Pin as featured."),
            ("readTime",         "string",  "No",  "≤40",     "Reading time label."),
        ]
    )

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (201, "Event created successfully."),
        (400, "Validation error."),
        (401, "Not authenticated."),
        (409, "Slug conflict or other business rule violation."),
    ])

    # ── 5.7 PUT  Update
    heading(doc, "5.7  PUT /api/events/{id}", 2)
    body(doc, "Update an existing event. Same approval workflow as creation applies.", italic=True)
    method_badge(doc, "PUT", "/api/events/{id}", "Update event", "Admin | SuperAdmin")

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Event updated."),
        (400, "Validation error."),
        (404, "Event not found."),
        (409, "Slug conflict."),
    ])

    # ── 5.8 DELETE
    heading(doc, "5.8  DELETE /api/events/{id}", 2)
    method_badge(doc, "DELETE", "/api/events/{id}", "Delete event", "Admin | SuperAdmin")

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Event deleted."),
        (401, "Not authenticated."),
        (404, "Event not found."),
        (409, "Event has dependent records that prevent deletion."),
    ])


# ─── SECTION 6: Categories ───────────────────────────────────────────────────

def section_categories(doc: Document):
    heading(doc, "6. Categories  —  /api/categories", 1)
    body(doc, "Manage event categories. GET is public; write operations require authentication.")

    heading(doc, "6.1  GET /api/categories", 2)
    method_badge(doc, "GET", "/api/categories", "List categories", "Public")
    heading(doc, "Response  200 OK  —  ApiResponse<CategoriesDto[]>", 3)

    heading(doc, "CategoriesDto", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",          "long",     "Unique category identifier."),
            ("name",        "string",   "Category name."),
            ("description", "string",   "Category description."),
            ("color",       "string",   "Colour used in the UI."),
            ("icon",        "string",   "Icon identifier."),
            ("tag",         "string[]", "Associated tag names."),
        ]
    )

    heading(doc, "6.2  POST /api/categories", 2)
    method_badge(doc, "POST", "/api/categories", "Create category", "Authenticated")
    heading(doc, "Request Body  —  CategoriesDto", 3)
    body(doc, "Supply all CategoriesDto fields (id is ignored on creation).")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Category created.")])

    heading(doc, "6.3  PUT /api/categories/{id}", 2)
    method_badge(doc, "PUT", "/api/categories/{id}", "Update category", "Authenticated")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Category updated.")])

    heading(doc, "6.4  DELETE /api/categories/{id}", 2)
    method_badge(doc, "DELETE", "/api/categories/{id}", "Delete category", "Authenticated")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Category deleted.")])


# ─── SECTION 7: Tags ─────────────────────────────────────────────────────────

def section_tags(doc: Document):
    heading(doc, "7. Tags  —  /api/tags", 1)
    body(doc, "Returns the global list of tag names. No authentication required.")

    heading(doc, "7.1  GET /api/tags", 2)
    method_badge(doc, "GET", "/api/tags", "List all tags", "Public")
    heading(doc, "Response  200 OK  —  ApiResponse<string[]>", 3)
    body(doc, "An array of tag name strings.")
    code_block(doc, '{\n  "isSuccess": true,\n  "data": ["culture", "sports", "tourism", "tech"]\n}')


# ─── SECTION 8: Invitations ──────────────────────────────────────────────────

def section_invitations(doc: Document):
    heading(doc, "8. Invitations  —  /api/…/invitations", 1)
    body(doc,
         "Invite guests to events, track attendance via QR codes, and manage invitation lifecycle "
         "(sent → verified / cancelled). All write operations require Admin or SuperAdmin role.")

    # 8.1 Invite
    heading(doc, "8.1  POST /api/events/{eventId}/invitations", 2)
    body(doc, "Invite a guest to an event. Generates a unique token + QR code and sends an email.", italic=True)
    method_badge(doc, "POST", "/api/events/{eventId}/invitations", "Invite guest", "Admin | SuperAdmin")

    heading(doc, "Path Parameters", 3)
    field_table(doc, ["Parameter","Type","Description"], [("eventId","long","The target event.")])

    heading(doc, "Request Body  —  InviteGuestRequest", 3)
    field_table(doc,
        ["Field", "Type", "Required", "Constraints", "Description"],
        [
            ("fullName",     "string",   "Yes", "≤160",    "Guest's full name."),
            ("email",        "string",   "Yes", "email ≤200","Guest's email (invitation sent here)."),
            ("phone",        "string",   "No",  "≤40",     "Guest's phone number."),
            ("organization", "string",   "No",  "≤160",    "Guest's organization."),
            ("expiresAtUtc", "datetime", "No",  "—",       "Custom QR expiry. Defaults to event end date."),
        ]
    )

    heading(doc, "Response  201 Created  —  ApiResponse<InvitationDto>", 3)
    heading(doc, "InvitationDto", 4)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",               "long",     "Invitation ID."),
            ("eventId",          "long",     "Associated event."),
            ("eventTitle",       "string",   "Event title snapshot."),
            ("guestId",          "long",     "Guest record ID."),
            ("guestName",        "string",   "Guest's full name."),
            ("guestEmail",       "string",   "Guest's email."),
            ("guestPhone",       "string",   "Guest's phone."),
            ("guestOrganization","string",   "Guest's organization."),
            ("token",            "string",   "Opaque token encoded in the QR."),
            ("status",           "string",   "sent | verified | cancelled."),
            ("inviteUrl",        "string",   "Absolute URL for the guest landing page."),
            ("qrDataUri",        "string",   "PNG image as data URI (image/png;base64)."),
            ("sentAtUtc",        "datetime?","When the email was sent."),
            ("expiresAtUtc",     "datetime?","When the QR expires."),
            ("verifiedAtUtc",    "datetime?","When the guest was checked in."),
            ("createdAtUtc",     "datetime", "Record creation time."),
        ]
    )
    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (201, "Invitation created and email dispatched."),
        (400, "Validation error."),
        (404, "Event not found."),
        (409, "Guest is already invited to this event."),
    ])

    # 8.2 List for event
    heading(doc, "8.2  GET /api/events/{eventId}/invitations", 2)
    method_badge(doc, "GET", "/api/events/{eventId}/invitations", "List event invitations", "Admin | SuperAdmin")
    body(doc, "Returns all invitations for the given event.")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "List returned.")])

    # 8.3 Get by ID
    heading(doc, "8.3  GET /api/invitations/{id}", 2)
    method_badge(doc, "GET", "/api/invitations/{id}", "Get invitation", "Admin | SuperAdmin")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Invitation found."), (404, "Not found.")])

    # 8.4 Get by token (public)
    heading(doc, "8.4  GET /api/invitations/by-token/{token}", 2)
    body(doc, "Public endpoint used by the guest landing page to display invitation details.", italic=True)
    method_badge(doc, "GET", "/api/invitations/by-token/{token}", "Get by token", "Public")
    heading(doc, "Path Parameters", 3)
    field_table(doc, ["Parameter","Type","Description"], [("token","string","The opaque invitation token from the QR code or email link.")])
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Invitation returned."), (404, "Invalid / expired token.")])

    # 8.5 QR
    heading(doc, "8.5  GET /api/invitations/{id}/qr", 2)
    body(doc, "Returns the QR code as a raw PNG file (Content-Type: image/png).", italic=True)
    method_badge(doc, "GET", "/api/invitations/{id}/qr", "Download QR image", "Admin | SuperAdmin")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "PNG binary returned."), (404, "Invitation not found.")])

    # 8.6 Resend
    heading(doc, "8.6  POST /api/invitations/{id}/resend", 2)
    body(doc, "Re-send the invitation email (with the link and QR attachment).", italic=True)
    method_badge(doc, "POST", "/api/invitations/{id}/resend", "Resend invitation", "Admin | SuperAdmin")
    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Email re-sent."),
        (404, "Invitation not found."),
        (409, "Invitation is already cancelled or expired."),
    ])

    # 8.7 Scan
    heading(doc, "8.7  POST /api/invitations/scan", 2)
    body(doc,
         "Scan a QR token at the door. Returns guest details for a confirmation popup. "
         "Does NOT check the guest in — call /verify after confirming.", italic=True)
    method_badge(doc, "POST", "/api/invitations/scan", "Scan QR token", "Admin | SuperAdmin")

    heading(doc, "Request Body  —  ScanRequest", 3)
    field_table(doc, ["Field","Type","Required","Description"], [("token","string","Yes","The token read from the QR code.")])

    heading(doc, "Response  200 OK  —  ApiResponse<ScanResultDto>", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("result",     "string",      "One of: verified | previewed | alreadyused | expired | cancelled | invalid."),
            ("message",    "string",      "Human-readable message for the confirmation UI."),
            ("canVerify",  "boolean",     "True if the QR is valid and the admin may proceed with check-in."),
            ("invitation", "InvitationDto?","Full invitation record (null if token is invalid)."),
        ]
    )

    # 8.8 Verify
    heading(doc, "8.8  POST /api/invitations/{id}/verify", 2)
    body(doc, "Confirm check-in: marks the invitation as verified and expires the QR.", italic=True)
    method_badge(doc, "POST", "/api/invitations/{id}/verify", "Verify / check-in", "Admin | SuperAdmin")
    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Guest checked in (ScanResultDto returned)."),
        (404, "Invitation not found."),
    ])

    # 8.9 Cancel
    heading(doc, "8.9  DELETE /api/invitations/{id}", 2)
    body(doc, "Revoke an invitation. The QR becomes invalid and a cancellation email is sent.", italic=True)
    method_badge(doc, "DELETE", "/api/invitations/{id}", "Cancel invitation", "Admin | SuperAdmin")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Invitation cancelled."), (404, "Not found.")])


# ─── SECTION 9: Approvals ────────────────────────────────────────────────────

def section_approvals(doc: Document):
    heading(doc, "9. Approvals  —  /api/approvals", 1)
    body(doc,
         "Manage the event publish-approval workflow. When an Admin submits an event for publication "
         "and NeedsApproval is set, a pending EventApprovalRequest is created. A SuperAdmin or Admin "
         "with sufficient privileges then approves or rejects it.")

    heading(doc, "9.1  GET /api/approvals", 2)
    method_badge(doc, "GET", "/api/approvals", "List all approval requests", "Admin | SuperAdmin")
    heading(doc, "Response  200 OK  —  ApiResponse<EventApprovalDto[]>", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",              "long",     "Approval record ID."),
            ("eventId",         "long",     "Associated event ID."),
            ("eventTitle",      "string",   "Event title snapshot."),
            ("eventSlug",       "string",   "Event slug snapshot."),
            ("requestedByName", "string",   "Full name of the submitter."),
            ("action",          "string",   "Create | Update."),
            ("status",          "string",   "Pending | Approved | Rejected."),
            ("rejectionReason", "string?",  "Set when rejected."),
            ("reviewedByName",  "string?",  "Full name of the reviewer."),
            ("requestedAtUtc",  "datetime", "When the request was submitted."),
            ("reviewedAtUtc",   "datetime?","When the request was reviewed."),
        ]
    )

    heading(doc, "9.2  POST /api/approvals/{id}/approve", 2)
    method_badge(doc, "POST", "/api/approvals/{id}/approve", "Approve event", "Admin | SuperAdmin")
    body(doc, "Publishes the event and marks the approval as Approved. No request body required.")
    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Event approved and published."),
        (400, "Approval already processed."),
        (404, "Approval record not found."),
    ])

    heading(doc, "9.3  POST /api/approvals/{id}/reject", 2)
    method_badge(doc, "POST", "/api/approvals/{id}/reject", "Reject event", "Admin | SuperAdmin")
    heading(doc, "Request Body  —  RejectApprovalRequest", 3)
    field_table(doc, ["Field","Type","Required","Description"], [("rejectionReason","string","Yes","Reason shown to the submitter.")])
    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "Event rejected."),
        (400, "Validation error or approval already processed."),
    ])


# ─── SECTION 10: Notifications ───────────────────────────────────────────────

def section_notifications(doc: Document):
    heading(doc, "10. Notifications  —  /api/notifications", 1)
    body(doc,
         "In-app notification management. Notifications are created server-side and pushed to connected "
         "clients via SignalR (see §13). Each user only sees their own notifications.")

    heading(doc, "10.1  GET /api/notifications", 2)
    method_badge(doc, "GET", "/api/notifications", "Get current user's notifications", "Authenticated")
    heading(doc, "Response  200 OK  —  ApiResponse<NotificationDto[]>", 3)
    field_table(doc,
        ["Field", "Type", "Description"],
        [
            ("id",           "long",    "Notification ID."),
            ("type",         "string",  "Notification type key (e.g. approval_request)."),
            ("title",        "string",  "Short notification title."),
            ("message",      "string",  "Notification body."),
            ("link",         "string?", "Optional deep-link URL."),
            ("isRead",       "boolean", "Whether the notification has been read."),
            ("createdAtUtc", "datetime","When the notification was created."),
        ]
    )

    heading(doc, "10.2  POST /api/notifications/{id}/read", 2)
    method_badge(doc, "POST", "/api/notifications/{id}/read", "Mark notification as read", "Authenticated")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Marked as read.")])

    heading(doc, "10.3  POST /api/notifications/read-all", 2)
    method_badge(doc, "POST", "/api/notifications/read-all", "Mark all notifications as read", "Authenticated")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "All notifications marked as read.")])

    heading(doc, "10.4  DELETE /api/notifications/{id}", 2)
    method_badge(doc, "DELETE", "/api/notifications/{id}", "Delete a notification", "Authenticated")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Notification deleted.")])


# ─── SECTION 11: Upload ──────────────────────────────────────────────────────

def section_upload(doc: Document):
    heading(doc, "11. Upload  —  /api/upload", 1)
    body(doc,
         "Upload image or other binary files. Returns a relative URL that can be placed in the "
         "image[] field of an event or elsewhere in the UI. Files are stored in wwwroot/uploads/.")

    heading(doc, "11.1  POST /api/upload", 2)
    method_badge(doc, "POST", "/api/upload", "Upload a file", "Public (no auth required)")
    body(doc, "Content-Type: multipart/form-data")

    heading(doc, "Form Fields", 3)
    field_table(doc, ["Field","Type","Required","Description"], [("file","IFormFile","Yes","The file to upload.")])

    heading(doc, "Response  200 OK  —  ApiResponse<string>", 3)
    body(doc, "The data field contains the relative URL of the uploaded file:")
    code_block(doc, '{\n  "isSuccess": true,\n  "data": "/uploads/3ca8caa8-1c1a-4d8f-9272-7ee822fc7268.png"\n}')

    heading(doc, "Status Codes", 3)
    status_table(doc, [
        (200, "File uploaded. Relative URL returned."),
        (400, "No file supplied or file is empty."),
        (500, "Server-side I/O error."),
    ])


# ─── SECTION 12: Calendar Filter Settings ────────────────────────────────────

def section_calendar(doc: Document):
    heading(doc, "12. Calendar Filter Settings  —  /api/calendar-filter-settings", 1)
    body(doc,
         "Controls which filter sections are visible on the public event calendar. "
         "GET is public; PUT requires authentication.")

    heading(doc, "CalendarFilterSettingsDto", 2)
    field_table(doc,
        ["Field", "Type", "Default", "Description"],
        [
            ("showCategory",  "boolean", "true", "Show the category filter section."),
            ("showDateRange", "boolean", "true", "Show the date-range filter section."),
            ("showLocation",  "boolean", "true", "Show the location/region filter section."),
            ("showPrice",     "boolean", "true", "Show the price/entry-type filter section."),
            ("showTags",      "boolean", "true", "Show the tags filter section."),
        ]
    )

    heading(doc, "12.1  GET /api/calendar-filter-settings", 2)
    method_badge(doc, "GET", "/api/calendar-filter-settings", "Get filter visibility settings", "Public")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Settings returned.")])

    heading(doc, "12.2  PUT /api/calendar-filter-settings", 2)
    method_badge(doc, "PUT", "/api/calendar-filter-settings", "Update filter visibility settings", "Authenticated")
    body(doc, "Request body is a CalendarFilterSettingsDto with all fields.")
    heading(doc, "Status Codes", 3)
    status_table(doc, [(200, "Settings updated.")])


# ─── SECTION 13: SignalR Hub ─────────────────────────────────────────────────

def section_signalr(doc: Document):
    heading(doc, "13. Real-Time Hub (SignalR)", 1)
    body(doc,
         "The server exposes a SignalR hub at /hubs/notifications for pushing real-time notifications "
         "to authenticated admin clients.")

    heading(doc, "Connection URL", 2)
    code_block(doc, "ws://localhost:5000/hubs/notifications?access_token=<jwt>")

    heading(doc, "Authentication", 2)
    body(doc,
         "Supply the JWT access token as the access_token query parameter. The hub requires [Authorize].")

    heading(doc, "Groups", 2)
    field_table(doc,
        ["Group", "Members", "Description"],
        [("admins","Admin and SuperAdmin role users","Receive approval-request notifications when an event is submitted for review.")]
    )

    heading(doc, "Server → Client Events", 2)
    field_table(doc,
        ["Event Name", "Payload", "Description"],
        [
            ("ReceiveNotification", "NotificationDto (JSON)", "Fired when a new notification is created for the connected user."),
        ]
    )

    heading(doc, "Client → Server Methods", 2)
    body(doc,
         "No public client-to-server methods are defined. The hub is read-only from the client side; "
         "all management actions go through the REST API.")


# ─── SECTION 14: Health Check ────────────────────────────────────────────────

def section_health(doc: Document):
    heading(doc, "14. Health Check", 1)
    method_badge(doc, "GET", "/health", "Liveness probe", "Public")
    body(doc, "Returns HTTP 200 with a JSON body when the service is running:")
    code_block(doc, '{ "status": "ok" }')
    body(doc, "Use this endpoint for load-balancer / Kubernetes liveness probes.")


# ─── SECTION 15: Enumerations ────────────────────────────────────────────────

def section_enums(doc: Document):
    heading(doc, "15. Enumerations", 1)

    heading(doc, "15.1  EventLifecycleStatus  (status field on EventDto)", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("draft",           "Event is saved but not published or submitted."),
            ("pendingapproval", "Submitted for admin review; not publicly visible."),
            ("published",       "Publicly visible on the event calendar."),
        ]
    )

    heading(doc, "15.2  ApprovalStatus  (status field on EventApprovalDto)", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("Pending",  "Awaiting a reviewer's decision."),
            ("Approved", "Event was approved and published."),
            ("Rejected", "Event was rejected with a reason."),
        ]
    )

    heading(doc, "15.3  ApprovalAction  (action field on EventApprovalDto)", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("Create", "A new event was submitted for approval."),
            ("Update", "An edit to an existing event was submitted for approval."),
        ]
    )

    heading(doc, "15.4  InvitationStatus  (status field on InvitationDto)", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("sent",      "Email dispatched; QR not yet scanned."),
            ("verified",  "Guest checked in at the door."),
            ("cancelled", "Invitation was revoked by an admin."),
        ]
    )

    heading(doc, "15.5  ScanResult  (result field on ScanResultDto)", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("previewed",   "Token is valid and the QR has not been used yet."),
            ("verified",    "Check-in confirmed."),
            ("alreadyused", "The QR was already scanned and the guest was checked in."),
            ("expired",     "The QR token has passed its expiry date."),
            ("cancelled",   "The invitation was cancelled before scan."),
            ("invalid",     "Token does not match any known invitation."),
        ]
    )

    heading(doc, "15.6  UserRole", 2)
    field_table(doc,
        ["Value", "Description"],
        [
            ("SuperAdmin", "Unrestricted platform administrator."),
            ("Admin",      "Management user, may be subject to approval workflow."),
        ]
    )


# ─── SECTION 16: Error Handling ──────────────────────────────────────────────

def section_errors(doc: Document):
    heading(doc, "16. Error Handling", 1)
    body(doc,
         "All errors are returned using the standard ApiResponse envelope with isSuccess: false. "
         "The message field contains a human-readable description; the errors array (when present) "
         "lists individual validation failures.")

    code_block(doc, """{
  "isSuccess": false,
  "message": "Validation failed.",
  "errors": [
    "The Title field is required.",
    "Email is not a valid email address."
  ]
}""")

    heading(doc, "Common HTTP Status Codes", 2)
    field_table(doc,
        ["Code", "Meaning"],
        [
            ("200 OK",                  "Operation succeeded."),
            ("201 Created",             "Resource created. Location header points to the new resource."),
            ("400 Bad Request",         "Validation or business-rule error."),
            ("401 Unauthorized",        "Missing or invalid JWT token."),
            ("403 Forbidden",           "Authenticated but insufficient role/permission."),
            ("404 Not Found",           "Requested resource does not exist."),
            ("409 Conflict",            "Duplicate resource or state conflict."),
            ("500 Internal Server Error","Unexpected server-side failure."),
        ]
    )

    heading(doc, "Rate Limiting", 2)
    body(doc,
         "A rate-limiting middleware is configured on the API. Clients that exceed the request quota "
         "will receive HTTP 429 Too Many Requests. Retry after the interval specified in the "
         "Retry-After response header.")


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    doc = make_doc()
    cover(doc)
    toc(doc)
    section_overview(doc)
    section_base_url(doc)
    section_envelope(doc)
    section_auth(doc)
    section_events(doc)
    section_categories(doc)
    section_tags(doc)
    section_invitations(doc)
    section_approvals(doc)
    section_notifications(doc)
    section_upload(doc)
    section_calendar(doc)
    section_signalr(doc)
    section_health(doc)
    section_enums(doc)
    section_errors(doc)

    out_path = "/home/notcool/Desktop/ntb-event/NTB_Event_API_Documentation.docx"
    doc.save(out_path)
    print(f"Saved → {out_path}")


if __name__ == "__main__":
    main()
