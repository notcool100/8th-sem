"""Generate the four Calendar of Events project business documents as DOCX files:

  1. Quotation (dated 10 May 2026)
  2. Accepted Terms of Reference (ToR)
  3. Tax Invoice
  4. Report on Completion of Task

Vendor : Xenith Technology
Client : Nepal Tourism Board (NTB)
Project: Calendar of Events Web Application
Amount : NPR 99,000.00 (VAT inclusive)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── colour palette (Nepal Tourism Board brand — matches existing project docs) ──
TEAL         = RGBColor(0x1C, 0x5C, 0x6D)   # Himalayan Teal — primary
CRIMSON      = RGBColor(0xBD, 0x24, 0x2B)   # Crimson Red — accent / totals
MARIGOLD_HEX = "F8CE1C"                      # Marigold Yellow — fill accent
TEAL_HEX     = "1C5C6D"
BG_LIGHT     = "EAF2F3"
TEXT_DARK    = RGBColor(0x26, 0x30, 0x38)
TEXT_MUTED   = RGBColor(0x6B, 0x72, 0x80)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

OUT_DIR = "/home/notcool/Desktop/ntb-event/business-docs"

# ─── vendor / client / project constants ─────────────────────────────────────
VENDOR_NAME    = "Xenith Technology"
VENDOR_ADDR    = "[Company Address, Kathmandu, Nepal]"
VENDOR_PAN     = "[PAN / VAT Registration No.]"
VENDOR_EMAIL   = "anjaljoshi6@gmail.com"
VENDOR_PHONE   = "[Phone Number]"

CLIENT_NAME    = "Nepal Tourism Board (NTB)"
CLIENT_ADDR    = "Bhrikutimandap, Kathmandu, Nepal"
CLIENT_WEB     = "www.ntb.gov.np"

PROJECT_TITLE  = "Calendar of Events Web Application"

SUBTOTAL = 87610.62
VAT_RATE = 0.13
VAT_AMT  = 11389.38
TOTAL    = 99000.00

LINE_ITEMS = [
    ("1", "Requirement Analysis & UI/UX Design", 10000.00),
    ("2", "Backend API & Database Development", 22000.00),
    ("3", "Frontend Development – Public Calendar of Events & Registration Form", 20000.00),
    ("4", "Admin Panel – Event Management & Approval Workflow", 15000.00),
    ("5", "Testing, Quality Assurance & Bug Fixing", 8610.62),
    ("6", "Deployment, Documentation & Handover", 12000.00),
]

AMOUNT_WORDS = "Nepalese Rupees Ninety-Nine Thousand Only (VAT Inclusive)"


# ─── generic helpers (shared styling) ─────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph_border_bottom(para, color="1C5C6D", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font_of(run, size_pt, bold=False, italic=False, color: RGBColor = None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = para.add_run(text)
    if level == 1:
        font_of(run, 14, bold=True, color=TEAL)
        add_paragraph_border_bottom(para)
    else:
        font_of(run, 11.5, bold=True, color=TEXT_DARK)
    return para


def body(doc: Document, text: str, color: RGBColor = None, italic=False, size=10.5, bold=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    font_of(run, size, italic=italic, bold=bold, color=color or TEXT_DARK)
    return para


def bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    font_of(run, 10.5, color=TEXT_DARK)
    return para


def numbered(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    font_of(run, 10.5, color=TEXT_DARK)
    return para


def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tblPr.append(borders)


def make_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


def header_banner(doc: Document, title: str, subtitle: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_none(table)
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(5)

    left = table.rows[0].cells[0]
    p = left.paragraphs[0]
    r = p.add_run(title)
    font_of(r, 26, bold=True, color=TEAL)
    p2 = left.add_paragraph()
    r2 = p2.add_run(subtitle)
    font_of(r2, 12.5, color=TEXT_MUTED, italic=True)

    right = table.rows[0].cells[1]
    right.vertical_alignment = 1
    p3 = right.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run("Issued by\n")
    font_of(r3, 9.5, color=TEXT_MUTED)
    r4 = p3.add_run(VENDOR_NAME)
    font_of(r4, 13, bold=True, color=CRIMSON)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(14)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MARIGOLD_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def from_to_block(doc: Document, from_label="FROM", to_label="BILL TO"):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders_none(table)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    from_cell = table.rows[0].cells[0]
    p = from_cell.paragraphs[0]
    r = p.add_run(from_label)
    font_of(r, 9, bold=True, color=TEAL)
    for line, bold in [
        (VENDOR_NAME, True),
        (VENDOR_ADDR, False),
        (f"PAN/VAT: {VENDOR_PAN}", False),
        (VENDOR_EMAIL, False),
        (VENDOR_PHONE, False),
    ]:
        pp = from_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(line)
        font_of(rr, 10.5, bold=bold, color=TEXT_DARK)

    to_cell = table.rows[0].cells[1]
    p2 = to_cell.paragraphs[0]
    r2 = p2.add_run(to_label)
    font_of(r2, 9, bold=True, color=TEAL)
    for line, bold in [
        (CLIENT_NAME, True),
        (CLIENT_ADDR, False),
        (CLIENT_WEB, False),
    ]:
        pp = to_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(line)
        font_of(rr, 10.5, bold=bold, color=TEXT_DARK)

    doc.add_paragraph()


def meta_strip(doc: Document, labels):
    meta = doc.add_table(rows=1, cols=len(labels))
    set_table_borders_none(meta)
    for i, (label, value) in enumerate(labels):
        cell = meta.rows[0].cells[i]
        set_cell_bg(cell, BG_LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(label.upper())
        font_of(r, 8, bold=True, color=TEXT_MUTED)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(value)
        font_of(r2, 11, bold=True, color=TEAL)
    doc.add_paragraph()


def cost_table(doc: Document, title="Cost Breakdown", total_label="TOTAL PROJECT COST"):
    heading(doc, title)

    n_rows = 1 + len(LINE_ITEMS) + 3  # header + items + subtotal + vat + total
    table = doc.add_table(rows=n_rows, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(1.4)
    table.columns[1].width = Cm(10.2)
    table.columns[2].width = Cm(3.4)

    hdr = table.rows[0]
    headers = ["S.N.", "Description", "Amount (NPR)"]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, TEAL_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        font_of(r, 10, bold=True, color=WHITE)

    for ri, (sn, desc, amount) in enumerate(LINE_ITEMS):
        row = table.rows[ri + 1]
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_of(p0.add_run(sn), 10, color=TEXT_DARK)

        c1 = row.cells[1]
        font_of(c1.paragraphs[0].add_run(desc), 10, color=TEXT_DARK)

        c2 = row.cells[2]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        font_of(p2.add_run(f"{amount:,.2f}"), 10, color=TEXT_DARK)

        if ri % 2 == 1:
            for c in row.cells:
                set_cell_bg(c, "F7F7F7")

    # subtotal row
    sub_row = table.rows[len(LINE_ITEMS) + 1]
    sub_row.cells[0].merge(sub_row.cells[1])
    sc = sub_row.cells[0]
    p = sc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p.add_run("Subtotal (Excl. VAT)"), 10, bold=True, color=TEXT_DARK)
    sc2 = sub_row.cells[2]
    p2 = sc2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p2.add_run(f"{SUBTOTAL:,.2f}"), 10, bold=True, color=TEXT_DARK)

    # vat row
    vat_row = table.rows[len(LINE_ITEMS) + 2]
    vat_row.cells[0].merge(vat_row.cells[1])
    vc = vat_row.cells[0]
    p = vc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p.add_run("Add: VAT @ 13%"), 10, bold=True, color=TEXT_DARK)
    vc2 = vat_row.cells[2]
    p2 = vc2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p2.add_run(f"{VAT_AMT:,.2f}"), 10, bold=True, color=TEXT_DARK)

    # total row
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[1])
    tc = total_row.cells[0]
    set_cell_bg(tc, BG_LIGHT)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p.add_run(total_label), 11, bold=True, color=TEAL)

    tc2 = total_row.cells[2]
    set_cell_bg(tc2, BG_LIGHT)
    p2 = tc2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p2.add_run(f"NPR {TOTAL:,.2f}"), 12, bold=True, color=CRIMSON)

    doc.add_paragraph()

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    r = note.add_run("Amount in words: ")
    font_of(r, 10, bold=True, color=TEXT_DARK)
    r2 = note.add_run(AMOUNT_WORDS)
    font_of(r2, 10, italic=True, color=TEXT_MUTED)


def signature_block(doc: Document, left_label, right_label):
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    set_table_borders_none(table)

    labels_top = [left_label, right_label]
    for i in range(2):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(30)
        pPr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "9CA3AF")
        border.append(b)
        pPr.append(border)

    for i, label in enumerate(labels_top):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(label)
        font_of(r, 9.5, color=TEXT_MUTED, italic=True)


def footer_note(doc: Document, text: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font_of(r, 10, italic=True, color=TEXT_MUTED)


SCOPE_ITEMS = [
    "Public-facing “Calendar of Events” page (month / list view) aligned with the official NTB brand identity.",
    "Online event registration form with field validation and confirmation workflow.",
    "Secure admin panel with authentication and role-based access control for event management.",
    "Full event lifecycle management: create, edit, categorize, tag, and publish events, with an approval workflow.",
    "Responsive, mobile-friendly design across all public and admin pages.",
    "Testing, bug-fixing, deployment-ready build and basic handover documentation.",
]


# ═══════════════════════════════════════════════════════════════════════════
# 1. QUOTATION
# ═══════════════════════════════════════════════════════════════════════════

def build_quotation():
    doc = make_doc()
    header_banner(doc, "QUOTATION", PROJECT_TITLE)
    from_to_block(doc)
    meta_strip(doc, [
        ("Quotation No.", "XT/QTN/2082-83/001"),
        ("Date", "May 10, 2026"),
        ("Valid Until", "June 10, 2026"),
    ])

    heading(doc, "Project Summary")
    body(doc,
         "Design, development and delivery of a public “Calendar of Events” web application for "
         "the Nepal Tourism Board, comprising an event listing/calendar page with an online registration "
         "form, and a secure admin panel for creating, approving and publishing events. The system will "
         "be built to align with the official NTB brand identity and be fully responsive across devices.")

    heading(doc, "Scope of Work")
    for it in SCOPE_ITEMS:
        bullet(doc, it)

    cost_table(doc)

    heading(doc, "Terms & Conditions")
    for it in [
        "This quotation is valid for 30 days from the date of issue.",
        "Payment Terms: 50% advance to commence work, 50% upon final delivery and acceptance.",
        "The quoted amount covers the scope of work listed above. Additional features or major scope "
        "changes requested after project kickoff will be quoted separately.",
        "Estimated delivery timeline: 6–7 weeks from advance payment and finalization of requirements.",
        "Hosting, domain, SMTP/email service and any third-party service costs are not included.",
        "A reasonable warranty/bug-fix period will be provided after delivery.",
    ]:
        bullet(doc, it)

    signature_block(doc, "Prepared By (Xenith Technology)", "Accepted By (Nepal Tourism Board)")
    footer_note(doc, "Thank you for the opportunity to work with the Nepal Tourism Board.")

    out = os.path.join(OUT_DIR, "1_Quotation_CalendarOfEvents.docx")
    doc.save(out)
    print(f"Saved -> {out}")


# ═══════════════════════════════════════════════════════════════════════════
# 2. ACCEPTED TERMS OF REFERENCE
# ═══════════════════════════════════════════════════════════════════════════

def build_tor():
    doc = make_doc()
    header_banner(doc, "TERMS OF REFERENCE", PROJECT_TITLE)
    from_to_block(doc, from_label="SERVICE PROVIDER", to_label="CLIENT")
    meta_strip(doc, [
        ("ToR No.", "XT/TOR/2082-83/001"),
        ("Date", "May 18, 2026"),
        ("Ref. Quotation", "XT/QTN/2082-83/001"),
    ])

    heading(doc, "1. Background")
    body(doc,
         "Nepal Tourism Board (NTB) requires a public “Calendar of Events” web application to "
         "publish, promote and manage tourism-related events, with an online registration mechanism for "
         "attendees. Xenith Technology has been engaged to design, develop and deliver this application "
         "as per Quotation No. XT/QTN/2082-83/001 dated May 10, 2026.")

    heading(doc, "2. Objectives")
    for it in [
        "Provide the public with an easy-to-use, up-to-date calendar of NTB events.",
        "Enable online registration for events through a validated web form.",
        "Provide NTB administrators with tools to create, review, approve and publish events.",
    ]:
        bullet(doc, it)

    heading(doc, "3. Scope of Work")
    for it in SCOPE_ITEMS:
        bullet(doc, it)

    heading(doc, "4. Deliverables")
    for it in [
        "Fully functional Calendar of Events public web pages and registration form.",
        "Admin panel for event and registration management.",
        "Source code and basic technical/handover documentation.",
        "Deployment-ready build.",
    ]:
        bullet(doc, it)

    heading(doc, "5. Project Timeline")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(5)
    rows = [
        ("Phase", "Indicative Duration"),
        ("Requirement Analysis & UI/UX Design", "Week 1"),
        ("Backend & Frontend Development", "Week 2 – 4"),
        ("Admin Panel, Testing & QA", "Week 5 – 6"),
        ("Deployment, Documentation & Handover", "Week 7"),
    ]
    for ri, (a, b_) in enumerate(rows):
        row = table.rows[ri]
        if ri == 0:
            for c, txt in zip(row.cells, (a, b_)):
                set_cell_bg(c, TEAL_HEX)
                font_of(c.paragraphs[0].add_run(txt), 10, bold=True, color=WHITE)
        else:
            font_of(row.cells[0].paragraphs[0].add_run(a), 10, color=TEXT_DARK)
            font_of(row.cells[1].paragraphs[0].add_run(b_), 10, color=TEXT_DARK)
    doc.add_paragraph()

    heading(doc, "6. Roles & Responsibilities")
    body(doc, "Xenith Technology shall:", bold=True, size=10.5)
    for it in [
        "Design, develop, test and deliver the application as per the agreed scope.",
        "Provide reasonable post-delivery bug-fix support during the warranty period.",
    ]:
        bullet(doc, it)
    body(doc, "Nepal Tourism Board shall:", bold=True, size=10.5)
    for it in [
        "Provide timely feedback, content and approvals required for development.",
        "Make payments as per the agreed payment terms.",
        "Provide hosting/domain access where required for deployment.",
    ]:
        bullet(doc, it)

    heading(doc, "7. Payment Terms")
    body(doc, f"Total project cost: NPR {TOTAL:,.2f} (VAT inclusive) — 50% advance, 50% on delivery and acceptance.")

    heading(doc, "8. Acceptance")
    body(doc,
         "By signing below, both parties confirm their acceptance of the terms of reference outlined "
         "in this document as the agreed basis for execution of the project.")

    signature_block(doc, "For Xenith Technology", "For Nepal Tourism Board (NTB)")
    footer_note(doc, "This Terms of Reference is accepted and binding upon signature by both parties.")

    out = os.path.join(OUT_DIR, "2_Accepted_Terms_of_Reference.docx")
    doc.save(out)
    print(f"Saved -> {out}")


# ═══════════════════════════════════════════════════════════════════════════
# 3. TAX INVOICE
# ═══════════════════════════════════════════════════════════════════════════

def build_tax_invoice():
    doc = make_doc()
    header_banner(doc, "TAX INVOICE", PROJECT_TITLE)
    from_to_block(doc, from_label="SELLER", to_label="BUYER")
    meta_strip(doc, [
        ("Invoice No.", "XT/INV/2082-83/001"),
        ("Invoice Date", "June 30, 2026"),
        ("Fiscal Year", "2082/083"),
    ])
    meta_strip(doc, [
        ("Ref. Quotation", "XT/QTN/2082-83/001"),
        ("Ref. ToR", "XT/TOR/2082-83/001"),
        ("Payment Mode", "Bank Transfer"),
    ])

    body(doc,
         f"This is to certify that the amount stated below has been charged for services rendered "
         f"to {CLIENT_NAME} for the {PROJECT_TITLE} project, in accordance with the Value Added Tax "
         f"Act, 2052 of Nepal.")

    cost_table(doc, title="Particulars", total_label="GRAND TOTAL (VAT INCL.)")

    heading(doc, "Payment Details")
    for label, value in [
        ("Bank Name", "[Bank Name]"),
        ("Account Name", "Xenith Technology"),
        ("Account Number", "[XXXXXXXXXXXX]"),
        ("Branch", "[Branch Name]"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{label}: ")
        font_of(r, 10.5, bold=True, color=TEXT_DARK)
        r2 = p.add_run(value)
        font_of(r2, 10.5, color=TEXT_DARK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(
        "This is a computer-generated tax invoice. Seller's PAN/VAT registration number is stated above."
    )
    font_of(r, 9, italic=True, color=TEXT_MUTED)

    signature_block(doc, "Authorized Signature (Xenith Technology)", "Received By (Nepal Tourism Board)")
    footer_note(doc, "Thank you for your business.")

    out = os.path.join(OUT_DIR, "3_Tax_Invoice_CalendarOfEvents.docx")
    doc.save(out)
    print(f"Saved -> {out}")


# ═══════════════════════════════════════════════════════════════════════════
# 4. REPORT ON COMPLETION OF TASK
# ═══════════════════════════════════════════════════════════════════════════

def build_completion_report():
    doc = make_doc()
    header_banner(doc, "COMPLETION REPORT", PROJECT_TITLE)
    from_to_block(doc, from_label="SUBMITTED BY", to_label="SUBMITTED TO")
    meta_strip(doc, [
        ("Report No.", "XT/CR/2082-83/001"),
        ("Date", "June 30, 2026"),
        ("Ref. ToR", "XT/TOR/2082-83/001"),
    ])

    heading(doc, "1. Project Overview")
    body(doc,
         f"This report confirms the successful completion of the {PROJECT_TITLE} project undertaken by "
         f"Xenith Technology for the Nepal Tourism Board, as per the accepted Terms of Reference "
         f"(XT/TOR/2082-83/001) and Quotation (XT/QTN/2082-83/001).")

    heading(doc, "2. Work Completed")
    for it in [
        "Requirement analysis and UI/UX design for the Calendar of Events module finalized.",
        "Backend API and database schema developed and deployed.",
        "Public Calendar of Events page and online event registration form developed and tested.",
        "Admin panel for event creation, approval and management delivered.",
        "End-to-end testing, quality assurance and bug fixing completed.",
        "Application deployed to the designated environment with handover documentation provided.",
    ]:
        bullet(doc, it)

    heading(doc, "3. Deliverables Handed Over")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(4)
    rows = [
        ("Deliverable", "Status"),
        ("Calendar of Events public web application", "Completed"),
        ("Event registration form", "Completed"),
        ("Admin panel for event management", "Completed"),
        ("Source code & handover documentation", "Completed"),
    ]
    for ri, (a, b_) in enumerate(rows):
        row = table.rows[ri]
        if ri == 0:
            for c, txt in zip(row.cells, (a, b_)):
                set_cell_bg(c, TEAL_HEX)
                font_of(c.paragraphs[0].add_run(txt), 10, bold=True, color=WHITE)
        else:
            font_of(row.cells[0].paragraphs[0].add_run(a), 10, color=TEXT_DARK)
            p = row.cells[1].paragraphs[0]
            font_of(p.add_run(b_), 10, bold=True, color=RGBColor(0x15, 0x80, 0x3D))
    doc.add_paragraph()

    heading(doc, "4. Financial Summary")
    body(doc, f"Total project cost invoiced: NPR {TOTAL:,.2f} (VAT inclusive), as per Tax Invoice No. "
              f"XT/INV/2082-83/001 dated June 30, 2026.")

    heading(doc, "5. Remarks")
    body(doc,
         "The project has been completed within the agreed scope and delivered to the satisfaction of "
         "the Nepal Tourism Board. Xenith Technology will provide a reasonable post-delivery support "
         "period for any minor bug fixes as agreed in the Terms of Reference.")

    heading(doc, "6. Acceptance")
    body(doc, "This report is submitted to confirm that all deliverables under the above-referenced "
              "project have been completed and handed over.")

    signature_block(doc, "Submitted By (Xenith Technology)", "Accepted By (Nepal Tourism Board)")
    footer_note(doc, "Thank you for the opportunity to work with the Nepal Tourism Board.")

    out = os.path.join(OUT_DIR, "4_Completion_Report_CalendarOfEvents.docx")
    doc.save(out)
    print(f"Saved -> {out}")


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    build_quotation()
    build_tor()
    build_tax_invoice()
    build_completion_report()


if __name__ == "__main__":
    main()
