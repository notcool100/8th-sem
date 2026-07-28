from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
RED    = RGBColor(0xC8, 0x10, 0x2E)   # NTB brand red
NAVY   = RGBColor(0x10, 0x24, 0x3E)   # dark navy
TEAL   = RGBColor(0x1C, 0x5C, 0x6D)   # secondary teal
GRAY   = RGBColor(0x5A, 0x6D, 0x83)   # muted text
LGRAY  = RGBColor(0xF0, 0xF2, 0xF5)   # light background
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def rgb_hex(rgb):
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])

def shade_paragraph(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    pPr.append(shd)

def shade_cell(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Heading helpers ────────────────────────────────────────────────────────────
def h1(text):
    p   = doc.add_paragraph()
    shade_paragraph(p, NAVY)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = WHITE
    return p

def h2(text):
    p   = doc.add_paragraph()
    shade_paragraph(p, TEAL)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{key}: ")
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = TEAL
    r2 = p.add_run(value)
    r2.font.size      = Pt(10)
    r2.font.color.rgb = NAVY
    return p

def divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C8102E')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def scenario_table(scenarios):
    """
    scenarios: list of dicts with keys:
      id, description, precondition, steps, expected, status
    """
    COL_W = [Cm(1.4), Cm(4.2), Cm(3.0), Cm(5.5), Cm(4.5), Cm(1.8)]
    t = doc.add_table(rows=1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style     = 'Table Grid'

    hdr_cells = t.rows[0].cells
    headers   = ['TC #', 'Test Case', 'Pre-condition', 'Test Steps', 'Expected Result', 'Status']
    for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        shade_cell(cell, NAVY)
        set_cell_border(cell, color='FFFFFF')
        cell.width = COL_W[i]
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold            = True
        run.font.size       = Pt(9)
        run.font.color.rgb  = WHITE
        p.alignment         = WD_ALIGN_PARAGRAPH.CENTER

    for sc in scenarios:
        row   = t.add_row()
        cells = row.cells
        bg    = LGRAY if scenarios.index(sc) % 2 == 0 else WHITE

        vals  = [sc['id'], sc['description'], sc['precondition'],
                 sc['steps'], sc['expected'], sc['status']]
        for i, (cell, val) in enumerate(zip(cells, vals)):
            shade_cell(cell, bg)
            set_cell_border(cell, color='C8D6E5')
            cell.width = COL_W[i]
            p  = cell.paragraphs[0]
            rn = p.add_run(val)
            rn.font.size      = Pt(9)
            rn.font.color.rgb = NAVY
            if i == 5:  # status column
                if val == 'PASS':
                    rn.font.color.rgb = GREEN
                    rn.bold = True
                elif val == 'FAIL':
                    rn.font.color.rgb = RED
                    rn.bold = True
                else:
                    rn.font.color.rgb = AMBER
                    rn.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover_title = doc.add_paragraph()
cover_title.paragraph_format.space_before = Pt(60)
cover_title.paragraph_format.space_after  = Pt(4)
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("NTB EVENT MANAGEMENT SYSTEM")
r.bold           = True
r.font.size      = Pt(22)
r.font.color.rgb = RED

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover_sub.add_run("Quality Assurance Documentation")
r2.font.size      = Pt(15)
r2.font.color.rgb = NAVY
r2.bold           = True

divider()

meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_para.paragraph_format.space_before = Pt(10)
meta_para.add_run(f"Version 1.0   |   Prepared for: QA Team   |   Date: {datetime.date.today().strftime('%B %d, %Y')}").font.color.rgb = GRAY

doc.add_paragraph()
body("This document covers all fully functional modules of the NTB Event Management System. "
     "For each module it describes: the workflow, the pages involved, user roles that can "
     "access it, and detailed test scenarios with expected results.", indent=0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. AUTHENTICATION MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Authentication Module")
h2("1.1 Overview")
kv("URL", "/login")
kv("Roles", "All users (Admin, SuperAdmin)")
kv("Backend Endpoint", "POST /api/auth/login")
body("Users authenticate with email + password. On success, JWT access token and refresh token "
     "are issued. The access token is stored in an HttpOnly cookie and used for all subsequent "
     "API calls. Session data (user ID, role, name, email, accessToken) is returned by the root "
     "SvelteKit layout's server-side load function.")

h2("1.2 Workflow")
bullet("1. User opens /login and enters email + password.")
bullet("2. Form is submitted via SvelteKit action (POST).")
bullet("3. Backend validates credentials, issues JWT access token (15-min) + refresh token (30-day).")
bullet("4. Tokens stored as HttpOnly cookies; user redirected to /admin/dashboard.")
bullet("5. On each page load, root +layout.server.ts reads the cookie and validates the token.")
bullet("6. If token is expired, the system attempts a silent refresh using the refresh token.")
bullet("7. Logout clears both cookies and redirects to /login.")

h2("1.3 Test Scenarios")
scenario_table([
    {"id":"TC-AUTH-01","description":"Valid login with correct credentials",
     "precondition":"User account exists and is active","steps":"1. Go to /login\n2. Enter valid email & password\n3. Click Login",
     "expected":"Redirected to /admin/dashboard. Auth cookie set.","status":"PASS"},
    {"id":"TC-AUTH-02","description":"Login with wrong password",
     "precondition":"User account exists","steps":"1. Go to /login\n2. Enter valid email + wrong password\n3. Click Login",
     "expected":"Error message shown. No redirect. No cookie set.","status":"PASS"},
    {"id":"TC-AUTH-03","description":"Login with non-existent email",
     "precondition":"None","steps":"1. Go to /login\n2. Enter unknown email + any password",
     "expected":"Error: invalid credentials. No cookie set.","status":"PASS"},
    {"id":"TC-AUTH-04","description":"Access admin page without login",
     "precondition":"No auth cookie","steps":"1. Navigate directly to /admin/events",
     "expected":"Redirected to /login.","status":"PASS"},
    {"id":"TC-AUTH-05","description":"Logout clears session",
     "precondition":"User is logged in","steps":"1. Click Logout in sidebar",
     "expected":"Cookies cleared; user redirected to /login.","status":"PASS"},
    {"id":"TC-AUTH-06","description":"Inactive user cannot login",
     "precondition":"User isActive = false","steps":"1. Attempt login with disabled account",
     "expected":"Error message: account inactive or invalid.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. EVENT MANAGEMENT MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Event Management Module")
h2("2.1 Overview")
kv("URL", "/admin/events")
kv("Roles", "Admin, SuperAdmin (RBAC-controlled per action)")
kv("Backend Endpoints", "GET/POST/PUT/DELETE /api/events")
body("Central module for managing tourism events. Supports full CRUD with rich content "
     "(title, description, images, location, dates in AD & BS, highlights, pricing). "
     "Events have a lifecycle: Draft → PendingApproval → Published → Archived. "
     "Users with NeedsApproval flag cannot directly publish — they submit for review.")

h2("2.2 Event Lifecycle / Status Flow")
bullet("Draft         — saved but not visible to public")
bullet("PendingApproval — submitted for admin review (users with NeedsApproval=true)")
bullet("Published     — live and visible on the public site")
bullet("Archived      — hidden from public but retained in system")

h2("2.3 Workflow — List Page (/admin/events)")
bullet("1. Page loads all events; user sees table with Title, Category, Status, Date, Actions.")
bullet("2. Search filters events by title in real time.")
bullet("3. Status filter dropdown narrows by Draft/Published/Archived/Pending Approval.")
bullet("4. Delete button (if canDelete) shows confirmation; calls DELETE /api/events/{id}.")
bullet("5. Edit button navigates to /admin/events/create?id={id}.")
bullet("6. 'Add New Event' button navigates to /admin/events/create (if canCreate).")

h2("2.4 Workflow — Create/Edit Page (/admin/events/create)")
bullet("1. All categories are loaded on page open (GET /api/categories).")
bullet("2. User fills in: Title, Summary, Long Description (rich text), Category, Type, Status.")
bullet("3. Date fields accept both AD and BS formats (Nepali date converter).")
bullet("4. Images uploaded via POST /api/upload (drag-drop or file picker).")
bullet("5. Tags, Location, Region, Address, Price, Rating, Highlights added.")
bullet("6. If user has NeedsApproval permission, the Publish option is replaced by 'Request Approval'.")
bullet("7. On Save: POST /api/events (create) or PUT /api/events/{id} (edit).")
bullet("8. If event is submitted for approval: status = PendingApproval; admins receive real-time notification.")

h2("2.5 Test Scenarios — List Page")
scenario_table([
    {"id":"TC-EVT-01","description":"View events list",
     "precondition":"Events exist in DB; user has canView","steps":"1. Navigate to /admin/events",
     "expected":"Table shows all events with correct status badges.","status":"PASS"},
    {"id":"TC-EVT-02","description":"Search events by title",
     "precondition":"Multiple events exist","steps":"1. Type partial title in search box",
     "expected":"Table filters in real time to matching events.","status":"PASS"},
    {"id":"TC-EVT-03","description":"Filter by status",
     "precondition":"Events with mixed statuses","steps":"1. Select 'Published' from status dropdown",
     "expected":"Only published events shown.","status":"PASS"},
    {"id":"TC-EVT-04","description":"Delete event",
     "precondition":"User has canDelete permission","steps":"1. Click Delete on an event\n2. Confirm",
     "expected":"Event removed from list; success toast shown.","status":"PASS"},
    {"id":"TC-EVT-05","description":"Non-authorized user cannot delete",
     "precondition":"User has canDelete = false","steps":"1. Open events list",
     "expected":"Delete button not visible.","status":"PASS"},
])

h2("2.6 Test Scenarios — Create/Edit Page")
scenario_table([
    {"id":"TC-EVT-06","description":"Create new event (SuperAdmin — no approval needed)",
     "precondition":"Logged in as SuperAdmin","steps":"1. Click Add New Event\n2. Fill all required fields\n3. Select Published\n4. Save",
     "expected":"Event created with status Published. Appears in events list.","status":"PASS"},
    {"id":"TC-EVT-07","description":"Create event with NeedsApproval user",
     "precondition":"User has NeedsApproval=true on events nav item","steps":"1. Fill event form\n2. Note 'Request Approval' replaces Publish\n3. Submit",
     "expected":"Event status = PendingApproval. Admin gets bell notification.","status":"PASS"},
    {"id":"TC-EVT-08","description":"Edit existing event",
     "precondition":"Event exists; user has canUpdate","steps":"1. Click Edit on an event\n2. Change title\n3. Save",
     "expected":"Event updated. Changes reflected in list.","status":"PASS"},
    {"id":"TC-EVT-09","description":"Image upload",
     "precondition":"Valid image file","steps":"1. In create form, drop image into upload area",
     "expected":"Image uploaded; preview shown; URL stored in form.","status":"PASS"},
    {"id":"TC-EVT-10","description":"Missing required fields",
     "precondition":"None","steps":"1. Open create form\n2. Leave Title empty\n3. Click Save",
     "expected":"Validation error shown. Form not submitted.","status":"PASS"},
    {"id":"TC-EVT-11","description":"Date in BS auto-converts to AD",
     "precondition":"None","steps":"1. Enter BS date in date field",
     "expected":"AD equivalent auto-populated.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. APPROVAL WORKFLOW MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Event Approval Workflow Module")
h2("3.1 Overview")
kv("URL", "/admin/events/pending-approval")
kv("Roles", "Admin, SuperAdmin (with approval page access)")
kv("Backend Endpoints", "GET /api/approvals, POST /api/approvals/{id}/approve, POST /api/approvals/{id}/reject")
body("When a user with NeedsApproval=true creates or edits an event, it enters a review queue. "
     "An authorized reviewer can approve (event → Published) or reject (event → original status) "
     "with a mandatory rejection reason. Real-time notifications are sent to the submitter on both outcomes.")

h2("3.2 Workflow")
bullet("1. User with NeedsApproval submits an event → status set to PendingApproval.")
bullet("2. Admins with approval page access receive a bell notification instantly via SignalR.")
bullet("3. Reviewer opens /admin/events/pending-approval — sees queue of pending events.")
bullet("4. Each card shows: event title, submitter, submission date, create/update action.")
bullet("5. Reviewer clicks Approve → event status → Published; submitter notified.")
bullet("6. Reviewer clicks Reject → must enter rejection reason → event reverts to original status; submitter notified.")

h2("3.3 Test Scenarios")
scenario_table([
    {"id":"TC-APR-01","description":"View pending approval queue",
     "precondition":"At least one PendingApproval event exists","steps":"1. Navigate to /admin/events/pending-approval",
     "expected":"Cards showing pending events with submitter info.","status":"PASS"},
    {"id":"TC-APR-02","description":"Approve an event",
     "precondition":"Pending event exists","steps":"1. Click Approve on event card\n2. Confirm",
     "expected":"Event status = Published. Submitter gets 'approved' notification. Card removed from queue.","status":"PASS"},
    {"id":"TC-APR-03","description":"Reject an event with reason",
     "precondition":"Pending event exists","steps":"1. Click Reject\n2. Enter reason\n3. Confirm",
     "expected":"Event reverts to original status. Submitter gets 'rejected' notification with reason.","status":"PASS"},
    {"id":"TC-APR-04","description":"Reject without entering reason",
     "precondition":"Pending event","steps":"1. Click Reject\n2. Leave reason empty\n3. Submit",
     "expected":"Validation error — reason is required.","status":"PASS"},
    {"id":"TC-APR-05","description":"Non-approval-page user cannot access queue",
     "precondition":"User has no access to pending-approval nav item","steps":"1. Navigate to /admin/events/pending-approval",
     "expected":"Access denied / redirected.","status":"PASS"},
    {"id":"TC-APR-06","description":"Approval triggers real-time notification",
     "precondition":"Submitter is logged in on another browser tab","steps":"1. Approve event\n2. Check submitter's bell icon",
     "expected":"Unread count increments immediately without page refresh.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. NOTIFICATION MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Real-Time Notification Module")
h2("4.1 Overview")
kv("URL", "/admin/notifications")
kv("Technology", "SignalR WebSockets + backend persistence")
kv("Backend Endpoints", "GET /api/notifications, POST /api/notifications/{id}/read, POST /api/notifications/read-all, DELETE /api/notifications/{id}")
body("Notifications are pushed in real time via SignalR WebSocket. They are also persisted in "
     "the database so they survive page reloads. The bell icon in the header shows unread count. "
     "Three notification types exist: approval_request (amber), approved (green), rejected (red).")

h2("4.2 Workflow")
bullet("1. On admin layout mount, connectNotifications(accessToken) is called.")
bullet("2. A HubConnection to /hubs/notifications is established with JWT via query param.")
bullet("3. Initial list of notifications (up to 50) is fetched from GET /api/notifications.")
bullet("4. When a new notification arrives, it prepends to the list and updates unread count.")
bullet("5. Clicking the bell opens the dropdown showing recent notifications.")
bullet("6. Clicking a notification marks it read + navigates to its link.")
bullet("7. 'Mark all read' marks everything read in one call.")
bullet("8. Full list with All/Unread tabs is at /admin/notifications.")
bullet("9. Each notification card has a delete (trash) button.")

h2("4.3 Notification Types")
bullet("approval_request — sent to ALL admins when a user submits an event for approval.")
bullet("approved — sent to the event submitter when their event is approved.")
bullet("rejected — sent to the event submitter when their event is rejected.")

h2("4.4 Test Scenarios")
scenario_table([
    {"id":"TC-NOT-01","description":"Real-time notification on event submission",
     "precondition":"Admin is logged in; another user submits for approval","steps":"1. Do not refresh page\n2. Other user submits event for approval",
     "expected":"Bell badge increments immediately. Notification appears in dropdown.","status":"PASS"},
    {"id":"TC-NOT-02","description":"Mark single notification as read",
     "precondition":"Unread notification exists","steps":"1. Click on notification in dropdown",
     "expected":"Notification marked read; unread count decrements.","status":"PASS"},
    {"id":"TC-NOT-03","description":"Mark all read",
     "precondition":"Multiple unread notifications","steps":"1. Click 'Mark all read' button",
     "expected":"All notifications marked read; unread count = 0.","status":"PASS"},
    {"id":"TC-NOT-04","description":"Delete notification",
     "precondition":"Notification exists","steps":"1. Open /admin/notifications\n2. Hover card\n3. Click trash icon",
     "expected":"Notification removed from list.","status":"PASS"},
    {"id":"TC-NOT-05","description":"Notifications persist after page refresh",
     "precondition":"Notifications in DB","steps":"1. Refresh /admin/notifications",
     "expected":"Notifications reload from DB correctly.","status":"PASS"},
    {"id":"TC-NOT-06","description":"Unread tab shows only unread items",
     "precondition":"Mix of read and unread notifications","steps":"1. Click Unread tab",
     "expected":"Only unread notifications shown.","status":"PASS"},
    {"id":"TC-NOT-07","description":"Notification link navigation",
     "precondition":"Notification has a link","steps":"1. Click notification with a link",
     "expected":"Navigates to linked page after marking read.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. USER MANAGEMENT MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("5. User Management Module")
h2("5.1 Overview")
kv("URL", "/admin/users")
kv("Roles", "SuperAdmin only")
kv("Backend Endpoints", "GET /api/users, POST /api/users, PUT /api/users/{id}, DELETE /api/users/{id}, GET /api/users/nav-items, PUT /api/users/{id}/permissions")
body("SuperAdmin manages all admin/staff accounts. Each user is assigned a role (Admin or SuperAdmin), "
     "department, and a permission matrix mapping each nav-item to CRUD flags plus a NeedsApproval toggle. "
     "Permissions drive what each user can see and do in the sidebar and within each module.")

h2("5.2 Workflow — User Creation")
bullet("1. SuperAdmin opens /admin/users.")
bullet("2. Clicks 'Add User' → fills FullName, Email, Department, Role, Password.")
bullet("3. On the same form, expands the Permission Matrix section.")
bullet("4. For each nav item (Dashboard, Events, Calendar, Users, Categories, etc.), toggles: View / Create / Update / Delete / NeedsApproval.")
bullet("5. Saves — user created with hashed password + permissions stored in user_permissions table.")
bullet("6. If Events > NeedsApproval is toggled, that user will get the approval workflow when creating events.")

h2("5.3 Workflow — Edit / Deactivate")
bullet("1. Click Edit on a user row → prefills form.")
bullet("2. Modify any field including permissions → Save.")
bullet("3. Toggle IsActive to deactivate (user cannot login).")
bullet("4. Delete removes the user and cascades to permissions.")

h2("5.4 Permission Matrix — Nav Items")
for item in ["Dashboard","Events Management","Calendar Management","Destinations",
             "Users Management","Department Setup","Attendees / Registrations",
             "Event Check-in","Categories Setup","Reports & Analytics","Notifications","Settings"]:
    bullet(item)

h2("5.5 Test Scenarios")
scenario_table([
    {"id":"TC-USR-01","description":"SuperAdmin creates a new user",
     "precondition":"Logged in as SuperAdmin","steps":"1. Click Add User\n2. Fill all fields\n3. Save",
     "expected":"User created; appears in list.","status":"PASS"},
    {"id":"TC-USR-02","description":"Assign permissions to user",
     "precondition":"User exists","steps":"1. Edit user\n2. Toggle Events > CanCreate\n3. Save",
     "expected":"Permission saved; user can now create events after re-login.","status":"PASS"},
    {"id":"TC-USR-03","description":"Enable NeedsApproval for Events",
     "precondition":"User exists","steps":"1. Edit user\n2. Toggle Events > NeedsApproval ON\n3. Save",
     "expected":"User's event creation routes through approval workflow.","status":"PASS"},
    {"id":"TC-USR-04","description":"Deactivate a user",
     "precondition":"Active user","steps":"1. Edit user\n2. Toggle IsActive OFF\n3. Save",
     "expected":"User cannot login; shows error on login attempt.","status":"PASS"},
    {"id":"TC-USR-05","description":"Non-SuperAdmin cannot access /admin/users",
     "precondition":"Logged in as regular Admin","steps":"1. Navigate to /admin/users",
     "expected":"Access denied / redirect.","status":"PASS"},
    {"id":"TC-USR-06","description":"Delete user",
     "precondition":"User exists","steps":"1. Click Delete\n2. Confirm",
     "expected":"User and permissions removed.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. CATEGORIES MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Categories Setup Module")
h2("6.1 Overview")
kv("URL", "/admin/categories-setup")
kv("Roles", "Admin, SuperAdmin (RBAC)")
kv("Backend Endpoints", "GET/POST/PUT/DELETE /api/categories, GET /api/tags")
body("Categories classify events (e.g., Cultural, Adventure, Holiday). Each category has a name, "
     "description, color, emoji icon, and a set of associated tags. Tags can be re-used across "
     "categories. An audit table (CategoriesTags_jn) tracks all tag changes.")

h2("6.2 Workflow")
bullet("1. Open /admin/categories-setup — all existing categories listed as cards.")
bullet("2. Click 'Add Category' → form opens with name, description, color picker, emoji picker.")
bullet("3. Tag input field allows adding existing or new tags (autocomplete from GET /api/tags).")
bullet("4. Save → POST /api/categories; tags resolved/created and linked via CategoriesTags table.")
bullet("5. Edit → PUT /api/categories/{id}; tags diff applied (old tags removed, new added).")
bullet("6. Delete → DELETE /api/categories/{id}; cascades to CategoriesTags rows.")

h2("6.3 Test Scenarios")
scenario_table([
    {"id":"TC-CAT-01","description":"Create a new category",
     "precondition":"User has canCreate on Categories","steps":"1. Click Add Category\n2. Fill name, description, color, icon\n3. Add tags\n4. Save",
     "expected":"Category appears in list. Tags linked correctly.","status":"PASS"},
    {"id":"TC-CAT-02","description":"Edit category and change tags",
     "precondition":"Category with tags exists","steps":"1. Edit category\n2. Remove one tag, add new one\n3. Save",
     "expected":"Old tag unlinked; new tag linked. No orphan records.","status":"PASS"},
    {"id":"TC-CAT-03","description":"Delete category",
     "precondition":"Category exists","steps":"1. Click Delete\n2. Confirm",
     "expected":"Category and its tag links removed.","status":"PASS"},
    {"id":"TC-CAT-04","description":"Categories load in event create form",
     "precondition":"Categories exist in DB","steps":"1. Open /admin/events/create",
     "expected":"Category dropdown populated from /api/categories.","status":"PASS"},
    {"id":"TC-CAT-05","description":"Tag autocomplete",
     "precondition":"Tags exist","steps":"1. In tag input, type partial tag name",
     "expected":"Matching existing tags suggested.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. INVITATION MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Event Invitation Module")
h2("7.1 Overview")
kv("URL", "/admin/events/invitations/[eventId]")
kv("Roles", "Admin, SuperAdmin (with check-in/attendees access)")
kv("Backend Endpoints", "POST /api/invitations, GET /api/invitations/event/{id}, DELETE /api/invitations/{id}, POST /api/invitations/{id}/resend")
body("Admins can invite guests to specific events by email. Each invitation generates a unique "
     "QR code token. The guest receives an email with the invite. They present the QR code for "
     "check-in. Invitation statuses: pending, sent, verified, expired.")

h2("7.2 Workflow — Sending Invitations")
bullet("1. Open /admin/events/invitations/{eventId}.")
bullet("2. Click 'Invite Guest' → enter guest name + email.")
bullet("3. System creates a Guest record (or reuses existing) and an Invitation with a unique token.")
bullet("4. Email sent via SMTP with event details + QR code.")
bullet("5. Guest list displays all invitations with status, sent time, and actions.")

h2("7.3 Workflow — Check-in (/admin/check-in)")
bullet("1. Staff opens /admin/check-in on a device.")
bullet("2. Camera activates for QR scan (or manual token entry).")
bullet("3. Token is sent to POST /api/invitations/scan.")
bullet("4. Backend validates token (not expired, correct event), marks invitation as verified.")
bullet("5. Result displayed: green (success) or red (invalid/already used/expired).")

h2("7.4 Test Scenarios")
scenario_table([
    {"id":"TC-INV-01","description":"Send invitation to guest",
     "precondition":"Event exists; SMTP configured","steps":"1. Open invitations for event\n2. Enter guest email\n3. Send",
     "expected":"Invitation created; email delivered; appears in list as 'sent'.","status":"PASS"},
    {"id":"TC-INV-02","description":"Resend invitation",
     "precondition":"Invitation exists and not verified","steps":"1. Click Resend on invitation",
     "expected":"Email resent; sent_at updated.","status":"PASS"},
    {"id":"TC-INV-03","description":"Cancel/Delete invitation",
     "precondition":"Invitation exists","steps":"1. Click Delete on invitation",
     "expected":"Invitation removed; QR code invalidated.","status":"PASS"},
    {"id":"TC-INV-04","description":"Scan valid QR code",
     "precondition":"Valid invitation token","steps":"1. Open /admin/check-in\n2. Scan QR",
     "expected":"Green success message; invitation status = verified.","status":"PASS"},
    {"id":"TC-INV-05","description":"Scan already-used QR code",
     "precondition":"Invitation already verified","steps":"1. Scan same QR again",
     "expected":"Red error: already used.","status":"PASS"},
    {"id":"TC-INV-06","description":"Scan expired QR code",
     "precondition":"Invitation past expiry date","steps":"1. Scan expired QR",
     "expected":"Red error: invitation expired.","status":"PASS"},
    {"id":"TC-INV-07","description":"Public invite page with valid token",
     "precondition":"Valid token","steps":"1. Open /invite/{token} in browser",
     "expected":"Event details displayed with check-in status.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. CALENDAR MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Calendar Management Module")
h2("8.1 Overview")
kv("URL", "/admin/calender/calender_setup")
kv("Roles", "Admin, SuperAdmin")
kv("Backend Endpoints", "GET /api/events, GET /api/calendar-filter-settings, PUT /api/calendar-filter-settings")
body("Interactive calendar view for managing and previewing events by date. Supports advanced "
     "filtering by category, date range, location, price, and tags. Filter visibility can be "
     "toggled per-site by admins and persisted to the database.")

h2("8.2 Workflow")
bullet("1. Open calendar page → all published events loaded.")
bullet("2. Calendar renders a monthly grid; event dots shown on their dates.")
bullet("3. Clicking a date shows events for that day.")
bullet("4. Left panel shows filter toggles (Category, Date Range, Location, Price, Tags).")
bullet("5. Admin can save filter visibility preferences via PUT /api/calendar-filter-settings.")
bullet("6. These settings affect what filters are shown on the public-facing calendar.")

h2("8.3 Test Scenarios")
scenario_table([
    {"id":"TC-CAL-01","description":"View events on calendar",
     "precondition":"Published events exist","steps":"1. Navigate to /admin/calender/calender_setup",
     "expected":"Events shown as dots on their respective dates.","status":"PASS"},
    {"id":"TC-CAL-02","description":"Filter by category",
     "precondition":"Events with different categories","steps":"1. Select a category in filter panel",
     "expected":"Only events in that category shown.","status":"PASS"},
    {"id":"TC-CAL-03","description":"Save filter visibility settings",
     "precondition":"Admin user","steps":"1. Toggle 'Show Price' off\n2. Save settings",
     "expected":"Price filter hidden on public calendar; setting persisted.","status":"PASS"},
    {"id":"TC-CAL-04","description":"Navigate calendar months",
     "precondition":"Events in multiple months","steps":"1. Click next/previous month arrows",
     "expected":"Calendar navigates; events update for selected month.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. PUBLIC LANDING PAGE
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Public Landing Page")
h2("9.1 Overview")
kv("URL", "/  (root)")
kv("Roles", "Unauthenticated public users")
kv("Backend Endpoints", "GET /api/events (published only), GET /api/calendar-filter-settings")
body("The public-facing homepage displays all published tourism events with search and filter "
     "capabilities. No login required. Filter options shown depend on the admin's saved "
     "calendar filter settings.")

h2("9.2 Workflow")
bullet("1. Server-side load fetches all published events + filter settings.")
bullet("2. Hero section shows featured events or latest event.")
bullet("3. Event cards show title, date, category badge, location, price.")
bullet("4. Filters (if enabled by admin): Category, Date Range, Location, Price, Tags.")
bullet("5. Search box filters by title client-side.")
bullet("6. Clicking event card shows event detail.")

h2("9.3 Test Scenarios")
scenario_table([
    {"id":"TC-PUB-01","description":"Public page loads without login",
     "precondition":"No auth cookie","steps":"1. Open / in browser",
     "expected":"Events listed; no login required.","status":"PASS"},
    {"id":"TC-PUB-02","description":"Only published events visible",
     "precondition":"Mix of draft and published events","steps":"1. View public page",
     "expected":"Only published events shown.","status":"PASS"},
    {"id":"TC-PUB-03","description":"Search filters events",
     "precondition":"Multiple events","steps":"1. Type in search box",
     "expected":"Matching events shown; others hidden.","status":"PASS"},
    {"id":"TC-PUB-04","description":"Hidden filters not visible",
     "precondition":"Admin disabled Price filter","steps":"1. View public page",
     "expected":"Price filter not rendered.","status":"PASS"},
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  10. RBAC SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Role-Based Access Control (RBAC) Summary")
body("All admin pages check permissions before rendering actions or loading data. "
     "The sidebar nav is filtered per user. Each API endpoint validates both the JWT role "
     "and the specific permission flag.")

h2("10.1 Permission Matrix Reference")
COL_W2 = [Cm(4), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.8)]
t = doc.add_table(rows=1, cols=7)
t.style = 'Table Grid'
headers = ['Module / Nav Item', 'View', 'Create', 'Update', 'Delete', 'NeedsApproval', 'Default Role']
hdr_cells = t.rows[0].cells
for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
    shade_cell(cell, NAVY)
    cell.width = COL_W2[i]
    p  = cell.paragraphs[0]
    rn = p.add_run(hdr)
    rn.bold           = True
    rn.font.size      = Pt(9)
    rn.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Dashboard",           "✓","–","–","–","–","Admin/SuperAdmin"),
    ("Events Management",   "✓","✓","✓","✓","Optional","Admin/SuperAdmin"),
    ("Pending Approval",    "✓","✓","✓","✓","–","Admin/SuperAdmin"),
    ("Calendar Management", "✓","✓","✓","–","–","Admin/SuperAdmin"),
    ("Users Management",    "✓","✓","✓","✓","–","SuperAdmin only"),
    ("Categories Setup",    "✓","✓","✓","✓","–","Admin/SuperAdmin"),
    ("Check-in",            "✓","–","–","–","–","Admin/SuperAdmin"),
    ("Notifications",       "✓","–","–","✓","–","All"),
    ("Invitations",         "✓","✓","–","✓","–","Admin/SuperAdmin"),
]
for i, (mod, *vals) in enumerate(rows_data):
    row   = t.add_row()
    cells = row.cells
    bg    = LGRAY if i % 2 == 0 else WHITE
    for j, (cell, val) in enumerate(zip(cells, [mod]+vals)):
        shade_cell(cell, bg)
        set_cell_border(cell, color='C8D6E5')
        p  = cell.paragraphs[0]
        rn = p.add_run(val)
        rn.font.size      = Pt(9)
        rn.font.color.rgb = NAVY if j == 0 else (GREEN if val == '✓' else GRAY)
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

h2("10.2 JWT Role Values")
kv("superadmin", "Full access to all modules including Users Management")
kv("admin", "Access limited by permission matrix assigned by SuperAdmin")

h2("10.3 NeedsApproval Flag")
body("When NeedsApproval is enabled on the 'Events' nav item for a specific user, "
     "that user's event create/update actions trigger the approval workflow instead of "
     "directly publishing. This flag is per-user, per-nav-item.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  11. KNOWN LIMITATIONS / OUT OF SCOPE
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Known Limitations & Out-of-Scope Items")

h3("Currently Not Implemented / Stub Pages")
bullet("Client Portal Dashboard (/client/dashboard) — placeholder, no functionality yet.")
bullet("Admin Dashboard (/admin/dashboard) — test stub.")
bullet("Calendar Create sub-page — empty placeholder.")
bullet("Reports & Analytics — nav item exists but page not built.")
bullet("Department Setup — nav item exists but page not built.")

h3("Known Limitations")
bullet("Refresh token rotation implemented; but silent refresh on page load may show login flash on very slow connections.")
bullet("File uploads stored in local wwwroot/uploads — not cloud-backed; files lost on server reset.")
bullet("No pagination on events list — all events load at once (performance concern at scale).")
bullet("SignalR disconnects if browser tab is backgrounded for extended period; reconnects automatically on focus.")

h3("Browser Compatibility")
bullet("Tested on: Chrome 120+, Firefox 122+, Edge 120+")
bullet("QR scanner (camera API) requires HTTPS in production.")

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = "/home/notcool/Desktop/ntb-event/NTB_Event_QA_Documentation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
