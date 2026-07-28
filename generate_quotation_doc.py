"""Generate the NTB Event Platform project quotation as a DOCX file."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── colour palette (Nepal Tourism Board brand) ─────────────────────────────
TEAL        = RGBColor(0x1C, 0x5C, 0x6D)   # Himalayan Teal — primary
CRIMSON     = RGBColor(0xBD, 0x24, 0x2B)   # Crimson Red — accent / totals
MARIGOLD_HEX= "F8CE1C"                      # Marigold Yellow — fill accent
TEAL_HEX    = "1C5C6D"
BG_LIGHT    = "EAF2F3"                      # very light teal table header
TEXT_DARK   = RGBColor(0x26, 0x30, 0x38)
TEXT_MUTED  = RGBColor(0x6B, 0x72, 0x80)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# ─── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph_border_bottom(para, color="1C5C6D", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font_of(run, size_pt, bold=False, italic=False, color: RGBColor = None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = para.add_run(text)
    if level == 1:
        font_of(run, 14, bold=True, color=TEAL)
        add_paragraph_border_bottom(para)
    else:
        font_of(run, 11.5, bold=True, color=TEXT_DARK)
    return para


def body(doc: Document, text: str, color: RGBColor = None, italic=False, size=10.5, bold=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    font_of(run, size, italic=italic, bold=bold, color=color or TEXT_DARK)
    return para


def bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    font_of(run, 10.5, color=TEXT_DARK)
    return para


def no_border_table(rows, cols):
    pass


def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tblPr.append(borders)


# ─── document setup ──────────────────────────────────────────────────────────

def make_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


# ─── HEADER BANNER ───────────────────────────────────────────────────────────

def header_banner(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_none(table)
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(5)

    left = table.rows[0].cells[0]
    p = left.paragraphs[0]
    r = p.add_run("QUOTATION")
    font_of(r, 28, bold=True, color=TEAL)
    p2 = left.add_paragraph()
    r2 = p2.add_run("NTB Event Management Platform")
    font_of(r2, 12.5, color=TEXT_MUTED, italic=True)

    right = table.rows[0].cells[1]
    right.vertical_alignment = 1
    p3 = right.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run("Prepared for\n")
    font_of(r3, 9.5, color=TEXT_MUTED)
    r4 = p3.add_run("Nepal Tourism Board")
    font_of(r4, 13, bold=True, color=CRIMSON)

    # gold accent rule under the banner
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(14)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MARIGOLD_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── META + FROM/TO BLOCK ────────────────────────────────────────────────────

def meta_block(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders_none(table)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    # From
    from_cell = table.rows[0].cells[0]
    p = from_cell.paragraphs[0]
    r = p.add_run("FROM")
    font_of(r, 9, bold=True, color=TEAL)
    for line, bold in [
        ("Anjal Joshi", True),
        ("Full-Stack Software Developer", False),
        ("anjaljoshi6@gmail.com", False),
    ]:
        pp = from_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(line)
        font_of(rr, 10.5, bold=bold, color=TEXT_DARK)

    # To
    to_cell = table.rows[0].cells[1]
    p2 = to_cell.paragraphs[0]
    r2 = p2.add_run("BILL TO")
    font_of(r2, 9, bold=True, color=TEAL)
    for line, bold in [
        ("Nepal Tourism Board (NTB)", True),
        ("Bhrikutimandap, Kathmandu, Nepal", False),
        ("www.ntb.gov.np", False),
    ]:
        pp = to_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(line)
        font_of(rr, 10.5, bold=bold, color=TEXT_DARK)

    doc.add_paragraph()

    # quotation meta strip
    meta = doc.add_table(rows=1, cols=3)
    set_table_borders_none(meta)
    labels = [
        ("Quotation No.", "NTB-EVT-Q-2026-06"),
        ("Date", "June 19, 2026"),
        ("Valid Until", "July 19, 2026"),
    ]
    for i, (label, value) in enumerate(labels):
        cell = meta.rows[0].cells[i]
        set_cell_bg(cell, BG_LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(label.upper())
        font_of(r, 8, bold=True, color=TEXT_MUTED)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(value)
        font_of(r2, 11, bold=True, color=TEAL)

    doc.add_paragraph()


# ─── PROJECT SUMMARY ─────────────────────────────────────────────────────────

def project_summary(doc: Document):
    heading(doc, "Project Summary")
    body(doc,
         "Design, development and delivery of a complete digital event management platform for the "
         "Nepal Tourism Board, comprising a public event calendar / listing site and a secure admin "
         "back-office for managing events, guest invitations with QR-code check-in, role-based "
         "approvals, real-time notifications and reporting. The system is built on a .NET 8 (Clean "
         "Architecture) backend with a PostgreSQL database and a SvelteKit frontend styled to the "
         "official NTB brand identity.")


# ─── COST BREAKDOWN TABLE ────────────────────────────────────────────────────

LINE_ITEMS = [
    ("1", "Requirement Analysis & System Architecture Design", 8000),
    ("2", "Backend API Development (.NET 8, Clean Architecture, PostgreSQL)", 20000),
    ("3", "Frontend Development (SvelteKit, Responsive UI/UX)", 18000),
    ("4", "Admin Dashboard, Roles & Permission Management", 10000),
    ("5", "Event Management Module (Categories, Tags, Calendar, Approval Workflow)", 10000),
    ("6", "Guest Invitation, QR Code Generation & Check-in System", 10000),
    ("7", "Reports, Analytics & Real-Time Notifications", 8000),
    ("8", "Public Website (NTB-Branded Theme), Testing, Deployment & Documentation", 6000),
]


def cost_table(doc: Document):
    heading(doc, "Cost Breakdown")

    table = doc.add_table(rows=1 + len(LINE_ITEMS) + 1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(1.4)
    table.columns[1].width = Cm(10.2)
    table.columns[2].width = Cm(3.4)

    # header row
    hdr = table.rows[0]
    headers = ["S.N.", "Description", "Amount (NPR)"]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, TEAL_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        font_of(r, 10, bold=True, color=WHITE)

    # data rows
    for ri, (sn, desc, amount) in enumerate(LINE_ITEMS):
        row = table.rows[ri + 1]
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_of(p0.add_run(sn), 10, color=TEXT_DARK)

        c1 = row.cells[1]
        font_of(c1.paragraphs[0].add_run(desc), 10, color=TEXT_DARK)

        c2 = row.cells[2]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        font_of(p2.add_run(f"{amount:,.2f}"), 10, color=TEXT_DARK)

        if ri % 2 == 1:
            for c in row.cells:
                set_cell_bg(c, "F7F7F7")

    # total row
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[1])
    tc = total_row.cells[0]
    set_cell_bg(tc, BG_LIGHT)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p.add_run("TOTAL PROJECT COST"), 11, bold=True, color=TEAL)

    tc2 = total_row.cells[2]
    set_cell_bg(tc2, BG_LIGHT)
    p2 = tc2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_of(p2.add_run("NPR 90,000.00"), 12, bold=True, color=CRIMSON)

    doc.add_paragraph()

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    r = note.add_run("Amount in words: ")
    font_of(r, 10, bold=True, color=TEXT_DARK)
    r2 = note.add_run("Nepalese Rupees Ninety Thousand Only.")
    font_of(r2, 10, italic=True, color=TEXT_MUTED)


# ─── SCOPE OF WORK ────────────────────────────────────────────────────────────

def scope_of_work(doc: Document):
    heading(doc, "Scope of Work")
    items = [
        "Public-facing event calendar and listing pages aligned with the official NTB brand identity.",
        "Secure admin panel with JWT authentication and role-based access control (SuperAdmin / Admin).",
        "Full event lifecycle management: create, edit, categorize, tag and publish events, with an "
        "optional approval workflow for non-SuperAdmin users.",
        "Guest invitation system with unique QR-code generation, email delivery and on-site QR check-in / verification.",
        "Real-time in-app notifications via SignalR for approval requests and check-in activity.",
        "Admin reports and analytics dashboard for event and attendance insights.",
        "Responsive, mobile-friendly design across all public and admin pages.",
        "Deployment-ready build with environment configuration, basic documentation and handover support.",
    ]
    for it in items:
        bullet(doc, it)


# ─── TERMS & CONDITIONS ───────────────────────────────────────────────────────

def terms(doc: Document):
    heading(doc, "Terms & Conditions")
    items = [
        "This quotation is valid for 30 days from the date of issue.",
        "Payment Terms: 50% advance to commence work, 50% upon final delivery and acceptance.",
        "The quoted amount covers the scope of work listed above. Additional features or major scope "
        "changes requested after project kickoff will be quoted separately.",
        "Estimated delivery timeline will be confirmed upon advance payment and finalization of requirements.",
        "Hosting, domain, SMTP/email service and any third-party service costs are not included in this quotation.",
        "A reasonable warranty/bug-fix period will be provided after delivery; details to be agreed upon at project start.",
    ]
    for it in items:
        bullet(doc, it)


# ─── SIGNATURE BLOCK ──────────────────────────────────────────────────────────

def signature_block(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    set_table_borders_none(table)

    labels_top = ["Prepared By", "Accepted By (Nepal Tourism Board)"]
    for i, label in enumerate(labels_top):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:bottom")
        # simple underline using border on paragraph
        p.paragraph_format.space_before = Pt(30)
        border = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "9CA3AF")
        border.append(b)
        pPr.append(border)

    for i, label in enumerate(labels_top):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(label)
        font_of(r, 9.5, color=TEXT_MUTED, italic=True)


def footer_note(doc: Document):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Thank you for the opportunity to work with the Nepal Tourism Board.")
    font_of(r, 10, italic=True, color=TEXT_MUTED)


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    doc = make_doc()
    header_banner(doc)
    meta_block(doc)
    project_summary(doc)
    scope_of_work(doc)
    cost_table(doc)
    terms(doc)
    signature_block(doc)
    footer_note(doc)

    out_path = "/home/notcool/Desktop/ntb-event/NTB_Event_Platform_Quotation.docx"
    doc.save(out_path)
    print(f"Saved → {out_path}")


if __name__ == "__main__":
    main()
